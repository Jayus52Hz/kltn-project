from __future__ import annotations

from pathlib import Path
import shutil

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
REPORTS_DIR = ROOT / "docs" / "reports"

SRC = REPORTS_DIR / "Report KLTN - 22133056 - Nguyen Quoc Thinh - 08-06 - chapter1 citations notes.docx"
OUT = REPORTS_DIR / "Report KLTN - 22133056 - Nguyen Quoc Thinh - 08-06 - v3 chuong1 ly thuyet chuyen sau.docx"


def clear_paragraph(paragraph):
    p = paragraph._p
    for child in list(p):
        if child.tag.endswith("}pPr"):
            continue
        p.remove(child)


def insert_paragraph_after(paragraph, text: str):
    new_p = paragraph._p.__copy__()
    wrapper = type(paragraph)(new_p, paragraph._parent)
    clear_paragraph(wrapper)
    wrapper.style = paragraph.style
    wrapper.add_run(text)
    paragraph._p.addnext(new_p)
    return wrapper


def insert_block_after_anchor(doc: Document, anchor_contains: str, paragraphs: list[str]):
    if any(paragraphs[0] in p.text for p in doc.paragraphs):
        return False
    for paragraph in doc.paragraphs:
        if anchor_contains in paragraph.text:
            cursor = paragraph
            for text in reversed(paragraphs):
                # addnext inserts immediately after the anchor; reversing preserves order
                cursor = insert_paragraph_after(paragraph, text)
            return True
    raise ValueError(f"Anchor not found: {anchor_contains}")


def main():
    if OUT.exists():
        OUT.unlink()
    shutil.copy2(SRC, OUT)
    doc = Document(str(OUT))

    expansions: list[tuple[str, list[str]]] = [
        (
            "Đối với hệ thống AGI Telesales, dữ liệu không chỉ là các bản ghi giao dịch",
            [
                "Về bản chất, ba mô hình Data Warehouse, Data Lake và Data Lakehouse khác nhau ở thời điểm áp đặt lược đồ, mức độ quản trị metadata và cách phục vụ workload phân tích. Data Warehouse ưu tiên dữ liệu đã chuẩn hóa và mô hình hóa trước khi ghi, nên phù hợp với báo cáo ổn định nhưng kém linh hoạt với JSON, transcript và log. Data Lake ưu tiên lưu trữ dữ liệu thô ở nhiều định dạng, nhưng nếu không có metadata, phân quyền và quy tắc chất lượng thì dễ mất khả năng kiểm soát. Lakehouse được xem như lớp trung gian kết hợp lưu trữ mở của Data Lake với quản trị bảng, transaction, schema enforcement và tối ưu truy vấn của Data Warehouse [11].",
                "Trong đề tài, phần lý thuyết này không chỉ dùng để định nghĩa khái niệm. Nó giải thích vì sao kiến trúc phải giữ raw document ở Bronze, chuẩn hóa và bảo vệ dữ liệu ở Silver, rồi mới mô hình hóa Star Schema ở Gold. Nếu bỏ qua Lakehouse và chỉ dùng MongoDB + dashboard, hệ thống sẽ khó tái xử lý transcript, khó truy vết thay đổi và dễ tạo tải phân tích trực tiếp lên OLTP. Nếu chỉ dùng Data Lake, dữ liệu có thể được lưu nhưng thiếu lớp bảng và metadata để phục vụ BI một cách có kiểm soát.",
            ],
        ),
        (
            "Change Data Capture là kỹ thuật phát hiện và truyền tải các thay đổi",
            [
                "Trong MongoDB, Change Streams cung cấp luồng sự kiện thay đổi ở mức collection, database hoặc deployment mà ứng dụng có thể theo dõi mà không phải tự tail oplog thủ công [30]. Cơ chế này chỉ phát ra các thay đổi đã được ghi nhận bền vững bởi đa số data-bearing members trong replica set, vì vậy nó phù hợp hơn cho pipeline downstream so với cách polling collection định kỳ. Mỗi event thường mang thông tin loại thao tác, namespace, document key, thời điểm cụm ghi nhận sự kiện và resume token để consumer có thể tiếp tục đọc sau gián đoạn.",
                "Đối với AGI Telesales, Change Streams giải quyết đúng điểm nghẽn của dữ liệu vận hành: cuộc gọi và trạng thái khách hàng thay đổi liên tục, nhưng pipeline phân tích không nên truy vấn lặp lại toàn bộ MongoDB. Tuy nhiên, Change Streams không tự động biến MongoDB thành hệ thống phân tích thời gian thực hoàn chỉnh. Pipeline vẫn phải xử lý các vấn đề như thời hạn giữ oplog, resume token hết hiệu lực, update event không có full document nếu cấu hình không phù hợp, và khả năng duplicate/replay khi consumer khởi động lại.",
            ],
        ),
        (
            "Debezium là nền tảng CDC mã nguồn mở",
            [
                "Debezium hoạt động như một lớp source connector trong Kafka Connect: connector đọc thay đổi từ hệ quản trị nguồn, đóng gói thành event có cấu trúc và đẩy vào Kafka topic [4]. Với MongoDB connector, Debezium sử dụng Change Streams để nhận các thao tác insert, update, delete; phần payload thường chứa thông tin `op`, `source`, timestamp, document key và dữ liệu trước/sau tùy chế độ capture. Cấu hình `change_streams_update_full` có ý nghĩa quan trọng vì update event có thể kèm trạng thái document đầy đủ trong trường `after`, giúp tầng Bronze lưu được raw document đủ ngữ cảnh thay vì chỉ lưu delta.",
                "Trong đề tài, Debezium không chỉ là công cụ copy dữ liệu. Nó là ranh giới kỹ thuật giữa OLTP và pipeline phân tích: MongoDB vẫn tập trung cho nghiệp vụ ghi/đọc vận hành, còn Debezium chịu trách nhiệm biến thay đổi thành event. Trade-off của lựa chọn này là pipeline phải hiểu Debezium envelope, phải kiểm soát connector status, và phải thiết kế Bronze job sao cho idempotent khi event được đọc lại hoặc khi connector khởi động lại.",
            ],
        ),
        (
            "Apache Kafka là nền tảng event streaming phân tán",
            [
                "Kafka tổ chức dữ liệu thành event được ghi vào topic, mỗi topic có thể chia thành nhiều partition để mở rộng throughput. Thứ tự event chỉ được đảm bảo trong phạm vi một partition; khi event có cùng key, Kafka có thể ghi chúng vào cùng partition để giữ thứ tự tương đối cho key đó [5]. Offset trong partition cho phép consumer biết đã đọc đến đâu, còn cơ chế retention cho phép downstream đọc lại dữ liệu trong một khoảng thời gian nhất định.",
                "Với pipeline CDC của đề tài, Kafka đóng vai trò buffer và replay log giữa Debezium và Spark. Điều này làm giảm coupling: Debezium chỉ cần ghi event vào topic, Spark Bronze job có thể đọc theo nhịp xử lý riêng, còn khi logic parse thay đổi có thể đọc lại event trong phạm vi retention. Đổi lại, Kafka không tự bảo đảm chất lượng dữ liệu nghiệp vụ; các vấn đề như duplicate event, ordering giữa nhiều collection, retention quá ngắn hoặc topic partition chưa hợp lý vẫn phải được xử lý ở tầng pipeline.",
            ],
        ),
        (
            "Apache Spark là engine xử lý dữ liệu phân tán",
            [
                "Spark cung cấp mô hình xử lý dựa trên DataFrame/Dataset và Spark SQL, cho phép cùng một engine thực hiện batch transformation, truy vấn SQL và một số workload streaming có khả năng chịu lỗi [6]. Về mặt thực thi, Spark chia dữ liệu thành partition, lập kế hoạch logical/physical plan, rồi phân phối task xuống executor. Cách tiếp cận này phù hợp với pipeline Lakehouse vì cùng một job có thể đọc nhiều nguồn, chuẩn hóa schema, join bảng, tính cột dẫn xuất và ghi ra định dạng bảng phân tích.",
                "Trong đề tài, Spark được dùng như compute layer cho Bronze, Silver và Gold. Bronze job đọc event từ Kafka và ghi raw document vào Iceberg; Silver job parse JSON, ép kiểu, masking PII, deduplicate và chạy NLP inference; Gold job tạo dimension/fact phục vụ BI. Trade-off là Spark đem lại khả năng mở rộng và API mạnh nhưng làm tăng độ phức tạp runtime: phải đồng bộ Python dependency, cấu hình memory/partition, kiểm soát UDF/Pandas UDF, và tránh đưa mô hình quá nặng vào batch job CPU-only.",
            ],
        ),
        (
            "Apache Iceberg là open table format cho bảng phân tích lớn",
            [
                "Iceberg không phải là database độc lập mà là table format quản lý metadata cho dữ liệu nằm trên object/file storage. Một bảng Iceberg gồm metadata file, snapshot, manifest list, manifest file và data/delete files. Snapshot cho phép bảng có lịch sử phiên bản; manifest giúp engine biết file nào thuộc snapshot nào; schema evolution và partition evolution cho phép thay đổi lược đồ hoặc cách phân vùng mà không phải rewrite toàn bộ dữ liệu [7].",
                "Vai trò của Iceberg trong đề tài là biến MinIO từ nơi chứa file thành Lakehouse table layer có thể truy vấn và cập nhật có kiểm soát. Tầng Bronze cần lưu raw document để audit; tầng Silver cần cập nhật theo khóa nghiệp vụ sau deduplicate; tầng Gold cần MERGE/overwrite các bảng dimension và fact. Nếu chỉ ghi Parquet rời rạc, pipeline sẽ khó quản lý snapshot, schema, partition và upsert. Trade-off là hệ thống phải cấu hình catalog/warehouse đúng, quản lý metadata, và hiểu rằng performance phụ thuộc vào số lượng file, partition layout và quá trình compaction.",
            ],
        ),
        (
            "Object storage lưu dữ liệu dưới dạng object",
            [
                "Object storage tổ chức dữ liệu theo bucket và object key, khác với file system truyền thống ở chỗ nó tối ưu cho lưu trữ dữ liệu bất biến/quy mô lớn và truy cập qua API. MinIO cung cấp API tương thích S3 và có thể chạy bằng container, vì vậy phù hợp để mô phỏng một lớp storage nội bộ cho Lakehouse trong môi trường local [9], [33]. Đối với Iceberg, object storage là nơi đặt data files và metadata files, còn Spark/Iceberg đảm nhiệm việc diễn giải bảng.",
                "Trong phạm vi prototype, MinIO giúp tách compute và storage: Spark có thể dừng/chạy lại mà dữ liệu Iceberg vẫn nằm trong warehouse. Cách này gần với mô hình cloud data lake nhưng vẫn chạy được local. Hạn chế là MinIO một node trong Docker Compose chưa đại diện cho production storage có replication, erasure coding, access policy đầy đủ và monitoring. Vì vậy báo cáo cần xem MinIO local như môi trường mô phỏng kiến trúc, không phải chứng minh độ sẵn sàng production.",
            ],
        ),
        (
            "Apache Airflow điều phối workflow theo DAG",
            [
                "Airflow mô hình hóa workflow bằng DAG, trong đó mỗi task là một đơn vị thực thi và các cạnh biểu diễn quan hệ phụ thuộc [8]. DAG không xử lý dữ liệu thay cho Spark hay Kafka; nó quản lý thứ tự chạy, retry, timeout, log, trạng thái task và khả năng trigger thủ công/theo lịch. Đây là khác biệt quan trọng: Airflow là orchestration layer, còn logic xử lý dữ liệu vẫn nằm trong Spark jobs, connector hoặc script chuyên trách.",
                "Trong đề tài, Airflow giúp biến nhiều bước rời rạc thành một pipeline có thể vận hành: chờ Debezium connector sẵn sàng, chạy Bronze ingestion, chạy Silver ETL, tạo Gold Star Schema và đồng bộ BigQuery. Điều này làm pipeline dễ demo, dễ kiểm tra lỗi và dễ chạy lại. Trade-off là Airflow LocalExecutor trong Docker Compose chưa đại diện cho production orchestration; nếu triển khai thật cần quản lý secrets, schedule, SLA, retry policy, alerting và phân quyền truy cập UI.",
            ],
        ),
        (
            "Apache Superset là nền tảng BI mã nguồn mở",
            [
                "Superset là lớp visualization/BI phía người dùng, có thể kết nối tới các SQL-speaking datastore thông qua DB-API driver và SQLAlchemy dialect [10]. Nó hỗ trợ dataset, chart, dashboard, filter và quyền truy cập. Về mặt kiến trúc, Superset không nên trực tiếp thay thế tầng Gold; nó tiêu thụ dữ liệu đã được mô hình hóa và tối ưu trước đó. Nếu dashboard truy vấn trực tiếp vào nguồn OLTP hoặc bảng raw, phần BI sẽ dễ chậm, khó bảo mật và khó giải thích.",
                "Trong đề tài, Superset được đặt sau Gold/BigQuery để minh họa KPI telesales như total calls, success rate, outcome breakdown, lead source performance và product performance. Cách đặt này giúp người dùng nghiệp vụ chỉ nhìn thấy dữ liệu đã qua masking và tổng hợp, thay vì truy cập raw transcript hoặc PII. Hạn chế là dashboard hiện tại chủ yếu chứng minh khả năng phục vụ phân tích; để chuyên nghiệp hơn cần bổ sung quyền truy cập theo role, data freshness indicator và dashboard kỹ thuật cho trạng thái pipeline.",
            ],
        ),
        (
            "BigQuery là kho dữ liệu phân tích dạng serverless",
            [
                "BigQuery là kho dữ liệu phân tích serverless trên Google Cloud, tách người dùng khỏi việc quản trị cụm máy chủ và tập trung vào lưu trữ/truy vấn dữ liệu phân tích [31]. Trong kiến trúc hybrid, BigQuery phù hợp làm serving warehouse cho dữ liệu đã sạch vì nó hỗ trợ truy vấn SQL, view, tích hợp BI và mở rộng theo workload phân tích. Tuy nhiên, việc đưa dữ liệu lên BigQuery cũng tạo ranh giới bảo mật mới: dữ liệu publish phải được tối giản, đã loại PII trực tiếp và có chính sách quyền rõ ràng.",
                "Looker Studio đóng vai trò lớp báo cáo tương tác trên dữ liệu đã kết nối, trong đó report có thể thêm data source và tạo chart/table/filter cho người dùng nghiệp vụ [32]. Với đề tài, BigQuery/Looker Studio được xem là hướng BI cloud-native, còn Superset là hướng BI mã nguồn mở/local. Việc trình bày cả hai giúp làm rõ tính hybrid: dữ liệu thô và nhạy cảm ở lại Lakehouse local, dữ liệu Gold đã kiểm soát có thể phục vụ dashboard nội bộ hoặc cloud tùy yêu cầu triển khai.",
            ],
        ),
        (
            "Các công nghệ được chọn theo tiêu chí",
            [
                "Docker Compose là công cụ định nghĩa và chạy ứng dụng nhiều container bằng file cấu hình, phù hợp để đóng gói một môi trường demo gồm nhiều service phụ thuộc nhau [33]. Trong đồ án, Compose giúp tái lập stack gồm MongoDB, Debezium, Kafka, Spark, MinIO, Airflow và Superset trên một máy local. Giá trị lý thuyết ở đây là khả năng mô hình hóa dependency và network giữa các service, không phải chứng minh tính sẵn sàng cao của production cluster.",
                "Khi đánh giá bộ công cụ, cần phân biệt rõ vai trò từng lớp: MongoDB là OLTP/source, Debezium/Kafka là ingestion/change transport, Spark là compute, Iceberg/MinIO là Lakehouse storage/table layer, Airflow là orchestration, BigQuery/Superset/Looker Studio là serving/BI. Sự chuyên nghiệp của kiến trúc nằm ở việc mỗi công cụ được đặt đúng vai trò, có dữ liệu đầu vào/đầu ra rõ ràng, có giới hạn được nêu trước và có bằng chứng kiểm thử tương ứng ở Chương 4.",
            ],
        ),
    ]

    for anchor, paragraphs in expansions:
        insert_block_after_anchor(doc, anchor, paragraphs)

    doc.save(str(OUT))
    print(OUT)


if __name__ == "__main__":
    main()
