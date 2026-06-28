# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document


SRC = Path(
    "docs/reports/Report KLTN - 22133056 - Nguyen Quoc Thinh - 20-06 - v5 airflow staged pipeline test.docx"
)
OUT = Path(
    "docs/reports/Report KLTN - 22133056 - Nguyen Quoc Thinh - 24-06 - v6 code report consistency.docx"
)


def set_cell(cell, text):
    cell.text = text


def replace_paragraph(doc, needle, replacement):
    hits = 0
    for para in doc.paragraphs:
        if needle in para.text:
            para.text = replacement
            hits += 1
    return hits


def replace_in_paragraphs(doc, old, new):
    hits = 0
    for para in doc.paragraphs:
        if old in para.text:
            para.text = para.text.replace(old, new)
            hits += 1
    return hits


def replace_in_tables(doc, old, new):
    hits = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if old in cell.text:
                    cell.text = cell.text.replace(old, new)
                    hits += 1
    return hits


doc = Document(SRC)
changes = []

# Staged Airflow DAG evidence and architecture wording.
changes.append(
    (
        "airflow_placeholder",
        replace_paragraph(
            doc,
            "[MINH CHỨNG HÌNH - Apache Airflow: Ảnh chụp DAG telesales_lakehouse_pipeline với các task wait_for_debezium_connector -> bronze_cdc_ingestion -> silver_etl -> gold_star_schema -> bq_sync_gold.]",
            "[MINH CHỨNG HÌNH - Apache Airflow: Ảnh chụp DAG telesales_lakehouse_pipeline sau khi tách stage, gồm wait_for_debezium_connector, các task Bronze/Silver riêng cho cust_dataset, offer_dataset, call_logs_dataset, các task Gold dim_customer/dim_offer/dim_date/fact_telesales_calls, nhánh callcenteren_external bronze/silver/gold và bq_sync_gold.]",
        ),
    )
)
changes.append(
    (
        "old_dag_chain_code_block",
        replace_paragraph(
            doc,
            "wait_for_debezium >> bronze_cdc_ingestion >> silver_etl >> gold_star_schema >> bq_sync_gold",
            "wait_for_debezium_connector >> [cust_dataset.bronze, offer_dataset.bronze, call_logs_dataset.bronze] >> [cust_dataset.silver, offer_dataset.silver, call_logs_dataset.silver] >> [dim_customer, dim_offer, dim_date] >> fact_telesales_calls >> bq_sync_gold; nhánh callcenteren_external.bronze >> callcenteren_external.silver >> callcenteren_external.gold chạy trong cùng DAG để phục vụ so sánh multi-source.",
        ),
    )
)
changes.append(
    (
        "airflow_pipeline_paragraph",
        replace_paragraph(
            doc,
            "Airflow DAG telesales_lakehouse_pipeline điều phối pipeline theo thứ tự wait_for_debezium_connector, bronze_cdc_ingestion, silver_etl, gold_star_schema và bq_sync_gold. DAG có schedule_interval 0 2 * * * và catchup=false, đồng thời có thể trigger thủ công qua Airflow UI. Bronze vốn là Spark Structured Streaming chạy liên tục, nhưng trong Airflow được cấu hình TRIGGER_ONCE=true để task có điểm kết thúc và phù hợp với mô hình batch kiểm thử.",
            "Airflow DAG telesales_lakehouse_pipeline điều phối pipeline theo mô hình staged task group: wait_for_debezium_connector đứng đầu, sau đó là các task Bronze/Silver riêng cho cust_dataset, offer_dataset và call_logs_dataset; tầng Gold tách thành dim_customer, dim_offer, dim_date và fact_telesales_calls; nhánh callcenteren_external có Bronze/Silver/Gold riêng trước bước bq_sync_gold. DAG có schedule_interval 0 2 * * * và catchup=false, đồng thời có thể trigger thủ công qua Airflow UI. Bronze được cấu hình TRIGGER_ONCE=true để task streaming có điểm kết thúc và phù hợp với mô hình batch kiểm thử.",
        ),
    )
)
changes.append(
    (
        "airflow_pipeline_paragraph_actual",
        replace_paragraph(
            doc,
            "Airflow DAG telesales_lakehouse_pipeline điều phối pipeline theo thứ tự wait_for_debezium_connector, bronze_cdc_ingestion, silver_etl, gold_star_schema và bq_sync_gold. DAG có schedule_interval 0 2 * * * và catchup=false, đồng thời có thể trigger thủ công qua Airflow UI. Bronze vốn là Spark Structured Streaming job, nhưng trong Airflow được chạy với TRIGGER_ONCE=true để xử lý hết Kafka offsets hiện có rồi dừng, phù hợp với mô hình task batch.",
            "Airflow DAG telesales_lakehouse_pipeline điều phối pipeline theo mô hình staged task group: wait_for_debezium_connector đứng đầu, sau đó là các task Bronze/Silver riêng cho cust_dataset, offer_dataset và call_logs_dataset; tầng Gold tách thành dim_customer, dim_offer, dim_date và fact_telesales_calls; nhánh callcenteren_external có Bronze/Silver/Gold riêng trước bước bq_sync_gold. DAG có schedule_interval 0 2 * * * và catchup=false, đồng thời có thể trigger thủ công qua Airflow UI. Bronze được cấu hình TRIGGER_ONCE=true để task streaming có điểm kết thúc và phù hợp với mô hình batch kiểm thử.",
        ),
    )
)
changes.append(
    (
        "merge_idempotency_task_names",
        replace_paragraph(
            doc,
            "Về mặt báo cáo, cần giải thích MERGE INTO không chỉ là chi tiết code mà là cơ chế đảm bảo bảng phân tích không bị append duplicate sau mỗi lần chạy. Nếu Airflow chạy lại silver_etl hoặc gold_star_schema, dữ liệu được update theo khóa thay vì nhân đôi. Đây là nền tảng cho tính idempotent của pipeline.",
            "Về mặt báo cáo, cần giải thích MERGE INTO không chỉ là chi tiết code mà là cơ chế đảm bảo bảng phân tích không bị append duplicate sau mỗi lần chạy. Nếu Airflow chạy lại các task Silver hoặc Gold theo entity/stage, dữ liệu được update theo khóa thay vì nhân đôi. Đây là nền tảng cho tính idempotent của pipeline.",
        ),
    )
)

# PII / publishing policy wording: code uses an explicit blocklist, not a BQ_INCLUDE_PII flag.
changes.append(
    (
        "pii_source_policy",
        replace_paragraph(
            doc,
            "Các trường PII như full_name, phone_number, national_id và address được giữ ở dữ liệu nguồn để mô phỏng rủi ro thực tế. Tầng Silver chịu trách nhiệm che giấu phone_number và national_id, còn khi đồng bộ BigQuery có cấu hình drop full_name và address nếu BQ_INCLUDE_PII=false. Việc giữ PII ở nguồn như vậy giúp kiểm thử được pipeline bảo vệ dữ liệu thay vì giả định dữ liệu đầu vào đã sạch.",
            "Các trường PII như full_name, phone_number, national_id và address được giữ ở dữ liệu nguồn để mô phỏng rủi ro thực tế. Tầng Silver chịu trách nhiệm che giấu phone_number và national_id; khi đồng bộ sang BigQuery, bq_sync_job.py sử dụng BLOCKED_BIGQUERY_COLUMNS để loại call_transcript, các cột nhãn trung gian/legacy, full_name và address trước khi publish. Việc giữ PII ở nguồn như vậy giúp kiểm thử được pipeline bảo vệ dữ liệu thay vì giả định dữ liệu đầu vào đã sạch.",
        ),
    )
)
changes.append(
    (
        "security_requirement",
        replace_paragraph(
            doc,
            "Dữ liệu telesales chứa PII như tên, số điện thoại, national_id và địa chỉ. Thiết kế hệ thống phải đảm bảo analyst không cần truy cập toàn bộ PII để phân tích KPI. Vì vậy phone_number và national_id được mask ở Silver, còn full_name và address bị loại khi sync sang BigQuery nếu BQ_INCLUDE_PII=false. Đây là điểm khác biệt quan trọng giữa Lakehouse local dùng cho xử lý kỹ thuật và serving dataset dùng cho BI.",
            "Dữ liệu telesales chứa PII như tên, số điện thoại, national_id và địa chỉ. Thiết kế hệ thống phải đảm bảo analyst không cần truy cập toàn bộ PII để phân tích KPI. Vì vậy phone_number và national_id được mask ở Silver, còn full_name và address bị loại trong bq_sync_job.py bằng BLOCKED_BIGQUERY_COLUMNS trước khi ghi BigQuery. Đây là điểm khác biệt quan trọng giữa Lakehouse local dùng cho xử lý kỹ thuật và serving dataset dùng cho BI.",
        ),
    )
)
changes.append(
    (
        "bq_sync_design",
        replace_paragraph(
            doc,
            "Thiết kế bq_sync_job.py có biến BQ_INCLUDE_PII. Khi false, dim_customer bị drop full_name và address trước khi ghi BigQuery. Đây là một điểm thiết kế quan trọng cần đưa vào báo cáo vì nó chứng minh security không chỉ nằm ở lý thuyết mà đã được hiện thực trong code.",
            "Thiết kế bq_sync_job.py có hằng BLOCKED_BIGQUERY_COLUMNS. Trước khi ghi BigQuery, job tự động drop các cột nhạy cảm hoặc không phục vụ phân tích như call_transcript, call_code_original, call_code_predicted, full_name và address nếu chúng tồn tại trong DataFrame. Đây là một điểm thiết kế quan trọng vì security không chỉ nằm ở lý thuyết mà đã được hiện thực trong code publish.",
        ),
    )
)
changes.append(
    (
        "bq_sync_optional_tables",
        replace_paragraph(
            doc,
            "bq_sync_job.py đọc bốn bảng Gold từ Iceberg và ghi sang BigQuery dataset kltn0710. Các bảng gồm dim_customer, dim_offer, dim_date và fact_telesales_calls. Nếu BQ_INCLUDE_PII=false, dim_customer sẽ drop full_name và address trước khi ghi. Đây là bước bảo vệ dữ liệu ở lớp publish, đặc biệt khi BigQuery hoặc dashboard có thể được chia sẻ cho người dùng phân tích.",
            "bq_sync_job.py đọc bốn bảng Gold chính từ Iceberg và ghi sang BigQuery dataset kltn0710, gồm dim_customer, dim_offer, dim_date và fact_telesales_calls. Job cũng đọc các bảng so sánh CallCenterEN nếu tồn tại. Trước khi ghi, BLOCKED_BIGQUERY_COLUMNS loại các cột nhạy cảm hoặc không phục vụ BI như call_transcript, các cột nhãn legacy, full_name và address. Đây là bước bảo vệ dữ liệu ở lớp publish, đặc biệt khi BigQuery hoặc dashboard có thể được chia sẻ cho người dùng phân tích.",
        ),
    )
)
changes.append(
    (
        "silver_gold_label_policy",
        replace_paragraph(
            doc,
            "Sau inference, Silver giữ call_transcript để phục vụ kiểm tra kỹ thuật và khả năng tái xử lý cục bộ. Gold và BigQuery không nhận transcript, không nhận call_code_original và không nhận call_code_predicted. Các tầng phân tích chỉ giữ call_code do mô hình sinh ra, cùng các cờ nghiệp vụ được suy ra từ call_code.",
            "Sau inference, Silver giữ call_transcript để phục vụ kiểm tra kỹ thuật và khả năng tái xử lý cục bộ. Gold và BigQuery không nhận transcript; bq_sync_job.py cũng loại các cột nhãn trung gian/legacy nếu chúng còn tồn tại. Các tầng phân tích chỉ giữ call_code do mô hình sinh ra, cùng các cờ nghiệp vụ được suy ra từ call_code.",
        ),
    )
)
changes.append(
    (
        "silver_gold_label_policy_actual",
        replace_paragraph(
            doc,
            "Sau inference, Silver giữ call_transcript để phục vụ kiểm tra kỹ thuật và khả năng tái xử lý cục bộ. Gold và BigQuery không nhận transcript, không nhận call_code_original và không nhận call_code_predicted. Các tầng phân tích chỉ giữ call_code do mô hình sinh ra, cùng các cờ nghiệp vụ được suy ra từ call_code như has_successful_sale, has_hard_rejection, has_soft_rejection, has_do_not_call và has_objection.",
            "Sau inference, Silver giữ call_transcript để phục vụ kiểm tra kỹ thuật và khả năng tái xử lý cục bộ. Gold và BigQuery không nhận transcript; bq_sync_job.py cũng loại các cột nhãn trung gian/legacy nếu chúng còn tồn tại. Các tầng phân tích chỉ giữ call_code do mô hình sinh ra, cùng các cờ nghiệp vụ được suy ra từ call_code như has_successful_sale, has_hard_rejection, has_soft_rejection, has_do_not_call và has_objection.",
        ),
    )
)
changes.append(
    (
        "rebuild_label_policy",
        replace_paragraph(
            doc,
            "Sau full rebuild, Silver call_logs có call_code là nhãn do mô hình sinh và vẫn giữ call_transcript để phục vụ inference nội bộ. Gold fact_telesales_calls có call_code và các cờ nghiệp vụ, nhưng không có call_transcript, call_code_original hoặc call_code_predicted. BigQuery sync cũng loại các cột nhạy cảm hoặc cột không phục vụ BI thông qua blocklist:",
            "Sau full rebuild, Silver call_logs có call_code là nhãn do mô hình sinh và vẫn giữ call_transcript để phục vụ inference nội bộ. Gold fact_telesales_calls có call_code và các cờ nghiệp vụ, nhưng không có call_transcript; BigQuery sync tiếp tục loại các cột nhạy cảm hoặc cột nhãn trung gian/legacy thông qua blocklist:",
        ),
    )
)
changes.append(
    (
        "rebuild_label_policy_actual",
        replace_paragraph(
            doc,
            "Sau full rebuild, Silver call_logs có call_code là nhãn do mô hình sinh và vẫn giữ call_transcript để phục vụ inference nội bộ. Gold fact_telesales_calls có call_code và các cờ nghiệp vụ, nhưng không có call_transcript, call_code_original hoặc call_code_predicted. BigQuery sync cũng loại các cột nhạy cảm trước khi ghi, bảo đảm transcript không đi vào lớp serving phân tích.",
            "Sau full rebuild, Silver call_logs có call_code là nhãn do mô hình sinh và vẫn giữ call_transcript để phục vụ inference nội bộ. Gold fact_telesales_calls có call_code và các cờ nghiệp vụ, nhưng không có call_transcript. BigQuery sync cũng loại các cột nhạy cảm hoặc cột nhãn trung gian/legacy trước khi ghi, bảo đảm transcript không đi vào lớp serving phân tích.",
        ),
    )
)
changes.append(
    (
        "callcenteren_serving_limit",
        replace_paragraph(
            doc,
            "Về hạ tầng serving, nhánh CallCenterEN đã được ghi vào Iceberg/MinIO và tạo bảng Gold so sánh, nhưng BigQuery sync cho các bảng so sánh mới chỉ được bổ sung ở mức optional code path, chưa chạy xác thực đầy đủ trên môi trường cloud. Đây là giới hạn vận hành cần ghi nhận nếu mở rộng báo cáo sang phần demo dashboard cloud.",
            "Về hạ tầng serving, nhánh CallCenterEN đã được ghi vào Iceberg/MinIO và bq_sync_job.py đã bổ sung đường đọc optional cho các bảng so sánh khi chúng tồn tại. Kết quả staged run ngày 20/06/2026 đã xác nhận các bảng BigQuery khớp số dòng với Lakehouse, nhưng việc tái chạy vẫn phụ thuộc credential cloud và trạng thái Docker cục bộ. Đây là giới hạn vận hành cần ghi nhận nếu mở rộng demo dashboard cloud.",
        ),
    )
)
changes.append(
    (
        "callcenteren_serving_limit_actual",
        replace_paragraph(
            doc,
            "Về hạ tầng serving, nhánh CallCenterEN đã được ghi vào Iceberg/MinIO và tạo bảng Gold so sánh, nhưng BigQuery sync cho các bảng so sánh mới chỉ được bổ sung ở mức optional code path, chưa chạy xác thực đầy đủ trên môi trường cloud. Đây là giới hạn vận hành cần ghi nhận nếu mở rộng báo cáo sang phần BI ngoài lakehouse local.",
            "Về hạ tầng serving, nhánh CallCenterEN đã được ghi vào Iceberg/MinIO và bq_sync_job.py đã bổ sung đường đọc optional cho các bảng so sánh khi chúng tồn tại. Kết quả staged run ngày 20/06/2026 đã xác nhận các bảng BigQuery khớp số dòng với Lakehouse, nhưng việc tái chạy vẫn phụ thuộc credential cloud và trạng thái Docker cục bộ. Đây là giới hạn vận hành cần ghi nhận nếu mở rộng demo dashboard cloud.",
        ),
    )
)

# Targeted table updates. Indices below are 0-based python-docx table indices.
t = doc.tables[3]
set_cell(t.rows[2].cells[2], "lakehouse.silver.call_logs có call_code do mô hình sinh")

t = doc.tables[18]
set_cell(t.rows[2].cells[0], "callcenteren_15k_candidate")
set_cell(t.rows[2].cells[1], "15,000")
set_cell(t.rows[2].cells[2], "Nguồn chính cho nhánh CallCenterEN trong pipeline multi-source")
set_cell(t.rows[3].cells[0], "pseudo_labels_gemini / callcenteren_labeled")
set_cell(t.rows[3].cells[1], "2,420 / 2,260")
set_cell(t.rows[3].cells[2], "Weak labels thô và tập hợp lệ sau merge/lọc, dùng để train/valid/test model riêng")

t = doc.tables[33]
set_cell(t.rows[1].cells[2], "External dataset branch và corpus đối chiếu domain shift")
set_cell(t.rows[2].cells[2], "15,000 candidate; 2,260 labeled split")
set_cell(t.rows[3].cells[2], "3746.2027")
set_cell(t.rows[4].cells[2], "639.8362")
set_cell(t.rows[5].cells[2], "349.5284")
set_cell(t.rows[7].cells[2], "0.9128")
set_cell(t.rows[8].cells[2], "Trung bình 55.5125 token PII/mẫu")

t = doc.tables[36]
set_cell(t.rows[1].cells[0], "Đọc JSON trực tiếp từ ZIP")
set_cell(t.rows[1].cells[1], "Không giải nén toàn bộ dataset gốc")
set_cell(t.rows[1].cells[2], "Bỏ __MACOSX, giữ source_zip/source_entry/source_domain")
set_cell(t.rows[2].cells[0], "Lọc chất lượng transcript")
set_cell(t.rows[2].cells[1], "Giữ mẫu đủ dài và confidence tốt")
set_cell(t.rows[2].cells[2], "15,000 candidate; confidence trung bình 0.9128")
set_cell(t.rows[3].cells[0], "Deduplicate")
set_cell(t.rows[3].cells[1], "Loại transcript trùng theo normalized text")
set_cell(t.rows[3].cells[2], "text_hash dùng làm khóa split deterministic")
set_cell(t.rows[4].cells[0], "Tạo baseline sample")
set_cell(t.rows[4].cells[1], "Phân tích cơ sở dữ liệu tham chiếu")
set_cell(t.rows[4].cells[2], "3,000 mẫu phục vụ phân tích prompt/schema")
set_cell(t.rows[5].cells[0], "Pseudo-labeling và split")
set_cell(t.rows[5].cells[1], "Gán weak label theo bộ call_code của đề tài")
set_cell(t.rows[5].cells[2], "2,420 pseudo-label; 2,260 dòng hợp lệ, train/valid/test = 1,598/315/347")
set_cell(t.rows[6].cells[0], "Inference model riêng")
set_cell(t.rows[6].cells[1], "Áp dụng model CallCenterEN đã tinh chỉnh cho toàn bộ candidate")
set_cell(t.rows[6].cells[2], "callcenteren_15k_with_model_callcodes.csv có 15,000 dòng")

t = doc.tables[38]
set_cell(t.rows[1].cells[0], "M0")
set_cell(t.rows[1].cells[1], "Dataset chính")
set_cell(t.rows[1].cells[2], "Validation/test của dataset chính")
set_cell(t.rows[1].cells[3], "Baseline BoW production cho pipeline Silver chính")
set_cell(t.rows[2].cells[0], "M1")
set_cell(t.rows[2].cells[1], "Dataset chính")
set_cell(t.rows[2].cells[2], "CallCenterEN labeled split")
set_cell(t.rows[2].cells[3], "Đo domain shift khi model chính đánh chéo sang CallCenterEN")
set_cell(t.rows[3].cells[0], "M2 / CallCenterEN riêng")
set_cell(t.rows[3].cells[1], "CallCenterEN train/valid")
set_cell(t.rows[3].cells[2], "CallCenterEN test")
set_cell(t.rows[3].cells[3], "Huấn luyện/tinh chỉnh threshold cho model riêng của nhánh external")

t = doc.tables[50]
rows = [
    ("wait_for_debezium_connector", "PythonSensor", "Chờ connector mongo-source và task Debezium RUNNING", "Timeout 300 giây, reschedule mode."),
    ("primary_telesales.*.bronze", "SparkSubmitOperator", "Chạy bronze_job.py riêng cho cust_dataset, offer_dataset và call_logs_dataset", "TRIGGER_ONCE=true để task streaming kết thúc được trong Airflow."),
    ("primary_telesales.*.silver", "SparkSubmitOperator", "Chạy silver_job.py riêng theo entity; call_logs thực hiện mask PII, dedup và BoW inference", "Airflow truyền NLP_MODEL_TYPE=bow; RoBERTa là baseline/thí nghiệm tùy chọn."),
    ("primary_telesales.gold.*", "SparkSubmitOperator", "Tạo dim_customer, dim_offer, dim_date và fact_telesales_calls", "Ghi bằng MERGE INTO Iceberg và giữ call_code làm nhãn phân tích."),
    ("callcenteren_external.* / bq_sync_gold", "SparkSubmitOperator / BashOperator", "Ghi nhánh CallCenterEN Bronze/Silver/Gold rồi đồng bộ Gold sang BigQuery", "bq_sync_job.py loại BLOCKED_BIGQUERY_COLUMNS trước khi publish."),
]
for idx, values in enumerate(rows, start=1):
    for col, value in enumerate(values):
        set_cell(t.rows[idx].cells[col], value)

t = doc.tables[53]
set_cell(t.rows[3].cells[1], "drop các cột trong BLOCKED_BIGQUERY_COLUMNS nếu tồn tại")
set_cell(t.rows[3].cells[2], "dim_customer publish không chứa full_name/address; fact publish không chứa transcript")
set_cell(t.rows[3].cells[3], "Giảm rủi ro khi dữ liệu phục vụ BI ra ngoài lakehouse local.")

t = doc.tables[57]
set_cell(
    t.rows[3].cells[1],
    "wait_for_debezium_connector -> primary_telesales.*.bronze -> primary_telesales.*.silver -> primary_telesales.gold.* -> callcenteren_external.* -> bq_sync_gold",
)
set_cell(t.rows[3].cells[2], "Luồng end-to-end có dependency rõ theo dataset và stage.")

t = doc.tables[59]
set_cell(t.rows[8].cells[1], "call_code là array do model sinh; không còn publish call_code_original/call_code_predicted")
set_cell(t.rows[8].cells[2], "Query mẫu một transcript và nhãn model.")

t = doc.tables[65]
set_cell(t.rows[1].cells[4], "0,2527")
set_cell(t.rows[1].cells[5], "0,7261")
set_cell(t.rows[1].cells[6], "0,6407")
set_cell(t.rows[2].cells[4], "0,0004")
set_cell(t.rows[2].cells[5], "0,3970")
set_cell(t.rows[2].cells[6], "0,1017")
set_cell(t.rows[3].cells[4], "0,1412")
set_cell(t.rows[3].cells[5], "0,7241")
set_cell(t.rows[3].cells[6], "0,1933")

t = doc.tables[67]
set_cell(t.rows[1].cells[0], "M0_primary_bow")
set_cell(t.rows[1].cells[1], "13,857")
set_cell(t.rows[1].cells[2], "1,732")
set_cell(t.rows[1].cells[3], "0.2633")
set_cell(t.rows[1].cells[4], "0.7309")
set_cell(t.rows[1].cells[5], "0.6501")
set_cell(t.rows[1].cells[6], "0.7326")
set_cell(t.rows[1].cells[7], "0.0688")
set_cell(t.rows[2].cells[0], "M4_combined_bow")
set_cell(t.rows[2].cells[1], "15,455")
set_cell(t.rows[2].cells[2], "1,732")
set_cell(t.rows[2].cells[3], "0.2564")
set_cell(t.rows[2].cells[4], "0.7283")
set_cell(t.rows[2].cells[5], "0.6429")
set_cell(t.rows[2].cells[6], "0.7304")
set_cell(t.rows[2].cells[7], "0.0696")

t = doc.tables[68]
set_cell(t.rows[1].cells[0], "M0_primary_bow")
set_cell(t.rows[1].cells[1], "13,857")
set_cell(t.rows[1].cells[2], "1,733")
set_cell(t.rows[1].cells[3], "0.2527")
set_cell(t.rows[1].cells[4], "0.7261")
set_cell(t.rows[1].cells[5], "0.6407")
set_cell(t.rows[1].cells[6], "0.7267")
set_cell(t.rows[1].cells[7], "0.0699")
set_cell(t.rows[2].cells[0], "M4_combined_bow")
set_cell(t.rows[2].cells[1], "15,455")
set_cell(t.rows[2].cells[2], "1,733")
set_cell(t.rows[2].cells[3], "0.2499")
set_cell(t.rows[2].cells[4], "0.7238")
set_cell(t.rows[2].cells[5], "0.6389")
set_cell(t.rows[2].cells[6], "0.7249")
set_cell(t.rows[2].cells[7], "0.0708")

t = doc.tables[71]
set_cell(t.rows[5].cells[1], "Parse Bronze, mask PII, deduplicate và sinh call_code bằng model.")
set_cell(t.rows[5].cells[3], "NLP_MODEL_TYPE mặc định là bow; RoBERTa chỉ bật bằng NLP_MODEL_TYPE=roberta cho thí nghiệm so sánh.")

changes.append(("table_call_code_predicted", replace_in_tables(doc, "call_code_predicted", "call_code")))
changes.append(("table_bq_flag", replace_in_tables(doc, "BQ_INCLUDE_PII=false", "BLOCKED_BIGQUERY_COLUMNS")))
changes.append(("paragraph_bq_flag", replace_in_paragraphs(doc, "BQ_INCLUDE_PII=false", "BLOCKED_BIGQUERY_COLUMNS")))

# Re-apply cells where the broad legacy-column replacement would make the prose awkward.
t = doc.tables[59]
set_cell(t.rows[8].cells[1], "call_code là array do model sinh; không còn publish các cột nhãn legacy")
set_cell(t.rows[8].cells[2], "Query mẫu một transcript và nhãn model.")

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)

print(f"Wrote {OUT}")
for name, count in changes:
    print(f"{name}: {count}")
