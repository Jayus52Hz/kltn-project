from __future__ import annotations

from pathlib import Path
import shutil
import zipfile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
REPORTS_DIR = ROOT / "docs" / "reports"

SRC = REPORTS_DIR / "Report KLTN - 22133056 - Nguyen Quoc Thinh - revised callcenteren full integrated.docx"
OUT = REPORTS_DIR / "Report KLTN - 22133056 - Nguyen Quoc Thinh - revised callcenteren full integrated - text fixes.docx"


def clear_paragraph(paragraph):
    p = paragraph._p
    for child in list(p):
        if child.tag == qn("w:pPr"):
            continue
        p.remove(child)


def set_paragraph_text(paragraph, text: str):
    style = paragraph.style
    clear_paragraph(paragraph)
    paragraph.style = style
    paragraph.add_run(text)


def replace_in_paragraph(paragraph, old: str, new: str) -> bool:
    text = paragraph.text
    if old not in text:
        return False
    set_paragraph_text(paragraph, text.replace(old, new))
    return True


def replace_cell_text(cell, text: str):
    if not cell.paragraphs:
        cell.add_paragraph(text)
        return
    set_paragraph_text(cell.paragraphs[0], text)
    for paragraph in cell.paragraphs[1:]:
        clear_paragraph(paragraph)


def iter_paragraphs(doc: Document):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def add_update_fields_on_open(docx_path: Path):
    with zipfile.ZipFile(docx_path, "a") as z:
        settings_name = "word/settings.xml"
        xml = z.read(settings_name).decode("utf-8")
        if "w:updateFields" not in xml:
            insert = '<w:updateFields w:val="true"/>'
            xml = xml.replace("</w:settings>", f"{insert}</w:settings>")
            z.writestr(settings_name, xml)


def find_paragraph(doc: Document, exact_text: str) -> int | None:
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == exact_text:
            return i
    return None


def replace_paragraph_block(doc: Document, start_exclusive: int, end_exclusive: int, entries: list[str]):
    block = doc.paragraphs[start_exclusive + 1:end_exclusive]
    for paragraph, entry in zip(block, entries):
        set_paragraph_text(paragraph, entry)
    for paragraph in block[len(entries):]:
        clear_paragraph(paragraph)


def main():
    if OUT.exists():
        OUT.unlink()
    shutil.copy2(SRC, OUT)

    doc = Document(str(OUT))

    replacements = {
        "1.5. KẾT QUẢ DỰ KIẾN ĐẠT ĐƯỢC": "1.5. KẾT QUẢ ĐẠT ĐƯỢC",
        "Kết quả dự kiến của đề tài gồm": "Kết quả đạt được của đề tài gồm",
        "RoBERTa vẫn có thể bật lại bằng NLP_MODEL_TYPE=bow để phục vụ thí nghiệm hoặc so sánh": "RoBERTa vẫn có thể bật lại bằng NLP_MODEL_TYPE=roberta để phục vụ thí nghiệm hoặc so sánh",
        "silver_job.py mặc định chạy RoBERTa, đồng thời giữ BoW làm fallback bằng biến môi trường.": "silver_job.py mặc định chạy BoW + Logistic Regression; RoBERTa được giữ như cấu hình thí nghiệm có thể bật lại bằng biến môi trường.",
        "tích hợp RoBERTa ở Silver": "tích hợp BoW + Logistic Regression ở Silver",
        "Tích hợp mô hình NLP/RoBERTa để phân loại nội dung cuộc gọi và sinh nhãn dự đoán phục vụ phân tích nghiệp vụ.": "Tích hợp mô hình NLP BoW + Logistic Regression để phân loại nội dung cuộc gọi; RoBERTa được dùng làm baseline học sâu để so sánh.",
        "BoW + Logistic Regression là mô hình chính và BoW là baseline": "BoW + Logistic Regression là mô hình production và RoBERTa là baseline học sâu",
        "Kết quả dự kiến đạt được": "Kết quả đạt được",
    }

    citation_replacements = {
        "Phần lý thuyết được xây dựng từ tài liệu chính thức của các công nghệ như MongoDB, Debezium, Kafka, Spark, Iceberg, Airflow, Docker, MinIO và Superset.": "Phần lý thuyết được xây dựng từ tài liệu chính thức của các công nghệ như MongoDB, Debezium, Kafka, Spark, Iceberg, Airflow, Docker, MinIO và Superset [4]-[10].",
        "Nghiên cứu Lakehouse được công bố tại CIDR 2021 đề xuất một thế hệ nền tảng dữ liệu mở nhằm thống nhất Data Warehousing và Advanced Analytics.": "Nghiên cứu Lakehouse được công bố tại CIDR 2021 đề xuất một thế hệ nền tảng dữ liệu mở nhằm thống nhất Data Warehousing và Advanced Analytics [11].",
        "Mô hình chính của hệ thống là Bag-of-Words kết hợp Logistic Regression theo chiến lược One-vs-Rest.": "Mô hình chính của hệ thống là Bag-of-Words kết hợp Logistic Regression theo chiến lược One-vs-Rest [12]-[16].",
        "RoBERTa vẫn được trình bày trong đề tài với vai trò baseline học sâu.": "RoBERTa vẫn được trình bày trong đề tài với vai trò baseline học sâu [17]-[20].",
        "RoBERTa được giữ làm baseline học sâu để so sánh.": "RoBERTa được giữ làm baseline học sâu để so sánh [17]-[20].",
        "Về lý thuyết, RoBERTa có khả năng biểu diễn ngữ cảnh tốt hơn BoW nhờ self-attention.": "Về lý thuyết, RoBERTa có khả năng biểu diễn ngữ cảnh tốt hơn BoW nhờ self-attention và kiến trúc Transformer/BERT [21], [22].",
        "Mô hình sinh dữ liệu chính là models/gemma-3-27b-it; notebook cũng có phiên bản thử nghiệm ban đầu với models/gemini-2.5-flash.": "Mô hình sinh dữ liệu chính là models/gemma-3-27b-it [1], [2]; notebook cũng có phiên bản thử nghiệm ban đầu với models/gemini-2.5-flash [3].",
        "CallCenterEN là một public dataset chứa transcript hội thoại call center tiếng Anh đã được xử lý PII.": "CallCenterEN là một public dataset chứa transcript hội thoại call center tiếng Anh đã được xử lý PII [25], [26].",
        "Theo Gururangan và cộng sự (2020), tiếp tục pretrain mô hình ngôn ngữ trên dữ liệu cùng miền có thể giúp cải thiện biểu diễn trước khi fine-tune cho tác vụ cụ thể.": "Theo Gururangan và cộng sự (2020), tiếp tục pretrain mô hình ngôn ngữ trên dữ liệu cùng miền có thể giúp cải thiện biểu diễn trước khi fine-tune cho tác vụ cụ thể [27].",
        "Theo Ratner và cộng sự (2016), các nguồn nhãn yếu có thể giúp tạo tập huấn luyện quy mô lớn hơn, nhưng cần xem chúng là noisy labels.": "Theo Ratner và cộng sự (2016), các nguồn nhãn yếu có thể giúp tạo tập huấn luyện quy mô lớn hơn, nhưng cần xem chúng là noisy labels [28].",
        "Theo khảo sát self-training của Amini và cộng sự (2022), các mẫu pseudo-label có độ tin cậy cao có thể được bổ sung vào tập huấn luyện, trong khi tập kiểm thử vẫn phải giữ ground truth độc lập.": "Theo khảo sát self-training của Amini và cộng sự (2022), các mẫu pseudo-label có độ tin cậy cao có thể được bổ sung vào tập huấn luyện, trong khi tập kiểm thử vẫn phải giữ ground truth độc lập [29].",
        "CallCenterEN là tập dữ liệu transcript call center tiếng Anh được giới thiệu trong bài báo Real-World En Call Center Transcripts Dataset with PII Redaction của Dao, Chawla, Banda và DeLeeuw, công bố trên arXiv năm 2025. Dataset card được phát hành công khai trên Hugging Face bởi AIxBlock với tên 92k-real-world-call-center-scripts-english.": "CallCenterEN là tập dữ liệu transcript call center tiếng Anh được giới thiệu trong bài báo Real-World En Call Center Transcripts Dataset with PII Redaction của Dao, Chawla, Banda và DeLeeuw, công bố trên arXiv năm 2025 [25]. Dataset card được phát hành công khai trên Hugging Face bởi AIxBlock với tên 92k-real-world-call-center-scripts-english [26].",
        "CallCenterEN là tập transcript call center tiếng Anh được công bố công khai trong bài báo Real-World En Call Center Transcripts Dataset with PII Redaction.": "CallCenterEN là tập transcript call center tiếng Anh được công bố công khai trong bài báo Real-World En Call Center Transcripts Dataset with PII Redaction [25], [26].",
    }

    for paragraph in iter_paragraphs(doc):
        for old, new in {**replacements, **citation_replacements}.items():
            replace_in_paragraph(paragraph, old, new)

    # Fix duplicated reviewer form label if present.
    seen_advisor_form = False
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if "PHIẾU NHẬN XÉT CỦA GIÁO VIÊN HƯỚNG DẪN" in text:
            if seen_advisor_form:
                set_paragraph_text(paragraph, "PHIẾU NHẬN XÉT CỦA GIÁO VIÊN PHẢN BIỆN")
            seen_advisor_form = True
        replace_in_paragraph(paragraph, "Giáo viên hướng dẫn\n\n(Ký & ghi rõ họ tên)", "Giáo viên phản biện\n\n(Ký & ghi rõ họ tên)")

    # Table-specific fixes from the review.
    for table in doc.tables:
        for row in table.rows:
            row_text = " ".join(cell.text for cell in row.cells)
            if "RoBERTa là mô hình chính nhưng runtime nặng hơn BoW" in row_text:
                replace_cell_text(row.cells[1], "BoW là mô hình production hiện tại; RoBERTa chỉ là baseline học sâu và hướng nâng cấp khi có model serving/GPU.")
                replace_cell_text(row.cells[2], "Tối ưu BoW theo threshold/label, thử model serving cho RoBERTa hoặc transformer nhỏ hơn trong hướng phát triển.")
            if "NLP_MODEL_TYPE mặc định là roberta" in row_text:
                replace_cell_text(row.cells[-1], "NLP_MODEL_TYPE mặc định là bow; RoBERTa chỉ bật bằng NLP_MODEL_TYPE=roberta cho thí nghiệm so sánh.")
            if "Airflow truyền NLP_MODEL_TYPE=roberta" in row_text:
                replace_cell_text(row.cells[-1], "Airflow truyền NLP_MODEL_TYPE=bow; Dockerfile vẫn giữ dependency RoBERTa để phục vụ thí nghiệm so sánh khi cần.")
            if "BoW + Logistic Regression là mô hình chính và BoW là baseline" in row_text:
                replace_cell_text(row.cells[0], "Các artifact chính đều đã có vị trí thuyết minh tương ứng trong báo cáo. Phần dữ liệu bám vào generate_data.ipynb và các CSV nguồn; phần NLP bám vào NLP_model.ipynb, trong đó BoW + Logistic Regression là mô hình production và RoBERTa là baseline học sâu; phần pipeline bám vào các job Bronze/Silver/Gold, DAG Airflow và Dockerfile.")

    # Replace the reference list with a consistent numbered IEEE-like list.
    refs_heading_idx = None
    appendix_heading_idx = None
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == "TÀI LIỆU THAM KHẢO":
            refs_heading_idx = i

    if refs_heading_idx is not None:
        for i in range(refs_heading_idx + 1, len(doc.paragraphs)):
            if doc.paragraphs[i].text.strip() == "PHỤ LỤC":
                appendix_heading_idx = i
                break

    references = [
        "[1] Google Developers Blog, \"Introducing Gemma 3: The Developer Guide,\" https://developers.googleblog.com/en/introducing-gemma3/",
        "[2] Google AI for Developers, \"Gemma documentation,\" https://ai.google.dev/gemma/docs",
        "[3] Google Cloud, \"Gemini 2.5 Flash model documentation,\" https://cloud.google.com/vertex-ai/generative-ai/docs/models/gemini/2-5-flash",
        "[4] Debezium Documentation, \"MongoDB connector,\" https://debezium.io/documentation/reference/stable/connectors/mongodb.html",
        "[5] Apache Kafka Documentation, https://kafka.apache.org/documentation/",
        "[6] Apache Spark Documentation, \"Structured Streaming Programming Guide,\" https://spark.apache.org/docs/latest/structured-streaming-programming-guide.html",
        "[7] Apache Iceberg Documentation, https://iceberg.apache.org/docs/latest/",
        "[8] Apache Airflow Documentation, \"DAGs and task dependencies,\" https://airflow.apache.org/docs/apache-airflow/stable/core-concepts/dags.html",
        "[9] MinIO Documentation, \"Object storage for S3-compatible workloads,\" https://min.io/docs/minio/container/index.html",
        "[10] Apache Superset Documentation, https://superset.apache.org/docs/intro",
        "[11] M. Armbrust, A. Ghodsi, R. Xin, and M. Zaharia, \"Lakehouse: A New Generation of Open Platforms that Unify Data Warehousing and Advanced Analytics,\" CIDR, 2021. https://www.vldb.org/cidrdb/papers/2021/cidr2021_paper17.pdf",
        "[12] scikit-learn Documentation, https://scikit-learn.org/stable/",
        "[13] scikit-learn Documentation, \"CountVectorizer,\" https://scikit-learn.org/stable/modules/generated/sklearn.feature_extraction.text.CountVectorizer.html",
        "[14] scikit-learn Documentation, \"LogisticRegression,\" https://scikit-learn.org/stable/modules/generated/sklearn.linear_model.LogisticRegression.html",
        "[15] scikit-learn Documentation, \"OneVsRestClassifier,\" https://scikit-learn.org/stable/modules/generated/sklearn.multiclass.OneVsRestClassifier.html",
        "[16] scikit-learn Documentation, \"MultiLabelBinarizer,\" https://scikit-learn.org/stable/modules/generated/sklearn.preprocessing.MultiLabelBinarizer.html",
        "[17] Y. Liu et al., \"RoBERTa: A Robustly Optimized BERT Pretraining Approach,\" arXiv:1907.11692, 2019. https://arxiv.org/abs/1907.11692",
        "[18] Hugging Face, \"FacebookAI/roberta-base model card,\" https://huggingface.co/FacebookAI/roberta-base",
        "[19] Hugging Face Transformers, \"Text classification task guide,\" https://huggingface.co/docs/transformers/tasks/sequence_classification",
        "[20] PyTorch Documentation, \"BCEWithLogitsLoss,\" https://docs.pytorch.org/docs/stable/generated/torch.nn.BCEWithLogitsLoss.html",
        "[21] A. Vaswani et al., \"Attention Is All You Need,\" arXiv:1706.03762, 2017. https://arxiv.org/abs/1706.03762",
        "[22] J. Devlin et al., \"BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding,\" arXiv:1810.04805, 2018. https://arxiv.org/abs/1810.04805",
        "[23] Hugging Face Transformers, \"Trainer documentation,\" https://huggingface.co/docs/transformers/main_classes/trainer",
        "[24] Hugging Face Datasets, \"Use with Pandas,\" https://huggingface.co/docs/datasets/use_with_pandas",
        "[25] H. Dao, G. Chawla, R. Banda, and C. DeLeeuw, \"Real-World En Call Center Transcripts Dataset with PII Redaction,\" arXiv:2507.02958, 2025. https://arxiv.org/abs/2507.02958",
        "[26] AIxBlock, \"92k Real-World Call Center Scripts English,\" Hugging Face Dataset. https://huggingface.co/datasets/AIxBlock/92k-real-world-call-center-scripts-english",
        "[27] S. Gururangan et al., \"Don't Stop Pretraining: Adapt Language Models to Domains and Tasks,\" ACL, 2020. https://aclanthology.org/2020.acl-main.740/",
        "[28] A. Ratner, C. De Sa, S. Wu, D. Selsam, and C. Ré, \"Data Programming: Creating Large Training Sets, Quickly,\" NeurIPS, 2016. https://papers.neurips.cc/paper/6523-data-programming-creating-large-training-sets-quickly",
        "[29] M.-R. Amini et al., \"Self-Training: A Survey,\" arXiv:2202.12040, 2022. https://arxiv.org/abs/2202.12040",
    ]

    if refs_heading_idx is not None and appendix_heading_idx is not None:
        target = doc.paragraphs[refs_heading_idx + 1:appendix_heading_idx]
        for paragraph in target:
            clear_paragraph(paragraph)
        for paragraph, ref in zip(target, references):
            set_paragraph_text(paragraph, ref)
        for paragraph in target[len(references):]:
            clear_paragraph(paragraph)

    # The front matter TOC in this DOCX behaves as static text after prior edits.
    # Rewrite it to remove child headings that were listed before their parents.
    toc_idx = find_paragraph(doc, "MỤC LỤC")
    fig_idx = find_paragraph(doc, "DANH MỤC HÌNH ẢNH")
    if toc_idx is not None and fig_idx is not None and toc_idx < fig_idx:
        toc_entries = [
            "PHẦN MỞ ĐẦU",
            "Chương 1: CƠ SỞ LÝ THUYẾT",
            "1.1. Tổng quan về kiến trúc dữ liệu hiện đại",
            "1.2. Kiến trúc Medallion",
            "1.3. Thu thập dữ liệu thay đổi theo thời gian thực",
            "1.4. Công nghệ xử lý và lưu trữ dữ liệu lớn",
            "1.5. Bảo mật dữ liệu cá nhân và khai phá dữ liệu hội thoại",
            "1.6. Điều phối pipeline và phục vụ phân tích",
            "1.7. Lựa chọn kiến trúc theo yêu cầu dữ liệu Telesales",
            "1.8. Đặc điểm dữ liệu cuộc gọi và transcript",
            "1.9. Vấn đề tách tải OLTP và OLAP",
            "1.10. Vai trò của schema evolution trong Lakehouse",
            "1.11. Vai trò của metadata trong truy vết dữ liệu",
            "1.12. Data quality trong kiến trúc Medallion",
            "1.13. Bảo mật PII trong pipeline phân tích",
            "1.14. Machine Learning trong pipeline dữ liệu",
            "1.15. Mô hình Star Schema cho dashboard vận hành",
            "1.16. Lớp serving và semantic layer cho BI",
            "1.17. Tổng hợp tiêu chí lựa chọn công nghệ",
            "1.18. Kết luận chương",
            "Chương 2: BỘ DỮ LIỆU, MÔ HÌNH SINH DỮ LIỆU VÀ MÔ HÌNH NLP",
            "2.1. Vai trò của dữ liệu trong đề tài",
            "2.2. Nguồn gốc và quy trình sinh dữ liệu",
            "2.3. Thiết kế schema và tính chất dữ liệu",
            "2.4. Chuẩn hóa dữ liệu thành các thực thể nguồn",
            "2.5. Chuẩn bị dữ liệu huấn luyện NLP",
            "2.6. Mô hình NLP phân loại call_code",
            "2.7. Data quality, đạo đức dữ liệu và giới hạn của dữ liệu tổng hợp",
            "Chương 3: PHÂN TÍCH VÀ THIẾT KẾ KIẾN TRÚC HYBRID DATA LAKEHOUSE",
            "3.1. Yêu cầu hệ thống AGI Telesales",
            "3.2. Kiến trúc tổng thể",
            "3.3. Thiết kế nguồn dữ liệu và CDC",
            "3.4. Thiết kế các tầng Medallion",
            "3.5. Thiết kế điều phối, serving và BI",
            "3.6. Thiết kế bảo mật và khả năng vận hành",
            "Chương 4: TRIỂN KHAI, THỰC NGHIỆM VÀ KIỂM THỬ HỆ THỐNG",
            "4.1. Môi trường triển khai",
            "4.2. Bootstrap dữ liệu và CDC",
            "4.3. Kiểm thử Bronze, Silver, Gold",
            "4.4. Kiểm thử BigQuery, Superset và dashboard",
            "4.5. Đánh giá NLP trong pipeline",
            "4.6. Kiểm thử khả năng chạy lại và phục hồi lỗi",
            "Chương 5: ĐÁNH GIÁ TỔNG HỢP, HẠN CHẾ VÀ HƯỚNG PHÁT TRIỂN",
            "5.1. Mức độ đáp ứng mục tiêu đề tài",
            "5.2. Đánh giá đóng góp kỹ thuật",
            "5.3. Hạn chế",
            "5.4. Hướng phát triển",
            "PHẦN KẾT LUẬN",
            "TÀI LIỆU THAM KHẢO",
            "PHỤ LỤC",
        ]
        replace_paragraph_block(doc, toc_idx, fig_idx, toc_entries)

    doc.save(str(OUT))
    print(OUT)


if __name__ == "__main__":
    main()
