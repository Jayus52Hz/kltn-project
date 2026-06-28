from __future__ import annotations

from pathlib import Path
import shutil

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
REPORTS_DIR = ROOT / "docs" / "reports"

SRC = REPORTS_DIR / "Report KLTN - 22133056 - Nguyen Quoc Thinh - 08-06.docx"
OUT = REPORTS_DIR / "Report KLTN - 22133056 - Nguyen Quoc Thinh - 08-06 - chapter1 citations notes.docx"


def clear_paragraph(paragraph):
    p = paragraph._p
    for child in list(p):
        if child.tag.endswith("}pPr"):
            continue
        p.remove(child)


def set_paragraph_text(paragraph, text: str):
    style = paragraph.style
    clear_paragraph(paragraph)
    paragraph.style = style
    paragraph.add_run(text)


def replace_once(doc: Document, old: str, new: str):
    for paragraph in doc.paragraphs:
        if old in paragraph.text:
            set_paragraph_text(paragraph, paragraph.text.replace(old, new, 1))
            return True
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if old in paragraph.text:
                        set_paragraph_text(paragraph, paragraph.text.replace(old, new, 1))
                        return True
    return False


def insert_after(paragraph, text: str):
    new_p = paragraph._p.addnext(paragraph._p.__class__())
    # The line above creates an empty w:p node of the same class. python-docx
    # can wrap it through the parent document's paragraph list after save/load,
    # but direct construction is not stable enough for rich formatting.


def add_note_after_paragraph(doc: Document, anchor_contains: str, note_text: str):
    if any(note_text in paragraph.text for paragraph in doc.paragraphs):
        return False
    for paragraph in doc.paragraphs:
        if anchor_contains in paragraph.text:
            p = paragraph._p
            new_p = p.__copy__()
            clear_paragraph(type(paragraph)(new_p, paragraph._parent))
            p.addnext(new_p)
            inserted = type(paragraph)(new_p, paragraph._parent)
            inserted.style = paragraph.style
            run = inserted.add_run(note_text)
            run.italic = True
            return True
    return False


def append_references(doc: Document, refs: list[str]):
    # Add new numbered references before PHỤ LỤC. Use the blank paragraphs
    # before PHỤ LỤC when available; otherwise insert new paragraphs there.
    refs_start = None
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == "TÀI LIỆU THAM KHẢO":
            refs_start = i

    appendix_idx = None
    for i, paragraph in enumerate(doc.paragraphs):
        if refs_start is not None and i > refs_start and paragraph.text.strip() == "PHỤ LỤC":
            appendix_idx = i
            break
    if appendix_idx is None:
        return
    refs_block = "\n".join(
        paragraph.text for paragraph in doc.paragraphs[(refs_start or 0):appendix_idx]
    )
    refs_to_add = [ref for ref in refs if ref.split("]", 1)[0] + "]" not in refs_block]
    if not refs_to_add:
        return

    anchor = doc.paragraphs[appendix_idx]
    for ref in refs_to_add:
        p = anchor._p
        new_p = p.__copy__()
        clear_paragraph(type(anchor)(new_p, anchor._parent))
        p.addprevious(new_p)
        inserted = type(anchor)(new_p, anchor._parent)
        inserted.style = doc.paragraphs[appendix_idx - 1].style
        inserted.add_run(ref)


def main():
    if OUT.exists():
        OUT.unlink()
    shutil.copy2(SRC, OUT)
    doc = Document(str(OUT))

    replacements = {
        "Với MongoDB, Change Streams dựa trên oplog của replica set, cho phép hệ thống downstream theo dõi sự thay đổi dữ liệu mà không cần quét lại toàn bộ collection.": "Với MongoDB, Change Streams dựa trên oplog của replica set, cho phép hệ thống downstream theo dõi sự thay đổi dữ liệu mà không cần quét lại toàn bộ collection [30].",
        "Debezium là nền tảng CDC mã nguồn mở, thường triển khai trên Kafka Connect.": "Debezium là nền tảng CDC mã nguồn mở, thường triển khai trên Kafka Connect [4].",
        "Apache Kafka là nền tảng event streaming phân tán, tổ chức dữ liệu theo topic và partition.": "Apache Kafka là nền tảng event streaming phân tán, tổ chức dữ liệu theo topic và partition [5].",
        "Apache Spark là engine xử lý dữ liệu phân tán, hỗ trợ batch, SQL, structured streaming và machine learning.": "Apache Spark là engine xử lý dữ liệu phân tán, hỗ trợ batch, SQL, structured streaming và machine learning [6].",
        "Apache Iceberg là open table format cho bảng phân tích lớn trên data lake.": "Apache Iceberg là open table format cho bảng phân tích lớn trên data lake [7].",
        "MinIO là object storage tương thích S3, có thể triển khai nội bộ bằng Docker.": "MinIO là object storage tương thích S3, có thể triển khai nội bộ bằng Docker [9], [33].",
        "Apache Airflow điều phối workflow theo DAG.": "Apache Airflow điều phối workflow theo DAG [8].",
        "Apache Superset là nền tảng BI mã nguồn mở, phù hợp cho việc xây dựng dashboard nội bộ trực tiếp trên dữ liệu đã mô hình hóa.": "Apache Superset là nền tảng BI mã nguồn mở, phù hợp cho việc xây dựng dashboard nội bộ trực tiếp trên dữ liệu đã mô hình hóa [10].",
        "BigQuery là kho dữ liệu phân tích dạng serverless trên Google Cloud, phù hợp để lưu trữ dữ liệu phục vụ báo cáo khi cần tách lớp BI khỏi hạ tầng lakehouse local.": "BigQuery là kho dữ liệu phân tích dạng serverless trên Google Cloud, phù hợp để lưu trữ dữ liệu phục vụ báo cáo khi cần tách lớp BI khỏi hạ tầng lakehouse local [31].",
        "Looker Studio có thể kết nối trực tiếp với BigQuery để xây dựng báo cáo tương tác mà không cần triển khai thêm máy chủ BI.": "Looker Studio có thể kết nối trực tiếp với BigQuery để xây dựng báo cáo tương tác mà không cần triển khai thêm máy chủ BI [32].",
        "Hybrid Data Lakehouse phù hợp với phạm vi đề tài vì dữ liệu nhạy cảm được xử lý trong môi trường local bằng MinIO/Iceberg, còn dữ liệu phục vụ dashboard có thể đồng bộ sang BigQuery/Superset.": "Hybrid Data Lakehouse phù hợp với phạm vi đề tài vì dữ liệu nhạy cảm được xử lý trong môi trường local bằng MinIO/Iceberg [7], [9], còn dữ liệu phục vụ dashboard có thể đồng bộ sang BigQuery/Superset [10], [31].",
        "Các công nghệ được chọn theo tiêu chí: khả năng chạy local bằng Docker, hỗ trợ CDC, lưu được dữ liệu raw và dữ liệu phân tích, hỗ trợ upsert/merge, tích hợp được NLP và phục vụ dashboard.": "Các công nghệ được chọn theo tiêu chí: khả năng chạy local bằng Docker/Compose [33], hỗ trợ CDC, lưu được dữ liệu raw và dữ liệu phân tích, hỗ trợ upsert/merge, tích hợp được NLP và phục vụ dashboard.",
    }
    for old, new in replacements.items():
        replace_once(doc, old, new)

    notes = [
        ("Trong hệ thống Telesales, CDC đặc biệt quan trọng", "[NOTE ẢNH - MongoDB Change Streams: Chèn screenshot hoặc sơ đồ MongoDB ReplicaSet/oplog/change stream event; nhấn mạnh dữ liệu thay đổi được phát hiện từ nguồn vận hành.]"),
        ("Debezium MongoDB connector đọc thay đổi từ replica set", "[NOTE ẢNH - Debezium: Chèn sơ đồ Debezium Connect đọc MongoDB change stream và đẩy event sang Kafka; có thể kèm screenshot connector status RUNNING.]"),
        ("Apache Kafka là nền tảng event streaming phân tán", "[NOTE ẢNH - Apache Kafka: Chèn sơ đồ topic/partition cho ba collection cust, offer, call_logs; có thể kèm screenshot kafka-topics --list.]"),
        ("Apache Spark là engine xử lý dữ liệu phân tán", "[NOTE ẢNH - Apache Spark: Chèn hình Spark driver/executor hoặc Spark job Bronze/Silver/Gold; ưu tiên screenshot Spark UI/job completed nếu có.]"),
        ("Apache Iceberg là open table format", "[NOTE ẢNH - Apache Iceberg: Chèn sơ đồ metadata.json, manifest list, manifest files và data files; liên hệ với snapshot/schema evolution/MERGE INTO.]"),
        ("Object storage lưu dữ liệu dưới dạng object", "[NOTE ẢNH - MinIO/Object Storage: Chèn screenshot MinIO bucket warehouse/bronze/silver/gold hoặc sơ đồ Iceberg warehouse đặt trên S3-compatible object storage.]"),
        ("Apache Airflow điều phối workflow theo DAG", "[NOTE ẢNH - Apache Airflow: Chèn screenshot DAG telesales_lakehouse_pipeline với các task wait_for_debezium_connector -> bronze_cdc_ingestion -> silver_etl -> gold_star_schema -> bq_sync_gold.]"),
        ("Apache Superset là nền tảng BI mã nguồn mở", "[NOTE ẢNH - Apache Superset: Chèn screenshot dashboard Superset hoặc chart KPI đọc từ Gold/BigQuery; nhấn mạnh dashboard không truy vấn MongoDB trực tiếp.]"),
        ("Looker Studio có thể kết nối trực tiếp với BigQuery", "[NOTE ẢNH - BigQuery/Looker Studio: Chèn screenshot BigQuery table/view row count và dashboard Looker Studio; nếu chưa dùng Looker Studio thật thì ghi rõ là hướng serving BI cloud.]"),
        ("Các công nghệ được chọn theo tiêu chí", "[NOTE ẢNH - Docker Compose: Chèn sơ đồ container hoặc screenshot docker compose ps -a để minh họa môi trường local gồm MongoDB, Debezium, Kafka, Spark, Airflow, MinIO, Superset.]"),
    ]
    for anchor, note in notes:
        add_note_after_paragraph(doc, anchor, note)

    new_refs = [
        "[30] MongoDB Documentation, \"Change Streams,\" https://www.mongodb.com/docs/manual/changestreams/",
        "[31] Google Cloud Documentation, \"BigQuery overview,\" https://docs.cloud.google.com/bigquery/docs/introduction",
        "[32] Looker Studio Help, \"Tutorial: Create a new report,\" https://support.google.com/looker-studio/answer/06292570",
        "[33] Docker Documentation, \"Docker Compose,\" https://docs.docker.com/compose/",
    ]
    append_references(doc, new_refs)

    doc.save(str(OUT))
    print(OUT)


if __name__ == "__main__":
    main()
