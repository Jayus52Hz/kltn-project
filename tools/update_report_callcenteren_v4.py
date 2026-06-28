from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
INPUT_DOCX = (
    ROOT
    / "docs"
    / "reports"
    / "Report KLTN - 22133056 - Nguyen Quoc Thinh - 08-06 - v3 chuong1 ly thuyet chuyen sau.docx"
)
OUTPUT_DOCX = (
    ROOT
    / "docs"
    / "reports"
    / "Report KLTN - 22133056 - Nguyen Quoc Thinh - 20-06 - v4 callcenteren multisource lakehouse.docx"
)


def norm(text: str) -> str:
    return " ".join(text.split())


def find_paragraph(doc: Document, startswith: str) -> Paragraph:
    for para in doc.paragraphs:
        if norm(para.text).startswith(startswith):
            return para
    raise ValueError(f"Paragraph not found: {startswith}")


def find_heading(doc: Document, text: str) -> Paragraph:
    for para in doc.paragraphs:
        if norm(para.text) == text:
            return para
    raise ValueError(f"Heading not found: {text}")


def delete_paragraph(para: Paragraph) -> None:
    element = para._element
    element.getparent().remove(element)
    para._p = para._element = None


def add_paragraph_after(cursor: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    cursor._p.addnext(new_p)
    para = Paragraph(new_p, cursor._parent)
    if style:
        para.style = style
    if text:
        para.add_run(text)
    return para


def clear_and_set(para: Paragraph, text: str) -> None:
    para.clear()
    para.add_run(text)


def neutralize_figure_notes(doc: Document) -> None:
    replacements = [
        ("Chèn screenshot hoặc sơ đồ", "Ảnh chụp hoặc sơ đồ"),
        ("Chèn screenshot", "Ảnh chụp"),
        ("Chèn sơ đồ", "Sơ đồ"),
        ("Chèn hình", "Hình"),
        ("screenshot", "ảnh chụp"),
        ("có thể kèm", "kèm"),
        ("ưu tiên screenshot", "ảnh chụp"),
        ("nếu có", "khi khả dụng"),
        (
            "nếu chưa dùng Looker Studio thật thì ghi rõ là hướng serving BI cloud",
            "trường hợp Looker Studio chưa được triển khai thật được mô tả là hướng serving BI cloud",
        ),
    ]
    for para in doc.paragraphs:
        text = norm(para.text)
        if not text.startswith("[NOTE ẢNH - "):
            continue
        inner = text.removeprefix("[NOTE ẢNH - ").removesuffix("]")
        for old, new in replacements:
            inner = inner.replace(old, new)
        clear_and_set(para, f"[MINH CHỨNG HÌNH - {inner}]")


def insert_after_paragraph(doc: Document, startswith: str, text: str, style: str = "Normal") -> Paragraph:
    return add_paragraph_after(find_paragraph(doc, startswith), text, style)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold


def set_table_width(table: Table) -> None:
    table.autofit = True
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"), "5000")


def add_table_after(cursor: Paragraph, rows: list[list[str]], style: str = "Table Grid") -> Paragraph:
    doc = cursor._parent
    table = doc.add_table(rows=len(rows), cols=len(rows[0]), width=Inches(6.5))
    table.style = style
    set_table_width(table)
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            set_cell_text(table.cell(r_idx, c_idx), value, bold=(r_idx == 0))

    tbl = table._tbl
    tbl.getparent().remove(tbl)
    cursor._p.addnext(tbl)

    new_p = OxmlElement("w:p")
    tbl.addnext(new_p)
    return Paragraph(new_p, cursor._parent)


def section_body_paragraphs(doc: Document, heading_text: str, stop_heading_prefixes: tuple[str, ...]) -> list[Paragraph]:
    heading = find_heading(doc, heading_text)
    paragraphs = doc.paragraphs
    start = next(
        index for index, para in enumerate(paragraphs)
        if para._element is heading._element
    ) + 1
    result: list[Paragraph] = []
    for para in paragraphs[start:]:
        text = norm(para.text)
        style = para.style.name if para.style else ""
        if style.startswith("Heading") and any(text.startswith(prefix) for prefix in stop_heading_prefixes):
            break
        result.append(para)
    return result


def replace_section(doc: Document, heading_text: str, stop_heading_prefixes: tuple[str, ...], items: list[dict]) -> None:
    heading = find_heading(doc, heading_text)
    for para in section_body_paragraphs(doc, heading_text, stop_heading_prefixes):
        delete_paragraph(para)

    cursor = heading
    for item in items:
        kind = item["kind"]
        if kind == "p":
            cursor = add_paragraph_after(cursor, item["text"], item.get("style", "Normal"))
        elif kind == "caption":
            cursor = add_paragraph_after(cursor, item["text"], item.get("style", "Report Caption"))
        elif kind == "table":
            cursor = add_table_after(cursor, item["rows"])
        else:
            raise ValueError(f"Unsupported item kind: {kind}")


def main() -> None:
    doc = Document(str(INPUT_DOCX))

    # Keep the static table list aligned with captions inserted in this script.
    insert_after_paragraph(
        doc,
        "Bảng 2.10. Đề xuất rule kiểm soát chất lượng dữ liệu sau generation",
        "Bảng 2.11. Kết quả chuẩn bị split CallCenterEN sau pseudo-label",
    )
    table_list_cursor = insert_after_paragraph(
        doc,
        "Bảng 4.4. Định nghĩa KPI và biểu đồ dashboard",
        "Bảng 4.5. Kết quả huấn luyện và tinh chỉnh model riêng cho CallCenterEN",
    )
    table_list_cursor = add_paragraph_after(
        table_list_cursor,
        "Bảng 4.6. So sánh trực quan hai mô hình theo nguồn dữ liệu",
        "Normal",
    )
    add_paragraph_after(
        table_list_cursor,
        "Bảng 4.7. Kết quả ghi bảng Lakehouse cho nhánh CallCenterEN",
        "Normal",
    )

    # Opening section: make CallCenterEN visible as a completed contribution.
    clear_and_set(
        find_paragraph(doc, "Kết quả cuối cùng là một prototype end-to-end"),
        (
            "Kết quả cuối cùng là một prototype end-to-end có khả năng tái lập cục bộ, kiểm soát PII, tách tải OLTP/OLAP và mở rộng sang mô hình multi-source bằng nhánh CallCenterEN."
        ),
    )
    insert_after_paragraph(
        doc,
        "• Thiết kế dashboard và bộ KPI phục vụ theo dõi hiệu quả chiến dịch",
        (
            "• Mở rộng kiến trúc từ pipeline đơn nguồn sang multi-source Lakehouse bằng cách đưa CallCenterEN thành nhánh dataset thứ hai, có split, mô hình call_code riêng và các bảng Gold so sánh với dataset AGI Telesales."
        ),
    )

    # Remove report-writing instructions from body text and replace with completed-report prose.
    clear_and_set(
        find_paragraph(doc, "Superset phục vụ dashboard BI"),
        (
            "Superset phục vụ dashboard BI, còn dashboard tĩnh giúp xem nhanh file dashboard_data.json. Các biểu đồ được tổ chức quanh các KPI chính gồm tổng cuộc gọi, successful sales, success rate, outcome breakdown, lead source, product performance, credit tier, talk time band và agent performance. Đây là các phân tích trực tiếp xuất phát từ Star Schema."
        ),
    )
    clear_and_set(
        find_paragraph(doc, "Khi chụp hình dashboard"),
        (
            "Các minh chứng dashboard trong báo cáo được ưu tiên gắn với dữ liệu thật từ Gold hoặc BigQuery, thay vì dùng giao diện minh họa rỗng. Cách trình bày này giúp hình ảnh liên kết trực tiếp với kiến trúc dữ liệu đã triển khai và làm rõ dashboard đọc từ lớp serving nào."
        ),
    )
    clear_and_set(
        find_paragraph(doc, "Tuy nhiên, Bronze là append log của CDC"),
        (
            "Bronze đóng vai trò append log của CDC nên có thể tăng số bản ghi khi reset checkpoint hoặc replay Kafka từ earliest. Đây không phải lỗi nếu Bronze được xem là vùng raw/audit. Ngược lại, Silver và Gold phải duy trì tính ổn định nhờ deduplicate và MERGE INTO theo khóa tự nhiên, qua đó bảo đảm bảng sạch không bị nhân đôi khi pipeline chạy lại."
        ),
    )
    clear_and_set(
        find_paragraph(doc, "Notebook generate_data.ipynb hiện có cấu hình API key cục bộ"),
        (
            "Notebook generate_data.ipynb hiện có cấu hình API key cục bộ, đây là rủi ro bảo mật ở mức artifact phát triển. Phương án xử lý cho phiên bản công bố gồm loại bỏ khóa khỏi notebook, rotate khóa nếu từng chia sẻ và chuyển cấu hình sang biến môi trường hoặc secret manager. Báo cáo không ghi lại giá trị API key để tránh lộ thông tin nhạy cảm."
        ),
    )
    clear_and_set(
        find_paragraph(doc, "Phụ lục B liệt kê các sơ đồ và hình minh họa cần có"),
        (
            "Phụ lục B hệ thống hóa các nhóm sơ đồ và hình minh họa gắn với phần thuyết minh chính: kiến trúc tổng thể, luồng CDC, Medallion, Star Schema, Airflow DAG, PII flow, Docker Compose map, Debezium/Kafka status, BigQuery view, so sánh mô hình NLP và dashboard BI."
        ),
    )
    clear_and_set(
        find_paragraph(doc, "Khi bổ sung hình thật"),
        (
            "Các vị trí minh họa trong thân bài được chuẩn bị dưới dạng placeholder trung tính cho sơ đồ kiến trúc, screenshot vận hành, dashboard và kết quả truy vấn kiểm chứng. Nhóm placeholder này xác định loại bằng chứng hình ảnh tương ứng với từng phần thuyết minh và không làm thay đổi logic báo cáo."
        ),
    )
    neutralize_figure_notes(doc)
    fig_dashboard_outcome = find_paragraph(doc, "Hình 4.8. Dashboard phân tích outcome")
    fig_dashboard_superset = find_paragraph(doc, "Hình 4.3. Dashboard Superset/BI")
    clear_and_set(
        fig_dashboard_outcome,
        "Hình 4.3. Dashboard Superset/BI cho phân tích hiệu suất telesales",
    )
    clear_and_set(
        fig_dashboard_superset,
        "Hình 4.8. Dashboard phân tích outcome, lead source và product performance",
    )

    # Chapter 2: revise the role of CallCenterEN from an external-only baseline
    # to a second first-class dataset branch in the report narrative.
    replacements = {
        "Tập dữ liệu chính của đề tài vẫn là bộ dữ liệu telesales do tác giả xây dựng": (
            "Sau giai đoạn thực nghiệm bổ sung, đề tài được mở rộng từ pipeline đơn nguồn thành mô hình multi-source Hybrid Data Lakehouse. "
            "Dataset AGI Telesales vẫn là nhánh nghiệp vụ nội bộ, có đầy đủ customer, offer, call_logs, transcript và nhãn gốc phục vụ huấn luyện ban đầu. "
            "CallCenterEN được nâng lên thành nhánh dataset chính thứ hai để kiểm chứng khả năng tiếp nhận một nguồn call-center độc lập và so sánh khác biệt miền dữ liệu."
        ),
        "Tuy nhiên, bộ dữ liệu chính là dữ liệu được xây dựng cho bối cảnh đồ án": (
            "Việc đặt CallCenterEN ngang hàng không có nghĩa là trộn lẫn hai nguồn dữ liệu. Hai nhánh được xử lý độc lập, có split huấn luyện/kiểm thử riêng và có mô hình call_code riêng. "
            "Cách tổ chức này giúp báo cáo đánh giá được hai câu hỏi khác nhau: dataset chính vận hành tốt đến đâu trong bài toán AGI Telesales, và một nguồn call-center thực tế bên ngoài khác với dữ liệu nội bộ như thế nào."
        ),
        "Ngoài vai trò đối chiếu, một phần CallCenterEN được đưa vào thí nghiệm huấn luyện phụ": (
            "CallCenterEN thiếu các thực thể nghiệp vụ như customer, offer và campaign nên không thay thế trực tiếp schema nguồn của AGI Telesales. "
            "Thay vào đó, nhánh CallCenterEN có schema phân tích riêng, tập trung vào transcript, domain, hướng cuộc gọi, thời lượng, ASR confidence, metadata PII và call_code do mô hình riêng sinh ra. "
            "Ranh giới này cho phép so sánh dữ liệu và mô hình mà không làm sai quan hệ nghiệp vụ của dataset chính."
        ),
        "Ở bước đầu, nghiên cứu ưu tiên các nhóm gần với telesales và customer service": (
            "Ở bước triển khai mở rộng, nghiên cứu tạo candidate dataset gồm 15.000 dòng CallCenterEN từ nhiều domain như medicare, insurance, home_service_telecom, automotive, customer_service và medical_equipment. "
            "Tập này không còn chỉ là subset phụ trợ nhỏ, mà được dùng làm nguồn đầu vào chính cho nhánh CallCenterEN trong Lakehouse."
        ),
        "Kết quả lọc đã tạo hai tập con khác nhau": (
            "Sau khi dừng quá trình sinh pseudo-label, workspace ghi nhận 2.420 dòng đã được Gemini gán nhãn, trong đó 2.340 dòng vượt ngưỡng chất lượng mặc định. "
            "Khi merge với candidate dataset và lọc theo should_use_for_training=true, pseudo_label_confidence >= 0,80 và pseudo_call_code không rỗng, tập huấn luyện chính thức còn 2.260 dòng có text_hash duy nhất."
        ),
        "Tóm lại, CallCenterEN được giới thiệu và sử dụng theo ba tầng": (
            "Tóm lại, CallCenterEN được sử dụng theo hai lớp vai trò. Lớp thứ nhất là cơ sở học thuật để giải thích đặc trưng transcript, PII, duration và domain shift trong dữ liệu call center. "
            "Lớp thứ hai là nhánh dữ liệu thực nghiệm độc lập, có split, mô hình và bảng Lakehouse riêng để so sánh với dataset AGI Telesales."
        ),
        "Cần phân biệt rõ hai con số trong nghiên cứu": (
            "Cần phân biệt rõ các mốc dữ liệu trong nghiên cứu: 3.000 mẫu ban đầu phục vụ phân tích đặc trưng và thiết kế prompt; 15.000 dòng candidate là nguồn của nhánh CallCenterEN; "
            "2.260 dòng pseudo-label đạt ngưỡng được dùng để train/valid/test model riêng; sau đó model đã được áp dụng lại cho toàn bộ 15.000 dòng để sinh model_call_code."
        ),
        "Bảng so sánh cho thấy CallCenterEN phù hợp để làm cơ sở đối chiếu cấu trúc dữ liệu": (
            "Bảng so sánh cho thấy hai nguồn dữ liệu có vai trò khác nhau trong cùng kiến trúc multi-source. Dataset AGI Telesales có quan hệ customer-offer-call đầy đủ để phục vụ star schema nghiệp vụ, còn CallCenterEN có transcript thật, domain đa dạng và metadata PII phong phú để kiểm chứng tính tổng quát của thiết kế xử lý hội thoại. "
            "Vì vậy, CallCenterEN không thay thế dataset chính, nhưng được tổ chức như một nhánh dữ liệu chính ngang hàng trong lớp thực nghiệm so sánh."
        ),
    }
    for prefix, new_text in replacements.items():
        clear_and_set(find_paragraph(doc, prefix), new_text)

    replace_section(
        doc,
        "2.5.3. Lọc CallCenterEN subset và tạo pseudo-label",
        ("2.6.",),
        [
            {
                "kind": "p",
                "text": (
                    "Quy trình chuẩn bị CallCenterEN sau khi mở rộng được tổ chức thành ba bước: tạo candidate dataset 15.000 dòng, sinh pseudo-label cho một phần dữ liệu, sau đó tạo split deterministic cho tập đạt ngưỡng chất lượng. "
                    "Script prepare_callcenteren_splits.py merge callcenteren_15k_candidate.csv với pseudo_labels_gemini.csv qua text_hash, giúp tránh phụ thuộc vào thứ tự dòng hoặc tên file gốc."
                ),
            },
            {
                "kind": "p",
                "text": (
                    "Điều kiện lọc mặc định gồm should_use_for_training=true, pseudo_label_confidence >= 0,80 và pseudo_call_code không rỗng. "
                    "Sau khi lọc, tập dữ liệu còn 2.260 dòng duy nhất theo text_hash, có confidence trung bình 0,8901. "
                    "Split được tạo theo hash cố định với tỷ lệ 70/15/15, tương ứng 1.598 dòng train, 315 dòng validation và 347 dòng test."
                ),
            },
            {
                "kind": "caption",
                "text": "Bảng 2.11. Kết quả chuẩn bị split CallCenterEN sau pseudo-label",
            },
            {
                "kind": "table",
                "rows": [
                    ["Hạng mục", "Giá trị", "Ý nghĩa"],
                    ["Candidate dataset", "15.000 dòng", "Nguồn CallCenterEN dùng cho nhánh dữ liệu thứ hai"],
                    ["Pseudo-label đã sinh", "2.420 dòng", "Số dòng được Gemini gán nhãn trước khi dừng gen"],
                    ["Pseudo-label đạt ngưỡng", "2.340 dòng", "Dòng vượt điều kiện chất lượng ban đầu"],
                    ["Dòng đưa vào split", "2.260 dòng", "Dòng hợp lệ sau merge/lọc và không trùng text_hash"],
                    ["Train/Valid/Test", "1.598 / 315 / 347", "Split deterministic theo text_hash"],
                ],
            },
            {
                "kind": "p",
                "text": (
                    "Pseudo-label được chấp nhận là nhãn tin cậy ở mức pilot để huấn luyện mô hình riêng cho CallCenterEN, nhưng bản chất vẫn là weak label. "
                    "Do đó, báo cáo tách rõ đánh giá trên CallCenterEN test set với đánh giá trên dataset chính, không dùng test set của nguồn này để huấn luyện cho nguồn kia."
                ),
            },
        ],
    )

    replace_section(
        doc,
        "2.6.7. Cấu hình thí nghiệm auxiliary training",
        ("2.7.",),
        [
            {
                "kind": "p",
                "text": (
                    "Sau thí nghiệm multi-source ban đầu, hướng triển khai được điều chỉnh từ combined model sang hai mô hình riêng. "
                    "Dataset AGI Telesales tiếp tục sử dụng BoW + Logistic Regression hiện có cho pipeline Silver chính; CallCenterEN sử dụng một model BoW/TF-IDF riêng được huấn luyện và tinh chỉnh ngưỡng trên split CallCenterEN. "
                    "Quyết định này phù hợp với kết quả domain shift: model huấn luyện trên dataset chính khi đánh sang CallCenterEN chỉ đạt micro-F1 0,3970 và macro-F1 0,1017."
                ),
            },
            {
                "kind": "p",
                "text": (
                    "Script finetune_callcenteren_bow.py thử nhiều biến thể CountVectorizer/TF-IDF, Logistic Regression và tinh chỉnh threshold theo validation. "
                    "Model được chọn theo validation micro-F1 rồi Jaccard là count_word_lr_threshold. "
                    "Sau khi chọn model, artifact được lưu tại callcenteren_best_finetuned_model.pkl và được áp dụng cho toàn bộ 15.000 dòng CallCenterEN để sinh model_call_code cùng confidence."
                ),
            },
        ],
    )

    replace_section(
        doc,
        "4.5.5. Thực nghiệm auxiliary training với CallCenterEN",
        ("4.6.",),
        [
            {
                "kind": "p",
                "text": (
                    "Thực nghiệm CallCenterEN được chạy lại theo hướng nhánh dataset chính thứ hai, thay vì chỉ xem như tập phụ trợ. "
                    "Trước khi huấn luyện model riêng, model BoW đã huấn luyện trên dataset AGI Telesales được đánh giá trên 2.260 dòng CallCenterEN đạt exact match 0,0004, Jaccard trung bình 0,2606, micro-F1 0,3970 và macro-F1 0,1017. "
                    "Kết quả này là bằng chứng trực tiếp cho domain shift giữa hai nguồn dữ liệu."
                ),
            },
            {
                "kind": "p",
                "text": (
                    "Do đó, nghiên cứu chuyển sang huấn luyện và tinh chỉnh ngưỡng cho model riêng của CallCenterEN. "
                    "Ba cấu hình chính được thử gồm count_word_lr_threshold, tfidf_word_lr_threshold và tfidf_word_char_lr_threshold. "
                    "Model count_word_lr_threshold được chọn vì có validation micro-F1 cao nhất trong tiêu chí chọn model, dù một số cấu hình TF-IDF có weighted-F1 test nhỉnh hơn."
                ),
            },
            {
                "kind": "caption",
                "text": "Bảng 4.5. Kết quả huấn luyện và tinh chỉnh model riêng cho CallCenterEN",
            },
            {
                "kind": "table",
                "rows": [
                    ["Model", "Split", "Rows", "Exact match", "Avg Jaccard", "Micro-F1", "Macro-F1", "Weighted-F1"],
                    ["count_word_lr_threshold", "valid", "315", "0,1302", "0,6144", "0,7411", "0,2210", "0,7273"],
                    ["count_word_lr_threshold", "test", "347", "0,1412", "0,5943", "0,7241", "0,1933", "0,7128"],
                    ["tfidf_word_lr_threshold", "test", "347", "0,1326", "0,6049", "0,7360", "0,2080", "0,7426"],
                    ["tfidf_word_char_lr_threshold", "test", "347", "0,1326", "0,5983", "0,7312", "0,2147", "0,7396"],
                ],
            },
            {
                "kind": "p",
                "text": (
                    "Để phần thực nghiệm dễ quan sát hơn, kết quả được tổng hợp lại theo góc nhìn hai mô hình riêng. "
                    "Bảng dưới đây không chỉ so sánh chỉ số trong cùng miền dữ liệu, mà còn cho thấy khi model của dataset chính được đem đánh sang CallCenterEN thì hiệu năng giảm mạnh. "
                    "Đây là cơ sở thực nghiệm cho quyết định không dùng một combined model duy nhất ở giai đoạn hiện tại."
                ),
            },
            {
                "kind": "p",
                "text": (
                    "Dòng AGI Telesales đúng miền trong bảng sử dụng kết quả từ notebook NLP_model.ipynb; hai dòng liên quan đến CallCenterEN sử dụng kết quả từ các script thực nghiệm CallCenterEN. "
                    "Bảng được dùng để quan sát xu hướng domain shift và định vị hai model riêng, không thay thế một benchmark duy nhất được huấn luyện và đánh giá lại trong cùng một lần chạy."
                ),
            },
            {
                "kind": "caption",
                "text": "Bảng 4.6. So sánh trực quan hai mô hình theo nguồn dữ liệu",
            },
            {
                "kind": "table",
                "rows": [
                    ["Kịch bản", "Model", "Train", "Test", "Exact match", "Micro-F1", "Macro-F1", "Nhận xét"],
                    [
                        "Đúng miền dataset chính",
                        "AGI Telesales BoW + Logistic Regression",
                        "AGI Telesales train",
                        "AGI Telesales test",
                        "0,1789",
                        "0,7016",
                        "0,6208",
                        "Mô hình production cho pipeline Silver chính",
                    ],
                    [
                        "Đánh chéo sang CallCenterEN",
                        "AGI Telesales BoW + Logistic Regression",
                        "AGI Telesales train",
                        "CallCenterEN labeled",
                        "0,0004",
                        "0,3970",
                        "0,1017",
                        "Hiệu năng giảm rõ, thể hiện domain shift",
                    ],
                    [
                        "Đúng miền CallCenterEN",
                        "CallCenterEN count_word_lr_threshold",
                        "CallCenterEN train",
                        "CallCenterEN test",
                        "0,1412",
                        "0,7241",
                        "0,1933",
                        "Mô hình riêng phù hợp hơn cho nhánh CallCenterEN",
                    ],
                ],
            },
            {
                "kind": "p",
                "text": (
                    "Sau khi chọn model, script áp dụng model lên toàn bộ 15.000 dòng candidate để sinh schema hoàn chỉnh với model_call_code. "
                    "Tập kết quả có 2.340 dòng giữ được pseudo-label gốc từ Gemini để đối chiếu, confidence trung bình của model_call_code là 0,888. "
                    "Các nhãn phổ biến nhất gồm OPENING, NEEDS_ANALYSIS, FEE_DISCUSSION, PRODUCT_PITCH, CURIOUS_EXPLORATION và ACTIVE_LISTENING."
                ),
            },
            {
                "kind": "p",
                "text": (
                    "Spark ETL callcenteren_external_job.py đã đọc file schema 15.000 dòng và ghi nhánh CallCenterEN vào Iceberg/MinIO. "
                    "Kết quả kiểm chứng bằng spark-sql cho thấy bronze_external.callcenteren_raw, silver_external.callcenteren_clean và silver_external.callcenteren_labeled đều có 15.000 dòng; các bảng Gold so sánh cũng được tạo thành công."
                ),
            },
            {
                "kind": "caption",
                "text": "Bảng 4.7. Kết quả ghi bảng Lakehouse cho nhánh CallCenterEN",
            },
            {
                "kind": "table",
                "rows": [
                    ["Bảng Lakehouse", "Số dòng", "Vai trò"],
                    ["lakehouse.bronze_external.callcenteren_raw", "15.000", "Lưu dữ liệu CallCenterEN thô sau ingest"],
                    ["lakehouse.silver_external.callcenteren_clean", "15.000", "Chuẩn hóa transcript, duration, PII metadata và confidence"],
                    ["lakehouse.silver_external.callcenteren_labeled", "15.000", "Tách model_call_code thành mảng call_code phục vụ phân tích"],
                    ["lakehouse.gold_external.callcenteren_call_analytics", "8", "Tổng hợp theo source_domain và call_direction"],
                    ["lakehouse.gold.dataset_profile_comparison", "2", "So sánh profile dataset chính và CallCenterEN"],
                    ["lakehouse.gold.call_code_distribution_comparison", "59", "So sánh phân bố call_code giữa hai nguồn"],
                    ["lakehouse.gold.model_experiment_comparison", "6", "Lưu metrics huấn luyện và tinh chỉnh model CallCenterEN"],
                ],
            },
        ],
    )

    replace_section(
        doc,
        "5.2.4. Đóng góp từ external baseline và auxiliary corpus",
        ("5.3.",),
        [
            {
                "kind": "p",
                "text": (
                    "Việc đưa CallCenterEN vào nghiên cứu tạo ra đóng góp mới về thiết kế multi-source Lakehouse. "
                    "Thay vì chỉ có một dataset tổng hợp nội bộ, đồ án có thêm một nhánh dữ liệu call-center độc lập với candidate dataset 15.000 dòng, split pseudo-label riêng, model call_code riêng và các bảng Bronze/Silver/Gold riêng. "
                    "Điều này giúp kiến trúc thể hiện khả năng mở rộng sang nguồn dữ liệu ngoài mà không phá vỡ pipeline chính."
                ),
            },
            {
                "kind": "p",
                "text": (
                    "Đóng góp thứ hai là bằng chứng thực nghiệm về domain shift. "
                    "Model của dataset chính đánh sang CallCenterEN cho micro-F1 0,3970, trong khi model riêng CallCenterEN sau huấn luyện và tinh chỉnh đạt micro-F1 test 0,7241. "
                    "Kết quả này củng cố quyết định tách hai model theo nguồn dữ liệu, đồng thời cung cấp cơ sở định lượng để báo cáo không chỉ mô tả kiến trúc mà còn đánh giá được tác động của khác biệt miền dữ liệu."
                ),
            },
        ],
    )

    replace_section(
        doc,
        "5.3.4. Hạn chế khi sử dụng CallCenterEN và pseudo-label",
        ("5.4.",),
        [
            {
                "kind": "p",
                "text": (
                    "CallCenterEN là dữ liệu tiếng Anh và đa miền, trong đó nhiều domain như medicare, automotive hoặc home_service_telecom không trùng hoàn toàn với bối cảnh AGI Telesales tài chính. "
                    "Vì vậy, kết quả so sánh cần được hiểu như đánh giá domain shift và khả năng mở rộng kiến trúc, không phải bằng chứng rằng hai nguồn dữ liệu có thể thay thế trực tiếp cho nhau."
                ),
            },
            {
                "kind": "p",
                "text": (
                    "Pseudo-label được chấp nhận để triển khai pilot và huấn luyện model riêng, nhưng vẫn có nguy cơ nhiễu, đặc biệt với các nhãn outcome hoặc thái độ khách hàng khi transcript không biểu hiện rõ. "
                    "Do đó, các kết quả như micro-F1 0,7241 trên CallCenterEN test set nên được xem là kết quả trên weak labels có kiểm soát, chưa tương đương ground truth do annotator độc lập đánh giá."
                ),
            },
            {
                "kind": "p",
                "text": (
                    "Về hạ tầng serving, nhánh CallCenterEN đã được ghi vào Iceberg/MinIO và tạo bảng Gold so sánh, nhưng BigQuery sync cho các bảng so sánh mới chỉ được bổ sung ở mức optional code path, chưa chạy xác thực đầy đủ trên môi trường cloud. "
                    "Đây là giới hạn vận hành cần ghi nhận nếu mở rộng báo cáo sang phần BI ngoài lakehouse local."
                ),
            },
        ],
    )

    replace_section(
        doc,
        "4.6.2. Các lỗi thực nghiệm đã ghi nhận",
        ("Chương 5:",),
        [
            {
                "kind": "p",
                "text": (
                    "Trong quá trình rebuild pipeline, một nhóm lỗi vận hành liên quan đến dependency đã được ghi nhận, bao gồm thiếu hoặc chưa đồng bộ các thư viện scikit-learn, joblib, PyTorch và Transformers trong môi trường Spark/Airflow. "
                    "Các lỗi này không làm thay đổi thiết kế kiến trúc, nhưng cho thấy pipeline NLP phụ thuộc mạnh vào việc đóng gói artifact và môi trường runtime. Sau khi chuẩn hóa dependency và ưu tiên BoW cho production path, pipeline Silver có thể chạy ổn định hơn trong môi trường CPU-only."
                ),
            },
            {
                "kind": "p",
                "text": (
                    "Ngoài lỗi runtime, dữ liệu cũng có lỗi chất lượng như full_date null do timestamp parse lỗi. Thay vì loại bỏ âm thầm, báo cáo ghi nhận các lỗi này như một phần giới hạn thực nghiệm của dữ liệu synthetic và pipeline local. "
                    "Đối với nhánh CallCenterEN, quá trình chạy Spark ETL cũng phát sinh lỗi môi trường Ivy cache và lỗi aggregation khi tạo profile comparison; các lỗi này đã được xử lý trước khi bảng Bronze/Silver/Gold CallCenterEN được ghi thành công."
                ),
            },
        ],
    )

    heading_renames = {
        "2.6.7. Cấu hình thí nghiệm auxiliary training": "2.6.7. Mô hình riêng cho nhánh CallCenterEN",
        "4.5.5. Thực nghiệm auxiliary training với CallCenterEN": "4.5.5. Thực nghiệm CallCenterEN như nhánh dataset chính",
        "5.2.4. Đóng góp từ external baseline và auxiliary corpus": "5.2.4. Đóng góp từ nhánh CallCenterEN trong kiến trúc multi-source",
    }
    for old_heading, new_heading in heading_renames.items():
        clear_and_set(find_heading(doc, old_heading), new_heading)

    conclusion = find_paragraph(doc, "Đề tài đã xây dựng một prototype Hybrid Data Lakehouse cho AGI Telesales")
    clear_and_set(
        conclusion,
        (
            "Đề tài đã xây dựng một prototype Hybrid Data Lakehouse cho AGI Telesales theo hướng end-to-end: sinh dữ liệu tổng hợp, chuẩn hóa thành nguồn MongoDB, thu thập thay đổi bằng Debezium/Kafka, xử lý bằng Spark/Iceberg theo kiến trúc Medallion, tích hợp BoW + Logistic Regression ở Silver, xây dựng Star Schema ở Gold và phục vụ dashboard qua BigQuery/Superset. "
            "Sau phần mở rộng CallCenterEN, kiến trúc được nâng từ pipeline đơn nguồn thành multi-source Lakehouse: dataset AGI Telesales và CallCenterEN được tổ chức thành hai nhánh dữ liệu chính để so sánh, trong đó CallCenterEN có model call_code riêng và đã được ghi vào Bronze/Silver/Gold."
        ),
    )

    doc.save(str(OUTPUT_DOCX))
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
