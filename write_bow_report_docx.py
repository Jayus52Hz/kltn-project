from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt, RGBColor


INPUT_PATH = Path("REPORT_REVISION_DRAFT.md")
OUTPUT_PATH = Path("Report KLTN - 22133056 - Nguyen Quoc Thinh - revised bow model.docx")


def set_run_font(run, size=11, bold=False, italic=False, color=None, name="Arial"):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_code_block(doc, lines):
    para = doc.add_paragraph()
    para.style = "No Spacing"
    for i, line in enumerate(lines):
        run = para.add_run(line)
        set_run_font(run, size=9, name="Consolas", color=(40, 40, 40))
        if i != len(lines) - 1:
            run.add_break()
    para.paragraph_format.left_indent = Pt(12)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(6)


def add_table(doc, rows):
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    header = table.rows[0].cells
    for idx, value in enumerate(rows[0]):
        header[idx].text = value
        for paragraph in header[idx].paragraphs:
            for run in paragraph.runs:
                set_run_font(run, size=9, bold=True)
    for row_values in rows[1:]:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            cells[idx].text = value
            for paragraph in cells[idx].paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size=9)
    doc.add_paragraph()


def parse_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw = lines[i].strip().strip("|")
        cells = [cell.strip().replace("\\|", "|") for cell in raw.split("|")]
        if not all(set(cell) <= {"-", ":"} for cell in cells):
            rows.append(cells)
        i += 1
    return rows, i


def build_docx():
    md_lines = INPUT_PATH.read_text(encoding="utf-8").splitlines()
    doc = Document()

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)
    for style_name, size, color in [
        ("Heading 1", 16, (31, 78, 121)),
        ("Heading 2", 14, (46, 116, 181)),
        ("Heading 3", 12, (31, 78, 121)),
    ]:
        styles[style_name].font.name = "Arial"
        styles[style_name].font.size = Pt(size)
        styles[style_name].font.bold = True
        styles[style_name].font.color.rgb = RGBColor(*color)

    title = doc.add_paragraph()
    run = title.add_run("Nội dung chỉnh sửa report: BoW là mô hình chính")
    set_run_font(run, size=18, bold=True, color=(31, 78, 121))
    title.paragraph_format.space_after = Pt(10)

    subtitle = doc.add_paragraph()
    run = subtitle.add_run("Bản nháp cập nhật sau thực nghiệm full rebuild ngày 01/06/2026")
    set_run_font(run, size=11, italic=True, color=(90, 90, 90))
    subtitle.paragraph_format.space_after = Pt(12)

    i = 0
    in_code = False
    code_lines = []
    while i < len(md_lines):
        line = md_lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped.startswith("|"):
            table_rows, i = parse_table(md_lines, i)
            if table_rows:
                add_table(doc, table_rows)
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
        elif stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=3)
        elif stripped.startswith("- "):
            para = doc.add_paragraph(style="List Bullet")
            run = para.add_run(stripped[2:])
            set_run_font(run)
        elif stripped[0:2].isdigit() and ". " in stripped[:4]:
            para = doc.add_paragraph(style="List Number")
            run = para.add_run(stripped.split(". ", 1)[1])
            set_run_font(run)
        else:
            para = doc.add_paragraph()
            run = para.add_run(stripped)
            set_run_font(run)

        i += 1

    doc.add_page_break()
    note = doc.add_paragraph()
    run = note.add_run("Ghi chú: Bản này là draft để rà soát nội dung trước khi hợp nhất vào report chính.")
    set_run_font(run, size=10, italic=True, color=(90, 90, 90))

    doc.save(OUTPUT_PATH)
    print("created")


if __name__ == "__main__":
    build_docx()
