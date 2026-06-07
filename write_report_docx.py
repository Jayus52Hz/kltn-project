from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


INPUT_PATH = "Report KLTN - 22133056 - Nguyen Quoc Thinh - revised dataset architecture.docx"
OUTPUT_PATH = "Report KLTN - 22133056 - Nguyen Quoc Thinh - revised model theory.docx"


def insert_paragraph_after(paragraph, text="", style_name=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style_name:
        new_para.style = style_name
    if text:
        new_para.add_run(text)
    return new_para


def set_text(paragraph, text):
    paragraph.clear()
    paragraph.add_run(text)


def replace_range(doc, start_idx, end_idx, items):
    start = doc.paragraphs[start_idx]
    end = doc.paragraphs[end_idx]

    removing = []
    cursor = start._p.getnext()
    while cursor is not None and cursor is not end._p:
        removing.append(cursor)
        cursor = cursor.getnext()

    for element in removing:
        element.getparent().remove(element)

    anchor = start
    for style_name, text in items:
        anchor = insert_paragraph_after(anchor, text, style_name)


def append_references_after_prefix(doc, prefix, references):
    anchor = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            anchor = paragraph
            break
    if anchor is None:
        raise ValueError(f"Reference anchor not found: {prefix}")
    for ref in references:
        anchor = insert_paragraph_after(anchor, ref, "Normal")


phase1 = Document(INPUT_PATH)

# Rewrite Chapter 2 model sections from bottom to top so original indexes stay valid.
replace_range(
    phase1,
    653,
    690,
    [
        (
            "Normal",
            "Mô hình thứ hai trong notebook là roberta-base fine-tuned cho multi-label classification. Việc triển khai sử dụng hệ sinh thái Hugging Face và PyTorch: AutoTokenizer để token hóa transcript, AutoModelForSequenceClassification để nạp checkpoint pretrained và gắn classification head, Trainer để điều phối vòng huấn luyện, Datasets để chuẩn hóa dữ liệu đầu vào, còn PyTorch thực thi forward pass, backward pass và hàm loss. [13], [14], [15], [20], [21], [22]",
        ),
        (
            "Normal",
            "Về mặt lý thuyết, RoBERTa là mô hình encoder dựa trên Transformer. Transcript được chuyển thành token IDs, ánh xạ thành embeddings rồi đi qua nhiều tầng self-attention. Ở mỗi tầng, mô hình tính quan hệ giữa từng token với phần còn lại của chuỗi để tạo biểu diễn ngữ cảnh hóa. Nhờ cơ chế này, mô hình có thể hiểu tốt hơn các tình huống hội thoại mà ý nghĩa phụ thuộc vào toàn bộ ngữ cảnh thay vì chỉ dựa trên từ khóa cục bộ.",
        ),
        (
            "Normal",
            "Đề tài chọn hướng multi-label vì một transcript có thể mang nhiều call_code cùng lúc. Do đó đầu ra của mô hình là vector logits có số chiều bằng số lượng nhãn, thay vì một xác suất duy nhất cho một lớp. Các logits được đưa qua sigmoid độc lập theo từng chiều và so với ngưỡng dự đoán để quyết định nhãn nào được giữ lại.",
        ),
        ("Heading 4", "2.6.3.1. Cấu hình baseline dùng để so sánh"),
        (
            "Normal",
            "Baseline trong notebook là CountVectorizer + Logistic Regression theo chiến lược One-vs-Rest. CountVectorizer chuyển transcript thành ma trận số lần xuất hiện token/ngram [16]; OneVsRestClassifier huấn luyện một bộ phân loại nhị phân cho từng nhãn. Cấu hình này tạo ra mốc so sánh đơn giản nhưng có khả năng triển khai rất nhẹ trong môi trường Spark.",
        ),
        ("Report Code", "counter = CountVectorizer(max_features=5000, ngram_range=(1, 2), stop_words='english')"),
        ("Report Code", "X_train_count = counter.fit_transform(full_train_df['call_transcript'])"),
        ("Report Code", "base_lr = LogisticRegression(solver='liblinear', class_weight='balanced', random_state=42)"),
        ("Report Code", "counting_model = OneVsRestClassifier(base_lr)"),
        ("Report Code", "counting_model.fit(X_train_count, y_train_full)"),
        ("Heading 4", "2.6.3.2. Cấu hình RoBERTa fine-tuned"),
        (
            "Normal",
            "Checkpoint nền của mô hình là roberta-base. Theo paper RoBERTa [12], đây là biến thể được tối ưu hóa từ BERT; còn theo model card của Hugging Face [13], checkpoint này có thể được nạp trực tiếp bằng from_pretrained() để phục vụ downstream classification. Trong notebook của đề tài, mô hình được cấu hình với num_labels bằng số lượng call codes và problem_type='multi_label_classification' để phù hợp với bản chất dữ liệu.",
        ),
        (
            "Normal",
            "Dữ liệu huấn luyện được chuyển từ pandas sang Dataset của Hugging Face. Các nhãn call_code được mã hóa bằng MultiLabelBinarizer thành vector multi-hot 32 chiều. Transcript sau đó được tokenizer của roberta-base xử lý với truncation, padding và max_length=512 trước khi đưa vào mô hình. [17], [22]",
        ),
        ("Report Code", "model = AutoModelForSequenceClassification.from_pretrained("),
        ("Report Code", "    'roberta-base',"),
        ("Report Code", "    num_labels=len(mlb.classes_),"),
        ("Report Code", "    problem_type='multi_label_classification',"),
        ("Report Code", ")"),
        ("Report Code", "train_dataset = Dataset.from_pandas(train_df[['call_transcript', 'labels']])"),
        ("Report Code", "valid_dataset = Dataset.from_pandas(valid_df[['call_transcript', 'labels']])"),
        (
            "Report Code",
            "encoded = tokenizer(batch['call_transcript'], truncation=True, padding='max_length', max_length=512)",
        ),
        ("Report Caption", "Hình 2.6. So sánh baseline BoW và RoBERTa trong bài toán multi-label call_code"),
        ("Heading 4", "2.6.3.3. Hàm loss, ngưỡng dự đoán và ý nghĩa multi-label"),
        (
            "Normal",
            "Với multi-label classification, mỗi nhãn được xem như một bài toán nhị phân độc lập. Vì vậy mô hình không dùng softmax mà dùng sigmoid cho từng logit. Đề tài sử dụng BCEWithLogitsLoss; theo tài liệu PyTorch [15], hàm này kết hợp sigmoid và binary cross-entropy trong một lớp duy nhất, đồng thời ổn định số học hơn so với việc tách riêng hai bước.",
        ),
        (
            "Normal",
            "Trong notebook, labels được ép về float32 trước khi tính loss, và một custom Trainer được dùng để bảo đảm labels tương thích với BCEWithLogitsLoss. Ở bước suy luận, xác suất từng nhãn được so với ngưỡng 0.5. Nếu xác suất vượt ngưỡng, nhãn đó được gán cho transcript. Cơ chế này cho phép mô hình giữ lại đồng thời nhiều nhãn phù hợp trên cùng một cuộc gọi.",
        ),
        ("Report Code", "loss = torch.nn.BCEWithLogitsLoss()(outputs.logits, labels.float())"),
        ("Report Code", "probs = torch.sigmoid(torch.tensor(logits)).numpy()"),
        ("Report Code", "preds = (probs >= 0.5).astype(int)"),
        ("Report Code", "predicted_labels = mlb.inverse_transform(preds)"),
        ("Report Caption", "Hình 2.7. Cấu hình huấn luyện RoBERTa và luồng đánh giá train/valid/test"),
        ("Heading 4", "2.6.3.4. Kết quả so sánh và quyết định chọn mô hình"),
        (
            "Normal",
            "Trên tập test, RoBERTa đạt F1 Micro 73,46% so với 70,16% của BoW, tăng 3,30 điểm phần trăm. Quan trọng hơn, Precision của RoBERTa đạt 80,39%, cao hơn 14,47 điểm phần trăm so với baseline 65,92%. Dù Recall của BoW cao hơn trong lần thử nghiệm này, RoBERTa vẫn phù hợp hơn với mục tiêu downstream vì giúp giảm nguy cơ gán nhầm nhãn dương tính trong pipeline phân tích.",
        ),
        (
            "Normal",
            "Khi call_code do mô hình sinh ra trở thành nguồn dữ liệu cho Gold layer và BigQuery, false positive có thể làm sai lệch các KPI outcome như successful sale, hard rejection hoặc do-not-call. Vì vậy đề tài ưu tiên mô hình có Precision cao hơn, ngay cả khi phải đánh đổi một phần Recall. Đây là lý do thực tiễn khiến RoBERTa được chọn làm mô hình chính thay vì chỉ dừng ở baseline BoW.",
        ),
        ("Heading 4", "2.6.3.5. Tích hợp RoBERTa vào Silver job"),
        (
            "Normal",
            "Repository đã được đồng bộ theo quyết định này: silver_job.py mặc định đọc NLP_MODEL_TYPE=roberta, load thư mục roberta_saved cùng label_classes.json và sinh call_code trực tiếp từ transcript bằng Pandas UDF. Trong trường hợp cần chạy cấu hình nhẹ để demo hoặc đối chiếu, hệ thống vẫn giữ baseline BoW như một tùy chọn fallback, nhưng đó không còn là nguồn sự thật cho bước phân tích downstream.",
        ),
        ("Report Code", "NLP_MODEL_TYPE = os.getenv(\"NLP_MODEL_TYPE\", \"roberta\").lower()"),
        ("Report Code", "encoded = tokenizer(batch_texts, truncation=True, padding=True, max_length=512, return_tensors='pt')"),
        ("Report Code", "probs = torch.sigmoid(model(**encoded).logits)"),
        ("Report Code", "pred = [labels[i] for i, score in enumerate(row) if score >= ROBERTA_THRESHOLD]"),
        ("Report Caption", "Hình 2.8. Luồng RoBERTa inference ghi call_code vào Silver"),
        ("Report Caption", "Bảng 2.6. Kết quả mô hình NLP theo notebook NLP_model.ipynb"),
        ("Report Caption", "Hình 2.5. Pipeline huấn luyện và triển khai NLP trong tầng Silver"),
    ],
)

replace_range(
    phase1,
    650,
    653,
    [
        (
            "Normal",
            "Mô hình baseline dùng CountVectorizer với max_features=5000, ngram_range=(1,2), stop_words='english'. Sau bước biểu diễn đặc trưng, đề tài dùng OneVsRestClassifier với LogisticRegression solver='liblinear', class_weight='balanced' để học từng nhãn theo chiến lược one-vs-rest. Cách làm này phù hợp để thiết lập một mốc so sánh ban đầu vì dễ huấn luyện, dễ giải thích và chi phí tính toán thấp. [11], [16], [17]",
        ),
        (
            "Normal",
            "Tuy nhiên, baseline BoW phụ thuộc mạnh vào tần suất token/ngram nên chưa mô hình hóa tốt ngữ cảnh hội thoại. Nó khó biểu diễn được các hiện tượng như ngắt lời, chuyển ý, phản đối có điều kiện hoặc sắc thái câu nói. Vì vậy baseline được giữ với vai trò đối chứng chất lượng và engineering fallback, không phải mô hình được chọn làm nguồn tín hiệu chính cho pipeline phân tích.",
        ),
    ],
)

replace_range(
    phase1,
    647,
    650,
    [
        (
            "Normal",
            "Mô hình NLP trong đề tài không phải mô hình sinh dữ liệu. Nó là mô hình phân loại multi-label, nhận đầu vào là call_transcript và dự đoán danh sách call_code. Trong giai đoạn huấn luyện, call_code gốc của bộ dữ liệu tổng hợp chỉ đóng vai trò nhãn học có giám sát. Sau khi mô hình hoàn tất huấn luyện, pipeline vận hành không dùng lại call_code gốc ở các bước downstream.",
        ),
        (
            "Normal",
            "Về mặt kiến trúc, source of truth cho các bước phân tích sau này là call_code do mô hình sinh ra từ transcript. Điều này giúp pipeline phản ánh đúng bài toán thực tế: dữ liệu vận hành có transcript mới, mô hình suy luận nhãn nghiệp vụ, và các tầng Silver, Gold, BigQuery sử dụng kết quả dự đoán đó để phục vụ BI.",
        ),
        (
            "Normal",
            "Mục tiêu tích hợp NLP vào Silver là biến transcript phi cấu trúc thành tín hiệu phân tích có cấu trúc nhưng vẫn giữ ranh giới bảo mật. Transcript chỉ được dùng ở nơi cần cho inference, không được đẩy lên tầng Gold hoặc BigQuery; ngược lại, call_code sinh ra từ mô hình mới là trường được đưa tiếp vào các bước phân tích và tổng hợp KPI.",
        ),
    ],
)

append_references_after_prefix(
    phase1,
    "[17]",
    [
        "[18] Vaswani et al., Attention Is All You Need, https://arxiv.org/abs/1706.03762",
        "[19] Devlin et al., BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding, https://arxiv.org/abs/1810.04805",
        "[20] Hugging Face Transformers, AutoModelForSequenceClassification and RoBERTa documentation, https://huggingface.co/docs/transformers/main/model_doc/auto",
        "[21] Hugging Face Transformers, Trainer documentation, https://huggingface.co/docs/transformers/main_classes/trainer",
        "[22] Hugging Face Datasets, Use with Pandas, https://huggingface.co/docs/datasets/use_with_pandas",
    ],
)

for paragraph in phase1.paragraphs:
    text = paragraph.text.strip()
    if text == "Hình 2.8. Luồng RoBERTa inference ghi call_code_predicted vào Silver":
        set_text(paragraph, "Hình 2.8. Luồng RoBERTa inference ghi call_code vào Silver")
    elif text == "Hình 4.7. Đối chiếu call_code_original và call_code_predicted":
        set_text(paragraph, "Hình 4.7. Vai trò nhãn call_code gốc trong training và call_code do mô hình sinh trong production")
    elif "call_code_predicted" in text:
        set_text(paragraph, text.replace("call_code_predicted", "call_code"))

phase1.save(OUTPUT_PATH)


phase2 = Document(OUTPUT_PATH)

# Rewrite Chapter 1 model theory from bottom to top so indexes stay valid.
replace_range(
    phase2,
    481,
    497,
    [
        (
            "Normal",
            "Bài toán được đánh giá bằng các chỉ số F1 Micro, F1 Macro, Precision, Recall và Exact Match. F1 Micro phản ánh chất lượng tổng thể trên toàn bộ nhãn; F1 Macro cho thấy độ cân bằng giữa các nhãn; Precision phản ánh mức độ hạn chế dương tính giả; Recall phản ánh khả năng phát hiện đúng các nhãn thật; Exact Match là tỷ lệ dự đoán đúng hoàn toàn toàn bộ tập nhãn của một transcript.",
        ),
        (
            "Normal",
            "Trong đề tài, F1 Micro và Precision được ưu tiên hơn trong kết luận lựa chọn mô hình. Lý do là call_code do mô hình sinh ra sẽ tiếp tục được dùng làm nguồn sự thật cho Gold layer, BigQuery và dashboard. Nếu mô hình sinh nhiều nhãn dương tính giả, các chỉ số outcome, hành vi khách hàng hoặc hiệu quả chiến dịch có thể bị sai lệch đáng kể. Vì vậy, một mô hình có Precision cao hơn sẽ phù hợp hơn cho môi trường phân tích nghiệp vụ.",
        ),
        ("Normal", "Đoạn mã 1.11. Trích đoạn hiện thực cho tích hợp NLP/ML vào ETL"),
        ("Report Code", "NLP_MODEL_TYPE = os.getenv(\"NLP_MODEL_TYPE\", \"roberta\").lower()"),
        ("Report Code", "ROBERTA_PATH = os.path.join(MODELS_PATH, \"roberta_saved\")"),
        ("Report Code", "LABELS_PATH = os.path.join(MODELS_PATH, \"label_classes.json\")"),
        ("Report Code", "@F.pandas_udf(ArrayType(StringType()))"),
        ("Report Code", "def predict_call_codes(transcripts: pd.Series) -> pd.Series:"),
        ("Report Code", "    tokenizer, model, labels = _load_roberta_once()"),
        (
            "Report Code",
            "    encoded = tokenizer(list(transcripts.fillna('')), truncation=True, padding=True, max_length=512, return_tensors='pt')",
        ),
        ("Report Code", "    with torch.no_grad():"),
        ("Report Code", "        probs = torch.sigmoid(model(**encoded).logits).cpu().numpy()"),
        (
            "Report Code",
            "    return pd.Series([[labels[i] for i, score in enumerate(row) if score >= 0.5] for row in probs])",
        ),
    ],
)

replace_range(
    phase2,
    479,
    481,
    [
        (
            "Normal",
            "Mô hình chính của đề tài là RoBERTa fine-tuned cho multi-label sequence classification. Về nguồn gốc, RoBERTa là biến thể được tối ưu hóa mạnh hơn của BERT và được công bố trong nghiên cứu RoBERTa: A Robustly Optimized BERT Pretraining Approach của Liu et al. [12]. Xa hơn nữa, RoBERTa kế thừa nền tảng kiến trúc Transformer encoder từ công trình Attention Is All You Need của Vaswani et al. [18], trong đó self-attention là cơ chế cốt lõi để mô hình học quan hệ giữa các token trong chuỗi.",
        ),
        (
            "Normal",
            "Nếu BERT giới thiệu hướng biểu diễn ngôn ngữ hai chiều bằng Masked Language Modeling [19], thì RoBERTa chứng minh rằng hiệu quả downstream còn phụ thuộc rất mạnh vào cách pretraining. Theo paper gốc, RoBERTa cải thiện BERT bằng cách huấn luyện lâu hơn, dùng batch lớn hơn, dùng nhiều dữ liệu hơn, áp dụng dynamic masking và loại bỏ nhiệm vụ Next Sentence Prediction. Những thay đổi này giúp mô hình học biểu diễn ngữ cảnh ổn định hơn cho các tác vụ hiểu văn bản.",
        ),
        (
            "Normal",
            "Trong đề tài, checkpoint nền được sử dụng là roberta-base thông qua thư viện Hugging Face Transformers. Việc huấn luyện và triển khai được thực hiện bằng AutoTokenizer và AutoModelForSequenceClassification thay vì tự xây thủ công từng lớp Transformer. AutoTokenizer chịu trách nhiệm token hóa transcript bằng tokenizer của roberta-base; AutoModelForSequenceClassification nạp trọng số pretrained, gắn classification head và cấu hình mô hình cho bài toán multi-label. [13], [14], [20]",
        ),
        (
            "Normal",
            "Về cơ chế hoạt động, transcript trước hết được tokenizer chuyển thành token IDs, sau đó ánh xạ sang embeddings và đưa qua nhiều tầng Transformer encoder. Ở mỗi tầng, self-attention tính mức độ liên quan giữa từng token với toàn bộ token còn lại trong câu, từ đó tạo biểu diễn ngữ cảnh hóa. Đầu ra cuối cùng của mô hình là vector logits có kích thước bằng số lượng nhãn. Các logits này không đi qua softmax mà đi qua sigmoid độc lập theo từng chiều, vì một transcript có thể mang nhiều call_code cùng lúc.",
        ),
        (
            "Normal",
            "RoBERTa được chọn làm mô hình chính vì phù hợp hơn với tính chất dữ liệu hội thoại và kết quả thực nghiệm tốt hơn baseline ở các chỉ số quan trọng. So với BoW, RoBERTa hiểu ngữ cảnh tốt hơn, giảm phụ thuộc vào từ khóa đơn lẻ, đồng thời đạt F1 Micro và Precision cao hơn trên tập test. Trong ngữ cảnh Lakehouse của đề tài, call_code do mô hình sinh sẽ trở thành đầu vào phân tích downstream, nên việc hạn chế false positive là rất quan trọng.",
        ),
    ],
)

replace_range(
    phase2,
    477,
    479,
    [
        (
            "Normal",
            "Baseline được sử dụng trong đề tài là Bag-of-Words kết hợp Logistic Regression theo chiến lược One-vs-Rest. CountVectorizer biến transcript thành vector đặc trưng dựa trên tần suất token và n-gram, sau đó mỗi nhãn được học bởi một bộ phân loại nhị phân Logistic Regression riêng. Cách tiếp cận này có ưu điểm là đơn giản, dễ giải thích, huấn luyện nhanh và có thể triển khai nhẹ trong Spark bằng artifact joblib.",
        ),
        (
            "Normal",
            "Tuy nhiên, BoW có giới hạn quan trọng: mô hình dựa chủ yếu vào tần suất từ/ngram nên không nắm bắt tốt ngữ cảnh dài, sắc thái hội thoại, sự phủ định, chuyển ý hoặc quan hệ giữa các câu trong transcript. Với dữ liệu telesales, các yếu tố như ngắt lời, phản đối, đồng ý có điều kiện hoặc yêu cầu không gọi lại thường phụ thuộc vào ngữ cảnh hơn là chỉ phụ thuộc vào vài từ khóa rời rạc. Vì vậy BoW phù hợp làm baseline kỹ thuật, nhưng chưa phải lựa chọn tối ưu cho mô hình chính của pipeline. [11], [16], [17]",
        ),
    ],
)

replace_range(
    phase2,
    475,
    477,
    [
        (
            "Normal",
            "Trong đề tài, mô hình NLP không được dùng để sinh dữ liệu, không được dùng để viết báo cáo, mà được dùng để biến transcript cuộc gọi thành tín hiệu phân tích có cấu trúc. Cụ thể, mô hình nhận đầu vào là call_transcript và sinh ra danh sách call_code phản ánh nội dung, trạng thái và hành vi xuất hiện trong cuộc gọi. Kết quả này được dùng ở tầng Silver và tiếp tục đóng vai trò source of truth cho các bước phân tích downstream.",
        ),
        (
            "Normal",
            "Bài toán ở đây là multi-label text classification chứ không phải single-label classification. Một transcript có thể đồng thời chứa nhiều tín hiệu nghiệp vụ, ví dụ vừa có PRODUCT_PITCH, vừa có OBJECTION_HANDLING, vừa có FOLLOW_UP_EMAIL_REQUESTED. Vì vậy mô hình phải dự đoán đồng thời nhiều nhãn trên cùng một mẫu thay vì chỉ chọn một lớp duy nhất.",
        ),
        (
            "Normal",
            "Về mặt kiến trúc dữ liệu, bước NLP giúp chuyển dữ liệu phi cấu trúc thành trường phân tích có thể tổng hợp trong Lakehouse. Nếu không có bước này, transcript chỉ tồn tại như văn bản thô và dashboard khó phản ánh được bản chất cuộc gọi. Khi có call_code do mô hình sinh ra, hệ thống có thể xây dựng các cờ outcome, nhóm hành vi và KPI vận hành mà không cần đẩy transcript lên tầng Gold hoặc BigQuery.",
        ),
    ],
)

for paragraph in phase2.paragraphs:
    text = paragraph.text.strip()
    if text == "Hình 2.8. Luồng RoBERTa inference ghi call_code_predicted vào Silver":
        set_text(paragraph, "Hình 2.8. Luồng RoBERTa inference ghi call_code vào Silver")
    elif text == "Hình 4.7. Đối chiếu call_code_original và call_code_predicted":
        set_text(paragraph, "Hình 4.7. Vai trò nhãn call_code gốc trong training và call_code do mô hình sinh trong production")

phase2.save(OUTPUT_PATH)
print(OUTPUT_PATH)
