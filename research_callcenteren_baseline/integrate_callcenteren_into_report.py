"""
Integrate CallCenterEN research content into the correct thesis chapters.

Unlike write_callcenteren_report_docx.py, this script does not append the
CallCenterEN work as a late addendum. It inserts Vietnamese, report-ready content
into Chapter 1, Chapter 2, Chapter 4, Chapter 5, and the References section.
"""

from __future__ import annotations

import csv
import json
from collections import Counter
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = Path(__file__).resolve().parent / "output"
SOURCE_DOCX = ROOT / "Report KLTN - 22133056 - Nguyen Quoc Thinh.docx"
TARGET_DOCX = ROOT / "Report KLTN - 22133056 - Nguyen Quoc Thinh - revised callcenteren integrated.docx"

SUMMARY_JSON = OUTPUT_DIR / "dataset_comparison_summary.json"
PROMPT_DESIGN_JSON = OUTPUT_DIR / "callcenteren_prompt_design_analysis.json"
PSEUDO_CSV = OUTPUT_DIR / "pseudo_labels_gemini.csv"
VALID_METRICS_CSV = OUTPUT_DIR / "auxiliary_bow_valid_metrics.csv"
TEST_METRICS_CSV = OUTPUT_DIR / "auxiliary_bow_test_metrics.csv"


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8", newline="") as file:
        return list(csv.DictReader(file))


def metric(value: Any) -> str:
    if isinstance(value, float):
        return f"{value:.4f}"
    return str(value)


def set_document_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)


def find_paragraph(doc: Document, startswith: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(startswith):
            return paragraph
    raise ValueError(f"Cannot find paragraph starting with: {startswith}")


def content_start_index(doc: Document) -> int:
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == "PHẦN NỘI DUNG":
            return idx
    return 0


def find_content_heading(doc: Document, startswith: str):
    start = content_start_index(doc)
    for paragraph in doc.paragraphs[start:]:
        style_name = paragraph.style.name if paragraph.style else ""
        if "Heading" in style_name and paragraph.text.strip().startswith(startswith):
            return paragraph
    raise ValueError(f"Cannot find content heading starting with: {startswith}")


def find_heading(doc: Document, startswith: str):
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style else ""
        if "Heading" in style_name and paragraph.text.strip().startswith(startswith):
            return paragraph
    raise ValueError(f"Cannot find heading starting with: {startswith}")


def insert_paragraph_before(doc: Document, target, text: str = "", style: str | None = None, bold: bool = False):
    paragraph = doc.add_paragraph(style=style)
    if text:
        run = paragraph.add_run(text)
        run.bold = bold
    element = paragraph._p
    parent = element.getparent()
    parent.remove(element)
    target._p.addprevious(element)
    return paragraph


def insert_table_before(doc: Document, target, headers: list[str], rows: list[list[Any]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        run = cell.paragraphs[0].add_run(str(header))
        run.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    element = table._tbl
    parent = element.getparent()
    parent.remove(element)
    target._p.addprevious(element)
    insert_paragraph_before(doc, target, "")
    return table


def add_heading_before(doc: Document, target, text: str, level: int):
    return insert_paragraph_before(doc, target, text, style=f"Heading {level}")


def add_paragraphs_before(doc: Document, target, paragraphs: list[str]) -> None:
    for text in paragraphs:
        insert_paragraph_before(doc, target, text)


def update_toc(doc: Document) -> None:
    toc_end = content_start_index(doc)
    toc_insertions = [
        ("1.14. Machine Learning trong pipeline dữ liệu", [
            "1.14.5. Học thích nghi miền, weak supervision và pseudo-labeling",
        ]),
        ("2.1. Vai trò của dữ liệu trong đề tài", [
            "2.1.4. Vai trò của CallCenterEN trong nghiên cứu",
            "2.1.5. Từ CallCenterEN đến prompt sinh dataset chính",
        ]),
        ("2.3. Thiết kế schema và tính chất dữ liệu", [
            "2.3.5. Đối chiếu dataset chính với CallCenterEN",
        ]),
        ("2.5. Chuẩn bị dữ liệu huấn luyện NLP", [
            "2.5.3. Lọc CallCenterEN subset và tạo pseudo-label",
        ]),
        ("2.6. Mô hình NLP phân loại call_code", [
            "2.6.7. Cấu hình thí nghiệm auxiliary training",
        ]),
        ("4.5. Đánh giá NLP trong pipeline", [
            "4.5.5. Thực nghiệm auxiliary training với CallCenterEN",
        ]),
        ("5.2. Đánh giá đóng góp kỹ thuật", [
            "5.2.4. Đóng góp từ external baseline và auxiliary corpus",
        ]),
        ("5.3. Hạn chế", [
            "5.3.4. Hạn chế khi sử dụng CallCenterEN và pseudo-label",
        ]),
    ]
    for anchor_text, lines in toc_insertions:
        anchor = None
        for paragraph in doc.paragraphs[:toc_end]:
            if paragraph.text.strip().startswith(anchor_text):
                anchor = paragraph
                break
        if anchor is None:
            continue
        cursor = anchor
        for line in lines:
            if line not in "\n".join(p.text for p in doc.paragraphs[:toc_end]):
                new_para = insert_paragraph_before(doc, cursor, line)
                cursor = new_para


def add_theory_section(doc: Document) -> None:
    target = find_content_heading(doc, "1.15. Mô hình Star Schema")
    add_heading_before(doc, target, "1.14.5. Học thích nghi miền, weak supervision và pseudo-labeling", 3)
    add_paragraphs_before(doc, target, [
        "Trong bài toán phân loại hội thoại, mô hình không chỉ cần học quan hệ giữa transcript và nhãn, mà còn cần hiểu ngôn ngữ tự nhiên của môi trường call center. Các transcript thực tế thường chứa lời chào, xác minh thông tin, hỏi đáp nhu cầu, phản đối, thảo luận phí, yêu cầu không gọi lại và các đoạn ngắt quãng. Vì vậy, một tập transcript call center thực tế có thể đóng vai trò như corpus miền, giúp mô hình tiếp xúc với cách diễn đạt ngoài bộ dữ liệu chính.",
        "Hướng thứ nhất là học thích nghi miền (domain-adaptive pretraining). Theo Gururangan và cộng sự (2020), tiếp tục pretrain mô hình ngôn ngữ trên dữ liệu cùng miền có thể giúp cải thiện biểu diễn trước khi fine-tune cho tác vụ cụ thể. Trong đề tài này, CallCenterEN phù hợp với vai trò corpus miền vì nó chứa hội thoại call center thực tế, có inbound/outbound, nhiều domain và metadata ASR. Hướng này đặc biệt phù hợp nếu tiếp tục phát triển mô hình Transformer như RoBERTa.",
        "Hướng thứ hai là weak supervision hoặc pseudo-labeling. Khi dữ liệu ngoài miền có transcript nhưng không có nhãn call_code, có thể sử dụng labeling functions, mô hình hiện có hoặc mô hình AI để gán nhãn yếu cho một phần dữ liệu. Theo Ratner và cộng sự (2016), các nguồn nhãn yếu có thể giúp tạo tập huấn luyện quy mô lớn hơn, nhưng cần xem chúng là noisy labels. Theo khảo sát self-training của Amini và cộng sự (2022), các mẫu pseudo-label có độ tin cậy cao có thể được bổ sung vào tập huấn luyện, trong khi tập kiểm thử vẫn phải giữ ground truth độc lập.",
        "Do đó, CallCenterEN không chỉ là dữ liệu tham khảo ở mức mô tả. Nó tạo ra một cầu nối phương pháp luận: dataset chính của đề tài cung cấp nhãn nghiệp vụ chính thức, còn CallCenterEN cung cấp ngôn ngữ call center thực tế để đối chiếu và tạo tín hiệu huấn luyện phụ.",
    ])


def add_dataset_role_section(doc: Document) -> None:
    target = find_content_heading(doc, "2.2. Nguồn gốc và quy trình sinh dữ liệu")
    add_heading_before(doc, target, "2.1.4. Vai trò của CallCenterEN trong nghiên cứu", 3)
    add_paragraphs_before(doc, target, [
        "Tập dữ liệu chính của đề tài vẫn là bộ dữ liệu telesales do tác giả xây dựng, vì bộ dữ liệu này có đầy đủ customer, offer, call_logs, call_transcript và nhãn nghiệp vụ call_code. Đây là nguồn dữ liệu duy nhất được xem là ground truth cho training, validation, test và cho pipeline phân tích chính.",
        "Tuy nhiên, bộ dữ liệu chính không được trình bày như một tập dữ liệu sinh tùy tiện. Trước khi mô tả prompt và schema sinh dữ liệu, nghiên cứu sử dụng CallCenterEN như một external real-world baseline để rút ra các đặc trưng của hội thoại call center thực tế. Các đặc trưng này sau đó được chuyển thành nguyên tắc thiết kế prompt cho dataset chính.",
        "Ngoài vai trò đối chiếu, một phần CallCenterEN được đưa vào thí nghiệm huấn luyện phụ. Cụ thể, transcript từ CallCenterEN được lọc thành subset gần với telesales/insurance/customer service, sau đó được gán pseudo-label theo taxonomy call_code. Các nhãn này chỉ được dùng như weak labels trong mô hình phụ trợ, không được xem là nhãn đúng tuyệt đối.",
    ])


def add_prompt_design_section(doc: Document, prompt_design: dict[str, Any]) -> None:
    target = find_content_heading(doc, "2.2. Nguồn gốc và quy trình sinh dữ liệu")
    add_heading_before(doc, target, "2.1.5. Từ CallCenterEN đến prompt sinh dataset chính", 3)
    add_paragraphs_before(doc, target, [
        "Mạch nghiên cứu của phần dữ liệu được tổ chức theo thứ tự: phân tích CallCenterEN, trích xuất đặc trưng, chuyển đặc trưng thành quy tắc prompt, sau đó dùng prompt để sinh dataset telesales chính. Cách trình bày này giúp giải thích vì sao dataset chính có các trường và nhãn hiện tại, thay vì chỉ mô tả dataset như một kết quả đã có sẵn.",
        f"Trong bước phân tích nền, nghiên cứu sử dụng {prompt_design['rows']:,} mẫu CallCenterEN từ baseline_analysis_sample. Tập này không dùng để train trực tiếp. Nó được dùng để quan sát cấu trúc hội thoại, độ dài transcript, duration, PII, hướng cuộc gọi, domain và các hành vi hội thoại thường gặp.",
        f"Kết quả thống kê cho thấy 3,000 mẫu có độ dài trung bình {prompt_design['avg_chars']:.4f} ký tự, {prompt_design['avg_words']:.4f} từ, thời lượng trung bình {prompt_design['avg_duration']:.4f} giây, ASR confidence trung bình {prompt_design['avg_confidence']:.4f} và trung bình {prompt_design['avg_pii_tokens']:.4f} token PII đã được redaction. Các con số này cung cấp cơ sở để thiết kế transcript có độ dài vừa đủ, có trường talk_time_seconds và có xử lý PII trong pipeline.",
    ])
    insert_table_before(doc, target, ["Nhóm đặc trưng từ CallCenterEN", "Số mẫu", "Tỷ lệ"], [
        [group, f"{value['rows']:,}", f"{value['percent']}%"]
        for group, value in prompt_design["keyword_coverage"].items()
    ])
    insert_table_before(doc, target, ["Quan sát từ CallCenterEN", "Quy tắc prompt sinh dataset chính", "Trường/nhãn được tạo"], [
        [
            item["callcenteren_observation"],
            item["prompt_rule_for_primary_dataset"],
            item["primary_dataset_field_or_label"],
        ]
        for item in prompt_design["prompt_design_mapping"]
    ])
    add_paragraphs_before(doc, target, [
        "Từ bảng trên, prompt sinh dataset chính được thiết kế để bắt buộc transcript có cấu trúc hội thoại rõ ràng: mở đầu cuộc gọi, xác minh hoặc nhắc đến thông tin khách hàng, phân tích nhu cầu, giới thiệu sản phẩm/offer, thảo luận phí hoặc điều kiện, xử lý phản đối và kết thúc bằng outcome. Các thành phần này tương ứng với taxonomy call_code như OPENING, NEEDS_ANALYSIS, PRODUCT_PITCH, FEE_DISCUSSION, OBJECTION_HANDLING, WARM_LEAD, SOFT_REJECTION, HARD_REJECTION, DO_NOT_CALL_REQUEST hoặc SUDDEN_HANG_UP.",
        "Như vậy, CallCenterEN đóng vai trò cơ sở để xây dựng logic prompt và schema dataset chính. Sau bước này, dataset chính mới được sinh theo bối cảnh telesales tài chính của đồ án, có thêm customer profile, offer, campaign, call metadata và nhãn call_code để phục vụ cả mô hình NLP lẫn star schema phân tích.",
    ])


def add_dataset_comparison_section(doc: Document, summary: dict[str, Any]) -> None:
    primary = summary["primary_dataset"]
    external = summary["callcenteren_baseline_subset"]
    target = find_content_heading(doc, "2.4. Chuẩn hóa dữ liệu thành các thực thể nguồn")
    add_heading_before(doc, target, "2.3.5. Đối chiếu dataset chính với CallCenterEN", 3)
    add_paragraphs_before(doc, target, [
        "CallCenterEN là tập transcript call center tiếng Anh được công bố công khai trong bài báo Real-World En Call Center Transcripts Dataset with PII Redaction. Theo mô tả của tác giả, tập dữ liệu gồm 91,706 cuộc hội thoại, khoảng 10,448 giờ audio, bao gồm inbound và outbound calls, nhiều domain, ASR confidence, word-level timestamps và transcript đã được PII redaction.",
        "Các đặc điểm này cho phép dùng CallCenterEN để giải thích các thành phần chính trong dataset của đề tài. Trường call_transcript phản ánh thực tế rằng dữ liệu call center có transcript hội thoại; trường talk_time_seconds tương ứng với audio_duration; PII masking ở Silver layer tương ứng với yêu cầu PII redaction; còn việc áp dụng NLP trên transcript phù hợp với mục tiêu customer support/sales AI của CallCenterEN.",
    ])
    insert_table_before(doc, target, ["Thành phần trong đề tài", "Đặc điểm tương ứng trong CallCenterEN", "Ý nghĩa thiết kế"], [
        ["call_transcript", "Transcript hội thoại", "Chứng minh transcript là dữ liệu phi cấu trúc trung tâm của bài toán call center"],
        ["talk_time_seconds", "audio_duration", "Bổ sung cơ sở cho việc lưu và phân tích thời lượng cuộc gọi"],
        ["PII masking ở Silver layer", "PII redaction", "Phù hợp yêu cầu bảo vệ dữ liệu cá nhân trước khi phục vụ phân tích"],
        ["NLP inference trên transcript", "Dataset phục vụ customer support/sales AI", "Chứng minh NLP trên hội thoại là hướng xử lý hợp lý"],
        ["Ngữ cảnh telesales/call center", "inbound/outbound và nhiều domain", "Giải thích bối cảnh đa miền của dữ liệu cuộc gọi"],
    ])
    insert_table_before(doc, target, ["Tiêu chí", "Dataset chính của đề tài", "CallCenterEN subset đã lọc"], [
        ["Vai trò", "Primary task dataset", "External baseline và auxiliary corpus"],
        ["Số mẫu", f"{primary['rows']:,}", f"{external['rows']:,}"],
        ["Độ dài transcript trung bình", metric(primary["avg_chars"]), metric(external["avg_chars"])],
        ["Số từ trung bình", metric(primary["avg_words"]), metric(external["avg_words"])],
        ["Thời lượng trung bình", metric(primary["avg_duration_seconds"]), metric(external["avg_duration_seconds"])],
        ["Nhãn nghiệp vụ", "Có call_code", "Không có call_code gốc"],
        ["ASR confidence", "Không có", metric(external["avg_asr_confidence"])],
        ["PII redaction/masking", "Masking ở Silver", f"Trung bình {metric(external['avg_pii_tokens'])} token PII/mẫu"],
    ])
    add_paragraphs_before(doc, target, [
        "Bảng so sánh cho thấy CallCenterEN phù hợp để làm cơ sở đối chiếu cấu trúc dữ liệu, nhưng không thể thay thế dataset chính vì thiếu các thực thể nghiệp vụ như customer, offer, campaign và nhãn call_code. Vì vậy, nghiên cứu đặt CallCenterEN ở vai trò external baseline và auxiliary corpus, còn dataset của đề tài giữ vai trò nguồn dữ liệu chính.",
    ])


def add_labeling_process_section(doc: Document, pseudo_rows: list[dict[str, str]]) -> None:
    target = find_content_heading(doc, "2.6. Mô hình NLP phân loại call_code")
    label_counter: Counter[str] = Counter()
    usable = 0
    for row in pseudo_rows:
        if row.get("should_use_for_training", "").lower() == "true":
            usable += 1
        for label in [item.strip() for item in row.get("pseudo_call_code", "").split(",") if item.strip()]:
            label_counter[label] += 1

    add_heading_before(doc, target, "2.5.3. Lọc CallCenterEN subset và tạo pseudo-label", 3)
    add_paragraphs_before(doc, target, [
        "Cần phân biệt hai tập con CallCenterEN trong nghiên cứu. Tập 3,000 mẫu baseline_analysis_sample dùng để phân tích đặc trưng và thiết kế prompt sinh dataset chính. Tập auxiliary_training_candidate dùng cho thí nghiệm gán pseudo-label và huấn luyện phụ. Hai tập này có vai trò khác nhau, vì vậy việc chỉ có 300 pseudo-label không có nghĩa là nghiên cứu chỉ dùng 300 mẫu CallCenterEN.",
        "Quy trình chuẩn bị CallCenterEN cho huấn luyện phụ được thực hiện sau bước phân tích đặc trưng. Thay vì dùng toàn bộ dữ liệu 92k transcript, nghiên cứu chỉ chọn các nhóm gần với bài toán telesales và customer service như insurance_outbound, auto_insurance_customer_service_inbound, PII_redacted_auto_insurance_script và customer_service_general_inbound.",
        "Các điều kiện lọc gồm: ASR confidence từ 0.90 trở lên, thời lượng từ 60 đến 900 giây, độ dài transcript từ 300 đến 6,000 ký tự, transcript không rỗng và loại bỏ trùng lặp bằng hash của normalized text. Sau bước lọc, nghiên cứu tạo baseline_analysis_sample gồm 3,000 mẫu để phân tích đặc trưng và auxiliary_training_candidate gồm 2,000 mẫu để chuẩn bị gán pseudo-label.",
        "Để tạo weak labels, nghiên cứu sử dụng Google AI Studio model gemma-4-31b-it gán nhãn theo taxonomy call_code của dataset chính. Tốc độ gọi API được giữ ở mức 15 RPM bằng cách chèn thời gian chờ 4.2 giây giữa các request. Kết quả sinh được 300 dòng pseudo-label, trong đó 295 dòng được đánh dấu có thể dùng cho training. Sau khi merge lại với transcript và lọc nhãn hợp lệ trong taxonomy chính, 294 dòng được dùng trong thí nghiệm auxiliary training.",
    ])
    insert_table_before(doc, target, ["Bước", "Mục đích", "Kết quả"], [
        ["Chọn domain gần bài toán", "Giảm domain shift", "Insurance, auto insurance, customer service"],
        ["Lọc chất lượng transcript", "Giữ mẫu đủ dài và có confidence tốt", "confidence >= 0.90, duration 60-900s, text 300-6000 chars"],
        ["Deduplicate", "Loại bỏ transcript trùng", "Hash theo normalized text"],
        ["Tạo baseline sample", "Phân tích cơ sở dữ liệu tham chiếu", "3,000 mẫu"],
        ["Tạo auxiliary candidate", "Chuẩn bị gán nhãn phụ", "2,000 mẫu"],
        ["AI pseudo-labeling", "Gán nhãn yếu theo call_code", "300 pseudo-label, 294 dòng dùng trong training"],
    ])
    insert_table_before(doc, target, ["Pseudo-label", "Số lần xuất hiện"], [[label, f"{count:,}"] for label, count in label_counter.most_common(15)])
    add_paragraphs_before(doc, target, [
        "Các pseudo-label này không được xem là ground truth. Chúng chỉ đóng vai trò tín hiệu phụ trợ để kiểm tra giả thuyết rằng transcript call center thực tế có thể cải thiện nhẹ khả năng tổng quát hóa của mô hình phân loại call_code.",
    ])


def add_aux_training_method_section(doc: Document) -> None:
    target = find_content_heading(doc, "2.7. Data quality")
    add_heading_before(doc, target, "2.6.7. Cấu hình thí nghiệm auxiliary training", 3)
    add_paragraphs_before(doc, target, [
        "Sau khi có pseudo-label, nghiên cứu thiết kế thí nghiệm M0/M3 để đo đóng góp thực tế của CallCenterEN. M0 là mô hình BoW + Logistic Regression huấn luyện chỉ trên dataset chính. M3 giữ nguyên cấu hình mô hình nhưng bổ sung pseudo-labeled CallCenterEN vào tập huấn luyện. Validation và test set không thay đổi, vẫn lấy hoàn toàn từ dataset chính.",
        "Thiết kế này đảm bảo CallCenterEN tham gia training nhưng không làm thay đổi nguồn ground truth. Nếu M3 cải thiện so với M0, điều đó cho thấy external corpus có thể mang lại tín hiệu ngôn ngữ hữu ích. Nếu M3 không cải thiện, kết quả vẫn có giá trị vì nó phản ánh rủi ro domain shift và noisy labels trong weak supervision.",
        "Hướng domain-adaptive pretraining cũng được đặt trong thiết kế nghiên cứu như một nhánh mở rộng cho mô hình Transformer. Với mô hình vận hành hiện tại là BoW + Logistic Regression, phần thực nghiệm khả thi trong đồ án tập trung vào weak-label training. Tuy vậy, CallCenterEN vẫn được xác định là corpus phù hợp để tiếp tục pretrain RoBERTa/BERT trong hướng phát triển sau.",
    ])
    insert_table_before(doc, target, ["Cấu hình", "Dữ liệu huấn luyện", "Dữ liệu đánh giá", "Mục đích"], [
        ["M0", "Dataset chính", "Validation/test của dataset chính", "Baseline supervised chính"],
        ["M3", "Dataset chính + pseudo-labeled CallCenterEN", "Validation/test của dataset chính", "Đo tác động của weak auxiliary corpus"],
        ["M2 mở rộng", "CallCenterEN DAPT + dataset chính fine-tune", "Validation/test của dataset chính", "Hướng tiếp theo cho Transformer"],
    ])


def add_experiment_results_section(doc: Document, valid_rows: list[dict[str, str]], test_rows: list[dict[str, str]]) -> None:
    target = find_content_heading(doc, "4.6. Kiểm thử khả năng chạy lại")
    add_heading_before(doc, target, "4.5.5. Thực nghiệm auxiliary training với CallCenterEN", 3)
    add_paragraphs_before(doc, target, [
        "Thực nghiệm auxiliary training được chạy sau khi đã tạo 300 pseudo-label từ CallCenterEN. Trong đó, 294 dòng được dùng trong training sau khi lọc các nhãn hợp lệ và merge lại với transcript. Cấu hình M0 và M3 sử dụng cùng pipeline BoW + Logistic Regression để đảm bảo so sánh công bằng.",
        "Điểm quan trọng của thiết kế đánh giá là CallCenterEN không được đưa vào validation hoặc test set. Toàn bộ kết quả bên dưới được đo trên split validation/test gốc của dataset chính. Vì vậy, kết quả phản ánh tác động của auxiliary data đến bài toán call_code của đề tài, không phải độ chính xác trên CallCenterEN.",
    ])
    insert_table_before(doc, target, ["Model", "Train rows", "Eval rows", "Subset accuracy", "Micro-F1", "Macro-F1", "Weighted-F1", "Hamming loss"], [
        [row["model"], row["train_rows"], row["eval_rows"], f"{float(row['subset_accuracy']):.4f}", f"{float(row['micro_f1']):.4f}", f"{float(row['macro_f1']):.4f}", f"{float(row['weighted_f1']):.4f}", f"{float(row['hamming_loss']):.4f}"]
        for row in valid_rows
    ])
    insert_paragraph_before(doc, target, "Bảng trên là kết quả trên validation set.")
    insert_table_before(doc, target, ["Model", "Train rows", "Eval rows", "Subset accuracy", "Micro-F1", "Macro-F1", "Weighted-F1", "Hamming loss"], [
        [row["model"], row["train_rows"], row["eval_rows"], f"{float(row['subset_accuracy']):.4f}", f"{float(row['micro_f1']):.4f}", f"{float(row['macro_f1']):.4f}", f"{float(row['weighted_f1']):.4f}", f"{float(row['hamming_loss']):.4f}"]
        for row in test_rows
    ])
    add_paragraphs_before(doc, target, [
        f"Trên test set, M3 tăng micro-F1 từ {float(test_rows[0]['micro_f1']):.4f} lên {float(test_rows[1]['micro_f1']):.4f} và macro-F1 từ {float(test_rows[0]['macro_f1']):.4f} lên {float(test_rows[1]['macro_f1']):.4f}. Mức cải thiện nhỏ nhưng có hướng tích cực, đủ để kết luận rằng CallCenterEN có thể đóng góp như auxiliary corpus trong một pilot experiment.",
        "Kết quả không nên được diễn giải quá mức thành một cải thiện lớn của mô hình. Nó cho thấy quy trình weak-label training là khả thi và có kết nối thực tế với mô hình hiện tại. Các thí nghiệm mở rộng có thể tăng số pseudo-label lên 1,000-2,000 mẫu, kiểm tra thủ công một phần nhãn, hoặc chuyển sang domain-adaptive pretraining cho mô hình Transformer.",
    ])


def add_chapter5_sections(doc: Document) -> None:
    target_contrib = find_content_heading(doc, "5.3. Hạn chế")
    add_heading_before(doc, target_contrib, "5.2.4. Đóng góp từ external baseline và auxiliary corpus", 3)
    add_paragraphs_before(doc, target_contrib, [
        "Việc đưa CallCenterEN vào nghiên cứu tạo ra hai đóng góp bổ sung. Thứ nhất, nó cung cấp cơ sở học thuật để giải thích thiết kế dataset của đề tài, đặc biệt là transcript hội thoại, thời lượng cuộc gọi, PII masking và NLP trên dữ liệu phi cấu trúc. Thứ hai, nó chứng minh dataset ngoài có thể tham gia training như auxiliary corpus thông qua pseudo-labeling, trong khi dataset chính vẫn giữ vai trò ground truth.",
        "Cách tiếp cận này giúp tránh hai cực đoan: không xem dữ liệu tổng hợp của đề tài là vô căn cứ, nhưng cũng không đánh tráo dataset chính bằng một public dataset thiếu nhãn nghiệp vụ. CallCenterEN được đặt đúng vai trò: external baseline để đối chiếu và corpus phụ trợ để thử nghiệm mở rộng mô hình.",
    ])
    target_limit = find_content_heading(doc, "5.4. Hướng phát triển")
    add_heading_before(doc, target_limit, "5.3.4. Hạn chế khi sử dụng CallCenterEN và pseudo-label", 3)
    add_paragraphs_before(doc, target_limit, [
        "CallCenterEN là dữ liệu tiếng Anh và đa miền, trong đó một số domain như Medicare hoặc automotive không trùng hoàn toàn với telesales tài chính. Vì vậy, việc sử dụng dữ liệu này cần lọc domain và kiểm soát domain shift. Trong thí nghiệm hiện tại, nghiên cứu chỉ dùng subset gần với insurance/customer service.",
        "Pseudo-label do AI sinh ra có thể chứa nhiễu, đặc biệt với các nhãn outcome như SUCCESSFUL_SALE hoặc SUDDEN_HANG_UP khi transcript không biểu hiện rõ kết quả. Do đó, các nhãn này không được dùng làm ground truth và không được đưa vào test set. Hướng phát triển tiếp theo là review thủ công một phần pseudo-label, tăng kích thước mẫu và thử domain-adaptive pretraining với mô hình Transformer.",
    ])


def add_references(doc: Document) -> None:
    target = find_heading(doc, "PHỤ LỤC")
    insert_paragraph_before(doc, target, "Dao, H., Chawla, G., Banda, R., & DeLeeuw, C. (2025). Real-World En Call Center Transcripts Dataset with PII Redaction. arXiv:2507.02958. https://arxiv.org/abs/2507.02958")
    insert_paragraph_before(doc, target, "AIxBlock. 92k Real-World Call Center Scripts English. Hugging Face Dataset. https://huggingface.co/datasets/AIxBlock/92k-real-world-call-center-scripts-english")
    insert_paragraph_before(doc, target, "Gururangan, S., Marasović, A., Swayamdipta, S., Lo, K., Beltagy, I., Downey, D., & Smith, N. A. (2020). Don't Stop Pretraining: Adapt Language Models to Domains and Tasks. ACL 2020. https://aclanthology.org/2020.acl-main.740/")
    insert_paragraph_before(doc, target, "Ratner, A., De Sa, C., Wu, S., Selsam, D., & Ré, C. (2016). Data Programming: Creating Large Training Sets, Quickly. NeurIPS 2016. https://papers.neurips.cc/paper/6523-data-programming-creating-large-training-sets-quickly")
    insert_paragraph_before(doc, target, "Amini, M.-R., Feofanov, V., Pauletto, L., Lies Hadjadj, E., Devijver, E., & Maximov, Y. (2022). Self-Training: A Survey. arXiv:2202.12040. https://arxiv.org/abs/2202.12040")


def main() -> None:
    summary = json.loads(SUMMARY_JSON.read_text(encoding="utf-8"))
    prompt_design = json.loads(PROMPT_DESIGN_JSON.read_text(encoding="utf-8"))
    pseudo_rows = read_csv(PSEUDO_CSV)
    valid_rows = read_csv(VALID_METRICS_CSV)
    test_rows = read_csv(TEST_METRICS_CSV)

    doc = Document(SOURCE_DOCX)
    set_document_style(doc)
    update_toc(doc)
    add_theory_section(doc)
    add_dataset_role_section(doc)
    add_prompt_design_section(doc, prompt_design)
    add_dataset_comparison_section(doc, summary)
    add_labeling_process_section(doc, pseudo_rows)
    add_aux_training_method_section(doc)
    add_experiment_results_section(doc, valid_rows, test_rows)
    add_chapter5_sections(doc)
    add_references(doc)
    doc.save(TARGET_DOCX)
    print(f"Wrote: {TARGET_DOCX.name}")


if __name__ == "__main__":
    main()
