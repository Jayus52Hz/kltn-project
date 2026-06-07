"""
Insert the CallCenterEN -> feature extraction -> prompt design section into the
already integrated report without duplicating previous CallCenterEN sections.
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import Any

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = Path(__file__).resolve().parent / "output"
SOURCE_DOCX = Path("Report KLTN - 22133056 - Nguyen Quoc Thinh - revised callcenteren integrated.docx")
TARGET_DOCX = Path("Report KLTN - 22133056 - Nguyen Quoc Thinh - revised callcenteren prompt integrated.docx")
PROMPT_DESIGN_JSON = OUTPUT_DIR / "callcenteren_prompt_design_analysis.json"


def content_start_index(doc: Document) -> int:
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == "PHẦN NỘI DUNG":
            return idx
    return 0


def find_content_heading(doc: Document, startswith: str):
    start = content_start_index(doc)
    for paragraph in doc.paragraphs[start:]:
        style_name = paragraph.style.name if paragraph.style else ""
        if "Heading" in style_name and paragraph.text.strip().startswith(startswith):
            return paragraph
    raise ValueError(f"Cannot find content heading: {startswith}")


def find_toc_line(doc: Document, startswith: str):
    end = content_start_index(doc)
    for paragraph in doc.paragraphs[:end]:
        if paragraph.text.strip().startswith(startswith):
            return paragraph
    return None


def insert_paragraph_before(doc: Document, target, text: str = "", style: str | None = None, bold: bool = False):
    paragraph = doc.add_paragraph(style=style)
    if text:
        run = paragraph.add_run(text)
        run.bold = bold
    element = paragraph._p
    parent = element.getparent()
    parent.remove(element)
    target._p.addprevious(element)
    return paragraph


def insert_table_before(doc: Document, target, headers: list[str], rows: list[list[Any]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        run = table.rows[0].cells[idx].paragraphs[0].add_run(str(header))
        run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    element = table._tbl
    parent = element.getparent()
    parent.remove(element)
    target._p.addprevious(element)
    insert_paragraph_before(doc, target, "")


def add_toc_line_if_missing(doc: Document) -> None:
    if any("2.1.5. Từ CallCenterEN đến prompt sinh dataset chính" in p.text for p in doc.paragraphs[:content_start_index(doc)]):
        return
    anchor = find_toc_line(doc, "2.1.4. Vai trò của CallCenterEN")
    if anchor is None:
        anchor = find_toc_line(doc, "2.1. Vai trò của dữ liệu")
    if anchor is not None:
        insert_paragraph_before(doc, anchor, "2.1.5. Từ CallCenterEN đến prompt sinh dataset chính")


def add_prompt_section(doc: Document, prompt_design: dict[str, Any]) -> None:
    if any("2.1.5. Từ CallCenterEN đến prompt sinh dataset chính" in p.text and "Heading" in p.style.name for p in doc.paragraphs):
        return

    target = find_content_heading(doc, "2.2. Nguồn gốc và quy trình sinh dữ liệu")
    insert_paragraph_before(doc, target, "2.1.5. Từ CallCenterEN đến prompt sinh dataset chính", style="Heading 3")
    for text in [
        "Mạch nghiên cứu của phần dữ liệu được tổ chức theo thứ tự: phân tích CallCenterEN, trích xuất đặc trưng, chuyển đặc trưng thành quy tắc prompt, sau đó dùng prompt để sinh dataset telesales chính. Cách trình bày này giúp giải thích vì sao dataset chính có các trường và nhãn hiện tại, thay vì chỉ mô tả dataset như một kết quả đã có sẵn.",
        f"Trong bước phân tích nền, nghiên cứu sử dụng {prompt_design['rows']:,} mẫu CallCenterEN từ baseline_analysis_sample. Tập này không dùng để train trực tiếp. Nó được dùng để quan sát cấu trúc hội thoại, độ dài transcript, duration, PII, hướng cuộc gọi, domain và các hành vi hội thoại thường gặp.",
        f"Kết quả thống kê cho thấy 3,000 mẫu có độ dài trung bình {prompt_design['avg_chars']:.4f} ký tự, {prompt_design['avg_words']:.4f} từ, thời lượng trung bình {prompt_design['avg_duration']:.4f} giây, ASR confidence trung bình {prompt_design['avg_confidence']:.4f} và trung bình {prompt_design['avg_pii_tokens']:.4f} token PII đã được redaction.",
    ]:
        insert_paragraph_before(doc, target, text)

    insert_table_before(doc, target, ["Nhóm đặc trưng từ CallCenterEN", "Số mẫu", "Tỷ lệ"], [
        [group, f"{value['rows']:,}", f"{value['percent']}%"]
        for group, value in prompt_design["keyword_coverage"].items()
    ])
    insert_table_before(doc, target, ["Quan sát từ CallCenterEN", "Quy tắc prompt sinh dataset chính", "Trường/nhãn được tạo"], [
        [
            item["callcenteren_observation"],
            item["prompt_rule_for_primary_dataset"],
            item["primary_dataset_field_or_label"],
        ]
        for item in prompt_design["prompt_design_mapping"]
    ])
    for text in [
        "Từ bảng trên, prompt sinh dataset chính được thiết kế để bắt buộc transcript có cấu trúc hội thoại rõ ràng: mở đầu cuộc gọi, xác minh hoặc nhắc đến thông tin khách hàng, phân tích nhu cầu, giới thiệu sản phẩm/offer, thảo luận phí hoặc điều kiện, xử lý phản đối và kết thúc bằng outcome.",
        "Như vậy, CallCenterEN đóng vai trò cơ sở để xây dựng logic prompt và schema dataset chính. Sau bước này, dataset chính mới được sinh theo bối cảnh telesales tài chính của đồ án, có thêm customer profile, offer, campaign, call metadata và nhãn call_code để phục vụ cả mô hình NLP lẫn star schema phân tích.",
        "Cần phân biệt rõ hai con số trong nghiên cứu: 3,000 mẫu CallCenterEN được dùng cho phân tích đặc trưng và thiết kế prompt sinh dataset chính; 300 mẫu pseudo-label là thí nghiệm nhỏ hơn để kiểm tra khả năng dùng CallCenterEN như auxiliary training data.",
    ]:
        insert_paragraph_before(doc, target, text)


def main() -> None:
    prompt_design = json.loads(PROMPT_DESIGN_JSON.read_text(encoding="utf-8"))
    doc = Document(str(SOURCE_DOCX))
    add_toc_line_if_missing(doc)
    add_prompt_section(doc, prompt_design)
    doc.save(str(TARGET_DOCX))
    print(f"Wrote: {TARGET_DOCX}")


if __name__ == "__main__":
    main()
