"""
Append CallCenterEN baseline and auxiliary-training sections to the thesis DOCX.
"""

from __future__ import annotations

import csv
import json
from collections import Counter
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = Path(__file__).resolve().parent / "output"
SOURCE_DOCX = ROOT / "Report KLTN - 22133056 - Nguyen Quoc Thinh.docx"
TARGET_DOCX = ROOT / "Report KLTN - 22133056 - Nguyen Quoc Thinh - revised callcenteren baseline.docx"

SUMMARY_JSON = OUTPUT_DIR / "dataset_comparison_summary.json"
PSEUDO_CSV = OUTPUT_DIR / "pseudo_labels_gemini.csv"
VALID_METRICS_CSV = OUTPUT_DIR / "auxiliary_bow_valid_metrics.csv"
TEST_METRICS_CSV = OUTPUT_DIR / "auxiliary_bow_test_metrics.csv"


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8", newline="") as file:
        return list(csv.DictReader(file))


def add_page_break(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def set_normal_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)


def add_table(doc: Document, headers: list[str], rows: list[list[Any]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        run = header_cells[idx].paragraphs[0].add_run(str(header))
        run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    doc.add_paragraph()


def metric(value: Any) -> str:
    if isinstance(value, float):
        return f"{value:.4f}"
    return str(value)


def add_dataset_section(doc: Document, summary: dict[str, Any]) -> None:
    primary = summary["primary_dataset"]
    external = summary["callcenteren_baseline_subset"]

    doc.add_heading("Bo sung: CallCenterEN lam external baseline va auxiliary corpus", level=1)
    doc.add_paragraph(
        "Tap du lieu chinh cua do an van la tap telesales do tac gia xay dung, "
        "vi tap nay co day du customer, offer, call_logs, transcript va nhan nghiep vu call_code. "
        "CallCenterEN khong duoc dung de thay the tap du lieu chinh. Trong nghien cuu nay, "
        "CallCenterEN duoc su dung voi hai vai tro: external real-world baseline de bien minh "
        "thiet ke du lieu transcript, va auxiliary corpus de thu nghiem pseudo-label training."
    )
    doc.add_paragraph(
        "CallCenterEN la tap transcript call center tieng Anh duoc cong bo cong khai, "
        "bao gom 91,706 cuoc hoi thoai va 10,448 gio audio theo bai bao Dao et al. (2025). "
        "Tap du lieu nay co transcript, audio duration, ASR confidence, word-level timestamps "
        "va PII redaction. Cac dac trung nay cung cap co so hoc thuat de giai thich vi sao "
        "dataset cua do an can co call_transcript, talk_time_seconds va buoc masking PII trong Silver layer."
    )

    add_table(
        doc,
        ["Tieu chi", "Dataset chinh cua do an", "CallCenterEN subset", "Y nghia"],
        [
            ["Vai tro", "Primary task dataset", "External baseline / auxiliary corpus", "CallCenterEN ho tro doi chieu va training phu, khong thay the dataset chinh"],
            ["So mau", f"{primary['rows']:,}", f"{external['rows']:,}", "Dataset chinh lon hon subset tham chieu da loc"],
            ["Avg transcript chars", metric(primary["avg_chars"]), metric(external["avg_chars"]), "CallCenterEN dai hon, gan voi hoi thoai call center thuc te"],
            ["Avg word count", metric(primary["avg_words"]), metric(external["avg_words"]), "Hai tap cung co du lieu hoi thoai phi cau truc"],
            ["Avg duration seconds", metric(primary["avg_duration_seconds"]), metric(external["avg_duration_seconds"]), "talk_time_seconds co co so tu audio_duration trong call center"],
            ["Nhan nghiep vu", "call_code", "Khong co", "CallCenterEN khong the dung lam ground truth chinh"],
            ["ASR confidence", "Khong co", metric(external["avg_asr_confidence"]), "Co the mo rong de danh gia chat luong transcript"],
            ["Avg PII tokens", "PII field masking", metric(external["avg_pii_tokens"]), "Cung co co so cho privacy-preserving pipeline"],
        ],
    )

    doc.add_paragraph(
        "Ket qua so sanh cho thay dataset cua do an va CallCenterEN co cung lop du lieu cot loi: "
        "noi dung hoi thoai va thoi luong cuoc goi. Diem khac biet quan trong la CallCenterEN "
        "khong co business entities va call_code, trong khi day la hai thanh phan can thiet de "
        "xay dung star schema, dashboard va mo hinh phan loai nghiep vu trong do an."
    )


def add_pseudo_label_section(doc: Document, pseudo_rows: list[dict[str, str]]) -> None:
    label_counter: Counter[str] = Counter()
    usable = 0
    for row in pseudo_rows:
        if row.get("should_use_for_training", "").lower() == "true":
            usable += 1
        for label in [item.strip() for item in row.get("pseudo_call_code", "").split(",") if item.strip()]:
            label_counter[label] += 1

    doc.add_heading("Thi nghiem pseudo-label voi CallCenterEN", level=1)
    doc.add_paragraph(
        "Sau khi loc CallCenterEN subset, nghien cuu su dung Google AI Studio model gemma-4-31b-it "
        "de gan pseudo-label theo taxonomy call_code cua do an. Gioi han toc do duoc giu o muc 15 RPM; "
        "script chay voi sleep 4.2 giay giua cac request. Cac nhan nay duoc xem la weak labels, "
        "khong phai ground truth."
    )
    add_table(
        doc,
        ["Chi so", "Gia tri"],
        [
            ["Pseudo-label rows generated", f"{len(pseudo_rows):,}"],
            ["Rows marked usable for training", f"{usable:,}"],
            ["Model labeling", "gemma-4-31b-it"],
            ["API source", "Google AI Studio generateContent API"],
            ["RPM constraint", "15 requests/minute"],
        ],
    )
    add_table(
        doc,
        ["Pseudo-label", "So lan xuat hien"],
        [[label, f"{count:,}"] for label, count in label_counter.most_common(15)],
    )
    doc.add_paragraph(
        "Ve mat phuong phap, pseudo-label workflow dua tren weak supervision va self-training. "
        "Cac nhan do AI sinh ra co the chua noise, vi vay chi duoc dung lam du lieu training phu "
        "va khong duoc dua vao tap test chinh."
    )


def add_experiment_section(doc: Document, valid_rows: list[dict[str, str]], test_rows: list[dict[str, str]]) -> None:
    doc.add_heading("Danh gia tac dong cua auxiliary pseudo-label", level=1)
    doc.add_paragraph(
        "De kiem tra CallCenterEN co dong gop thuc te cho mo hinh hay khong, nghien cuu so sanh "
        "hai cau hinh BoW + Logistic Regression: M0 chi train tren dataset chinh va M3 train tren "
        "dataset chinh ket hop pseudo-labeled CallCenterEN. Tap validation va test van giu nguyen "
        "tu dataset chinh cua do an."
    )
    add_table(
        doc,
        ["Model", "Train rows", "Eval rows", "Subset accuracy", "Micro-F1", "Macro-F1", "Weighted-F1", "Hamming loss"],
        [
            [
                row["model"],
                row["train_rows"],
                row["eval_rows"],
                f"{float(row['subset_accuracy']):.4f}",
                f"{float(row['micro_f1']):.4f}",
                f"{float(row['macro_f1']):.4f}",
                f"{float(row['weighted_f1']):.4f}",
                f"{float(row['hamming_loss']):.4f}",
            ]
            for row in valid_rows
        ],
    )
    doc.add_paragraph("Bang tren la ket qua tren validation set.")
    add_table(
        doc,
        ["Model", "Train rows", "Eval rows", "Subset accuracy", "Micro-F1", "Macro-F1", "Weighted-F1", "Hamming loss"],
        [
            [
                row["model"],
                row["train_rows"],
                row["eval_rows"],
                f"{float(row['subset_accuracy']):.4f}",
                f"{float(row['micro_f1']):.4f}",
                f"{float(row['macro_f1']):.4f}",
                f"{float(row['weighted_f1']):.4f}",
                f"{float(row['hamming_loss']):.4f}",
            ]
            for row in test_rows
        ],
    )
    doc.add_paragraph(
        "Ket qua test cho thay auxiliary pseudo-label mang lai cai thien nhe: micro-F1 tang tu "
        f"{float(test_rows[0]['micro_f1']):.4f} len {float(test_rows[1]['micro_f1']):.4f}, "
        f"macro-F1 tang tu {float(test_rows[0]['macro_f1']):.4f} len {float(test_rows[1]['macro_f1']):.4f}. "
        "Do muc tang con nho, ket qua nen duoc trinh bay nhu pilot experiment, khong nen ket luan "
        "qua muc. Tuy nhien, no chung minh CallCenterEN co the tham gia vao model nhu auxiliary corpus "
        "va tao tin hieu huan luyen co lien quan den ngon ngu call center thuc te."
    )


def add_references(doc: Document) -> None:
    doc.add_heading("Tai lieu tham khao bo sung", level=1)
    refs = [
        "Dao, H., Chawla, G., Banda, R., & DeLeeuw, C. (2025). Real-World En Call Center Transcripts Dataset with PII Redaction. arXiv:2507.02958. https://arxiv.org/abs/2507.02958",
        "AIxBlock. 92k Real-World Call Center Scripts English. Hugging Face Dataset. https://huggingface.co/datasets/AIxBlock/92k-real-world-call-center-scripts-english",
        "Gururangan, S., Marasovic, A., Swayamdipta, S., Lo, K., Beltagy, I., Downey, D., & Smith, N. A. (2020). Don't Stop Pretraining: Adapt Language Models to Domains and Tasks. ACL 2020. https://aclanthology.org/2020.acl-main.740/",
        "Ratner, A., De Sa, C., Wu, S., Selsam, D., & Re, C. (2016). Data Programming: Creating Large Training Sets, Quickly. NeurIPS 2016. https://papers.neurips.cc/paper/6523-data-programming-creating-large-training-sets-quickly",
        "Amini, M.-R., Feofanov, V., Pauletto, L., Lies Hadjadj, E., Devijver, E., & Maximov, Y. (2022). Self-Training: A Survey. arXiv:2202.12040. https://arxiv.org/abs/2202.12040",
    ]
    for ref in refs:
        doc.add_paragraph(ref, style=None)


def main() -> None:
    summary = json.loads(SUMMARY_JSON.read_text(encoding="utf-8"))
    pseudo_rows = read_csv(PSEUDO_CSV)
    valid_rows = read_csv(VALID_METRICS_CSV)
    test_rows = read_csv(TEST_METRICS_CSV)

    doc = Document(SOURCE_DOCX)
    set_normal_style(doc)
    add_page_break(doc)
    add_dataset_section(doc, summary)
    add_pseudo_label_section(doc, pseudo_rows)
    add_experiment_section(doc, valid_rows, test_rows)
    add_references(doc)

    section = doc.sections[-1]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    doc.save(TARGET_DOCX)
    print(f"Wrote: {TARGET_DOCX.name}")


if __name__ == "__main__":
    main()
