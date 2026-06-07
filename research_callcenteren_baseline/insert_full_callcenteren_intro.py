"""
Add a full Vietnamese introduction of CallCenterEN into the integrated report.
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import Any

from docx import Document


OUTPUT_DIR = Path("research_callcenteren_baseline") / "output"
SOURCE_DOCX = Path("Report KLTN - 22133056 - Nguyen Quoc Thinh - revised callcenteren prompt integrated.docx")
TARGET_DOCX = Path("Report KLTN - 22133056 - Nguyen Quoc Thinh - revised callcenteren full integrated.docx")
SUBSET_SUMMARY = Path("92k-real-world-call-center-scripts-english") / "prepared_subset" / "subset_summary.json"
PROMPT_DESIGN_JSON = OUTPUT_DIR / "callcenteren_prompt_design_analysis.json"


def content_start_index(doc: Document) -> int:
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == "PHẦN NỘI DUNG":
            return idx
    return 0


def find_content_heading(doc: Document, startswith: str):
    start = content_start_index(doc)
    for paragraph in doc.paragraphs[start:]:
        style_name = paragraph.style.name if paragraph.style else ""
        if "Heading" in style_name and paragraph.text.strip().startswith(startswith):
            return paragraph
    raise ValueError(f"Cannot find content heading: {startswith}")


def insert_paragraph_before(doc: Document, target, text: str = "", style: str | None = None):
    paragraph = doc.add_paragraph(style=style)
    if text:
        paragraph.add_run(text)
    element = paragraph._p
    parent = element.getparent()
    parent.remove(element)
    target._p.addprevious(element)
    return paragraph


def insert_table_before(doc: Document, target, headers: list[str], rows: list[list[Any]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        run = table.rows[0].cells[idx].paragraphs[0].add_run(str(header))
        run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    element = table._tbl
    parent = element.getparent()
    parent.remove(element)
    target._p.addprevious(element)
    insert_paragraph_before(doc, target, "")


def already_inserted(doc: Document) -> bool:
    return any("2.1.4.1. Nguồn gốc, cách công bố và phạm vi sử dụng CallCenterEN" in p.text for p in doc.paragraphs)


def add_full_intro(doc: Document, subset_summary: dict[str, Any], prompt_design: dict[str, Any]) -> None:
    if already_inserted(doc):
        return

    target = find_content_heading(doc, "2.1.5. Từ CallCenterEN đến prompt sinh dataset chính")

    insert_paragraph_before(doc, target, "2.1.4.1. Nguồn gốc, cách công bố và phạm vi sử dụng CallCenterEN", style="Heading 4")
    for text in [
        "CallCenterEN là tập dữ liệu transcript call center tiếng Anh được giới thiệu trong bài báo Real-World En Call Center Transcripts Dataset with PII Redaction của Dao, Chawla, Banda và DeLeeuw, công bố trên arXiv năm 2025. Dataset card được phát hành công khai trên Hugging Face bởi AIxBlock với tên 92k-real-world-call-center-scripts-english. Trong nghiên cứu này, tác giả không thu thập trực tiếp các cuộc gọi này, mà sử dụng bản công bố công khai của CallCenterEN như một nguồn dữ liệu tham chiếu học thuật và corpus phụ trợ.",
        "Theo bài báo, CallCenterEN được xây dựng từ các quan hệ hợp tác với khoảng 9-10 BPO centers. Các bản ghi âm được thu thập từ hoạt động call center thực tế, sau đó được xử lý bằng dịch vụ ASR thương mại của AssemblyAI để tạo transcript. Tập dữ liệu bao gồm cả inbound và outbound calls, với tiếng Anh có các accent như Indian, Filipino và American. Do rủi ro nhận dạng sinh trắc học, audio gốc không được phát hành công khai; bản public release chỉ cung cấp transcript đã được loại bỏ PII.",
        "Tổng quy mô được paper công bố là 91,706 conversations, tương ứng khoảng 10,448 giờ audio trước khi audio bị loại khỏi bản public release. Mỗi transcript được lưu ở dạng JSON, có trường text chứa hội thoại đã redaction, confidence phản ánh độ tin cậy ASR, audio_duration tính bằng giây, danh sách words có timestamp/confidence theo từng từ, và danh sách redacted_pii_policies. Dataset được phát hành theo giấy phép CC BY-NC 4.0, phù hợp cho nghiên cứu học thuật và phát triển mô hình phi thương mại.",
        "Các hạng mục PII trong CallCenterEN rất rộng, bao gồm định danh cá nhân, số điện thoại, email, ngày sinh, địa chỉ, thông tin tài chính, thông tin y tế, số giấy tờ, vị trí, tổ chức, nghề nghiệp, thời gian và các chuỗi số. Đặc điểm này đặc biệt quan trọng với đề tài vì pipeline lakehouse của hệ thống telesales cũng cần xử lý transcript và dữ liệu khách hàng mà không để PII đi thẳng vào tầng phân tích.",
    ]:
        insert_paragraph_before(doc, target, text)

    insert_table_before(doc, target, ["Thuộc tính", "Mô tả CallCenterEN", "Ý nghĩa trong đề tài"], [
        ["Nguồn công bố", "Paper arXiv 2507.02958 và Hugging Face dataset card của AIxBlock", "Có nguồn học thuật và dataset card để trích dẫn"],
        ["Nguồn thu thập gốc", "Đối tác BPO centers, call center thực tế", "Làm baseline thực tế cho dữ liệu transcript"],
        ["Quy mô paper công bố", "91,706 conversations, khoảng 10,448 giờ audio", "Cho thấy transcript call center là nguồn dữ liệu quy mô lớn"],
        ["Bản public release", "Chỉ phát hành transcript, không phát hành audio", "Phù hợp phạm vi đồ án xử lý text transcript"],
        ["Call types", "Inbound và outbound", "Giải thích bối cảnh cuộc gọi đa chiều trong telesales/call center"],
        ["Ngôn ngữ/accent", "English; Indian, Filipino, American accents", "Có tính thực tế nhưng cần lưu ý domain/language shift"],
        ["Format", "JSON: text, confidence, audio_duration, words, redacted_pii_policies", "Đối chiếu với call_transcript, talk_time_seconds và xử lý PII"],
        ["PII", "Đã redaction nhiều nhóm PII", "Củng cố thiết kế masking ở Silver layer"],
        ["License", "CC BY-NC 4.0", "Chỉ dùng cho nghiên cứu/phi thương mại, không dùng làm dữ liệu thương mại"],
    ])

    insert_paragraph_before(doc, target, "2.1.4.2. Cách đưa CallCenterEN vào workspace và phạm vi dữ liệu được chọn", style="Heading 4")
    for text in [
        "Trong workspace của đề tài, CallCenterEN được đưa vào thư mục 92k-real-world-call-center-scripts-english dưới dạng các file ZIP theo domain/call type. Thay vì giải nén toàn bộ hoặc đưa nguyên dataset vào pipeline chính, nghiên cứu đọc trực tiếp JSON từ ZIP, loại bỏ thư mục __MACOSX, sau đó lọc subset theo tiêu chí chất lượng và độ gần domain.",
        "Ở bước đầu, nghiên cứu ưu tiên các nhóm gần với telesales và customer service: insurance_outbound, auto_insurance_customer_service_inbound, PII_redacted_auto_insurance_script và customer_service_general_inbound. Các nhóm như medicare_inbound, automotive_inbound hoặc home_service_inbound có thể hữu ích cho mở rộng sau, nhưng không được ưu tiên trong pilot vì có nguy cơ lệch domain so với bối cảnh telesales tài chính của đề tài.",
        "Kết quả lọc đã tạo hai tập con khác nhau. Tập baseline_analysis_sample gồm 3,000 mẫu dùng cho phân tích đặc trưng và thiết kế prompt sinh dataset chính. Tập auxiliary_training_candidate gồm 2,000 mẫu sạch hơn, dùng để gán pseudo-label và thử nghiệm huấn luyện phụ. Vì vậy, con số 3,000 và 300 không mâu thuẫn: 3,000 là mẫu phân tích thiết kế dataset/prompt; 300 là pilot pseudo-label training.",
    ]:
        insert_paragraph_before(doc, target, text)

    selected = subset_summary["baseline_output"]
    insert_table_before(doc, target, ["Tập con", "Số mẫu", "Vai trò"], [
        ["baseline_analysis_sample", "3,000", "Phân tích đặc trưng CallCenterEN và thiết kế prompt sinh dataset chính"],
        ["auxiliary_training_candidate", "2,000", "Ứng viên để gán pseudo-label và huấn luyện phụ"],
        ["pseudo_labels_gemini", "300", "Pilot weak-label training; không phải toàn bộ phần nghiên cứu CallCenterEN"],
    ])
    insert_table_before(doc, target, ["Đặc trưng từ 3,000 mẫu baseline", "Giá trị"], [
        ["Domain chính", ", ".join(f"{k}: {v}" for k, v in selected["by_domain"].items())],
        ["Hướng cuộc gọi", ", ".join(f"{k}: {v}" for k, v in selected["by_direction"].items())],
        ["Độ dài trung bình", f"{prompt_design['avg_chars']:.4f} ký tự; {prompt_design['avg_words']:.4f} từ"],
        ["Thời lượng trung bình", f"{prompt_design['avg_duration']:.4f} giây"],
        ["ASR confidence trung bình", f"{prompt_design['avg_confidence']:.4f}"],
        ["PII token trung bình", f"{prompt_design['avg_pii_tokens']:.4f}"],
    ])

    insert_paragraph_before(doc, target, "2.1.4.3. Các đặc trưng được rút ra từ CallCenterEN", style="Heading 4")
    insert_paragraph_before(
        doc,
        target,
        "Từ 3,000 mẫu baseline, nghiên cứu rút ra các đặc trưng hội thoại phổ biến: cuộc gọi gần như luôn có opening; phần lớn có product/offer discussion, identity verification, fee discussion và needs analysis; một tỷ lệ nhỏ hơn có objection/rejection hoặc follow-up/handoff. Những đặc trưng này được dùng để thiết kế prompt sinh dataset chính sao cho transcript không chỉ là đoạn văn ngẫu nhiên, mà có cấu trúc nghiệp vụ tương tự hội thoại call center.",
    )
    insert_table_before(doc, target, ["Nhóm đặc trưng", "Số mẫu", "Tỷ lệ"], [
        [group, f"{value['rows']:,}", f"{value['percent']}%"]
        for group, value in prompt_design["keyword_coverage"].items()
    ])

    insert_paragraph_before(
        doc,
        target,
        "Tóm lại, CallCenterEN được giới thiệu và sử dụng theo ba tầng: tầng nghiên cứu nguồn dữ liệu thực tế, tầng phân tích đặc trưng để thiết kế prompt/dataset chính, và tầng thử nghiệm auxiliary training bằng pseudo-label. Cách tổ chức này giúp dataset chính có cơ sở giải thích rõ ràng, đồng thời vẫn giữ ranh giới học thuật: dataset chính là nguồn ground truth, CallCenterEN là external baseline và auxiliary corpus.",
    )


def main() -> None:
    subset_summary = json.loads(SUBSET_SUMMARY.read_text(encoding="utf-8"))
    prompt_design = json.loads(PROMPT_DESIGN_JSON.read_text(encoding="utf-8"))
    doc = Document(str(SOURCE_DOCX))
    add_full_intro(doc, subset_summary, prompt_design)
    doc.save(str(TARGET_DOCX))
    print(f"Wrote: {TARGET_DOCX}")


if __name__ == "__main__":
    main()
