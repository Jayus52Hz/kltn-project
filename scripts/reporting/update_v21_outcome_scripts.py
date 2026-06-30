from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


REPORTS_DIR = Path("docs/reports")
INPUT = REPORTS_DIR / "Report KLTN - 22133056 - Nguyen Quoc Thinh - 30-06 - v20 conclusion compact.docx"
OUTPUT = REPORTS_DIR / "Report KLTN - 22133056 - Nguyen Quoc Thinh - 30-06 - v21 outcome scripts.docx"


def norm(text):
    return " ".join((text or "").split())


def find_paragraph(doc, needle, style_prefix=None):
    for paragraph in doc.paragraphs:
        text = norm(paragraph.text)
        if needle in text:
            if style_prefix is None or paragraph.style.name.startswith(style_prefix):
                return paragraph
    raise ValueError(f"Paragraph not found: {needle}")


def find_paragraph_index(doc, needle, style_prefix=None):
    for idx, paragraph in enumerate(doc.paragraphs):
        text = norm(paragraph.text)
        if needle in text:
            if style_prefix is None or paragraph.style.name.startswith(style_prefix):
                return idx
    raise ValueError(f"Paragraph index not found: {needle}")


def paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if style:
        new_paragraph.style = style
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def add_table_after(doc, paragraph, rows, style_source_index=None):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    if style_source_index is not None:
        table.style = doc.tables[style_source_index].style
    else:
        table.style = "Table Grid"
    for row_idx, values in enumerate(rows):
        for col_idx, value in enumerate(values):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = value
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    if row_idx == 0:
                        run.bold = True
    paragraph._p.addnext(table._tbl)
    return table


def append_row(table, values):
    row = table.add_row()
    for idx, value in enumerate(values):
        row.cells[idx].text = value
    return row


def main():
    doc = Document(INPUT)

    # Front table list: add the new verification table.
    table_list_anchor = find_paragraph(
        doc,
        "Bảng 4.8. Kết quả kiểm thử full pipeline tuần tự ngày 26/06/2026",
    )
    paragraph_after(
        table_list_anchor,
        "Bảng 4.9. Kết quả kiểm thử sinh kịch bản theo outcome_strategy ngày 30/06/2026",
        "Normal",
    )

    # Chapter 3: design section after the current dashboard subsection.
    ch3_anchor = find_paragraph(doc, "Superset và dashboard", "Heading")
    # Move to the last paragraph before 3.6 by walking from the heading text.
    ch3_idx = find_paragraph_index(doc, "Superset và dashboard", "Heading")
    ch3_insert_anchor = doc.paragraphs[ch3_idx]
    for paragraph in doc.paragraphs[ch3_idx + 1 :]:
        if norm(paragraph.text).startswith("3.6. Thiết kế bảo mật"):
            break
        ch3_insert_anchor = paragraph

    p = paragraph_after(
        ch3_insert_anchor,
        "3.5.4. Sinh kịch bản telesales theo outcome_strategy",
        "Heading 3",
    )
    p = paragraph_after(
        p,
        (
            "Sau lớp Gold và BigQuery serving, đồ án bổ sung một artifact phục vụ hành động nghiệp vụ: "
            "bảng kịch bản telesales theo outcome. Job outcome_script_job.py đọc fact_telesales_calls, "
            "dim_customer và dim_offer để tạo bảng lakehouse.gold.customer_outcome_scripts. Khi upstream "
            "chưa có outcome_strategy riêng, job ánh xạ outcome_category sang chiến lược hành động mặc định, "
            "ví dụ SALE -> CONFIRM_AND_COMPLETE, CALLBACK -> SCHEDULE_CALLBACK, DO_NOT_CALL -> SUPPRESS_CONTACT. "
            "Thiết kế này giúp chuyển kết quả phân tích cuộc gọi thành bước gợi ý vận hành có thể sử dụng bởi "
            "agent app hoặc dashboard, thay vì chỉ dừng ở thống kê outcome."
        ),
        "Normal",
    )
    p = paragraph_after(
        p,
        (
            "Cơ chế sinh nội dung dùng template rule cố định, không gọi LLM trong pipeline. Module "
            "outcome_script_rules.py định nghĩa sáu nhóm outcome gồm SALE, CALLBACK, SOFT_REJECTION, "
            "HARD_REJECTION, DO_NOT_CALL và IN_PROGRESS, kèm next_action tương ứng. Bảng đầu ra lưu các "
            "trường script_title, opening_line, main_pitch, objection_response, next_action, closing_line "
            "và variables_json. Các biến render lấy từ hồ sơ khách hàng và offer như product_name, loan_amount, "
            "interest_rate, income_band, credit_tier, is_existing_customer và lead_source."
        ),
        "Normal",
    )
    p = paragraph_after(
        p,
        (
            "Về bảo mật, artifact này không đưa phone_number, national_id hoặc call_transcript thô vào script. "
            "Tuy nhiên, do kịch bản có thể cá nhân hóa bằng tên khách hàng, bảng customer_outcome_scripts cần "
            "được xem như lớp serving vận hành có quyền truy cập hạn chế, khác với view phân tích tổng hợp chỉ "
            "phục vụ BI. Đây là một đánh đổi thiết kế giữa khả năng cá nhân hóa hội thoại và yêu cầu kiểm soát "
            "PII ở lớp publish dữ liệu."
        ),
        "Normal",
    )

    # Chapter 4: runtime verification after table 4.8.
    find_paragraph(doc, "PHẦN KẾT LUẬN", "Heading")
    conclusion_idx = find_paragraph_index(doc, "PHẦN KẾT LUẬN", "Heading")
    ch4_anchor = doc.paragraphs[conclusion_idx - 1]
    p = paragraph_after(
        ch4_anchor,
        "4.6.5. Kiểm thử sinh kịch bản theo outcome_strategy ngày 30/06/2026",
        "Heading 3",
    )
    p = paragraph_after(
        p,
        (
            "Sau khi triển khai job sinh kịch bản, Docker stack được khởi động lại bằng docker compose. "
            "Lần khởi động đầu ghi nhận Kafka lỗi do ZooKeeper còn ephemeral node /brokers/ids/1 từ phiên trước; "
            "sau khi restart riêng Kafka và Debezium Connect, các container chính gồm Kafka, Debezium, MinIO, "
            "MongoDB, Spark, Airflow và Superset đều ở trạng thái running hoặc healthy. Điều kiện này đủ để chạy "
            "job hậu Gold vì outcome_script_job.py chỉ đọc bảng Iceberg Gold trên MinIO và ghi lại vào namespace "
            "lakehouse.gold."
        ),
        "Normal",
    )
    p = paragraph_after(
        p,
        (
            "Lệnh spark-submit trong container spark-master đã chạy thành công outcome_script_job.py. Log Spark "
            "ghi nhận MERGE INTO lakehouse.gold.customer_outcome_scripts với 23.447 source records và thông báo "
            "Outcome script job completed successfully. Truy vấn kiểm chứng trên Spark SQL cho thấy bảng mới có "
            "đủ 23.447 dòng, bằng số dòng fact_telesales_calls, đồng thời phân bố theo outcome gồm CALLBACK 5.248, "
            "DO_NOT_CALL 4.503, HARD_REJECTION 2.268, IN_PROGRESS 5.137, SALE 3.933 và SOFT_REJECTION 2.358."
        ),
        "Normal",
    )
    p = paragraph_after(
        p,
        (
            "Bước đồng bộ BigQuery sau đó cũng hoàn tất. bq_sync_job.py đã ghi bảng "
            "project-ef0c6db5-0765-4391-845.kltn0710.customer_outcome_scripts với 23.447 dòng, sau đó file "
            "create_serving_views.sql tạo thành công view vw_customer_outcome_scripts. Truy vấn COUNT trên view "
            "trả về 23.447 dòng và phân bố next_action khớp với bảng Lakehouse. Riêng nhóm DO_NOT_CALL được kiểm "
            "tra mẫu cho thấy main_pitch chỉ ghi nhận yêu cầu ngừng liên hệ, next_action là ADD_TO_DO_NOT_CALL_LIST, "
            "không tiếp tục pitch sản phẩm."
        ),
        "Normal",
    )
    caption = paragraph_after(
        p,
        "Bảng 4.9. Kết quả kiểm thử sinh kịch bản theo outcome_strategy ngày 30/06/2026",
        "Report Caption",
    )
    add_table_after(
        doc,
        caption,
        [
            ["Nhóm kiểm thử", "Kết quả xác nhận", "Số dòng / trạng thái"],
            [
                "Spark Gold job",
                "outcome_script_job.py merge bảng customer_outcome_scripts sau fact_telesales_calls",
                "23.447 dòng",
            ],
            [
                "Outcome mapping",
                "Sáu nhóm outcome được ánh xạ sang next_action bằng template rule deterministic",
                "CALLBACK 5.248; DO_NOT_CALL 4.503; HARD_REJECTION 2.268; IN_PROGRESS 5.137; SALE 3.933; SOFT_REJECTION 2.358",
            ],
            [
                "DO_NOT_CALL guardrail",
                "Mẫu script không pitch tiếp sản phẩm và dùng next_action ADD_TO_DO_NOT_CALL_LIST",
                "Đạt",
            ],
            [
                "BigQuery sync",
                "bq_sync_job.py đồng bộ customer_outcome_scripts sang dataset kltn0710",
                "23.447 dòng",
            ],
            [
                "Serving view",
                "create_serving_views.sql tạo vw_customer_outcome_scripts và truy vấn COUNT thành công",
                "23.447 dòng",
            ],
            [
                "Airflow integration",
                "DAG import không có lỗi; task customer_outcome_scripts được đặt sau fact_telesales_calls và trước bq_sync_gold",
                "No import errors",
            ],
        ],
        style_source_index=72,
    )

    # Compact conclusion updates.
    total_anchor = find_paragraph(doc, "Đồng bộ dữ liệu Gold sang BigQuery")
    paragraph_after(
        total_anchor,
        (
            "• Bổ sung bước hậu Gold sinh kịch bản telesales theo outcome_strategy, tạo bảng "
            "customer_outcome_scripts và view vw_customer_outcome_scripts với 23.447 dòng đã kiểm chứng."
        ),
        "Normal",
    )
    contribution_anchor = find_paragraph(doc, "Đóng góp về vận hành")
    paragraph_after(
        contribution_anchor,
        (
            "• Đóng góp về lớp hành động nghiệp vụ: chuyển outcome phân tích thành kịch bản hội thoại "
            "có next_action rõ ràng, hỗ trợ nối kết giữa Lakehouse analytics và hoạt động telesales."
        ),
        "Normal",
    )
    limitation_anchor = find_paragraph(doc, "Cơ chế bảo mật đã có PII masking")
    paragraph_after(
        limitation_anchor,
        (
            "• Kịch bản hiện được sinh bằng template rule cố định, chưa có đánh giá chất lượng hội thoại "
            "từ người dùng nghiệp vụ; bảng script có nội dung cá nhân hóa nên cần bổ sung phân quyền và "
            "chính sách truy cập chặt hơn nếu đưa vào production."
        ),
        "Normal",
    )
    future_anchor = find_paragraph(doc, "Thử nghiệm triển khai trên hạ tầng cloud")
    paragraph_after(
        future_anchor,
        (
            "• Mở rộng lớp sinh kịch bản theo hướng quản trị template, A/B testing, kiểm duyệt nội dung "
            "và có thể kết hợp LLM có guardrail khi hạ tầng và yêu cầu kiểm soát rủi ro cho phép."
        ),
        "Normal",
    )

    # Appendix D code/report mapping table.
    appendix_table = doc.tables[-1]
    append_row(
        appendix_table,
        [
            "project/batch-etl/outcome_script_job.py",
            "Sinh bảng Gold customer_outcome_scripts từ fact_telesales_calls, dim_customer và dim_offer bằng template deterministic.",
            "Chương 3.5.4, Chương 4.6.5",
            "Spark log merge 23.447 dòng; schema gồm script_id, outcome_strategy, script_title, opening_line, main_pitch, next_action, closing_line và variables_json.",
        ],
    )
    append_row(
        appendix_table,
        [
            "project/batch-etl/outcome_script_rules.py",
            "Tập trung rule ánh xạ outcome_category sang outcome_strategy, script_template_id và next_action.",
            "Chương 3.5.4, Chương 4.6.5",
            "Unit test test_outcome_script_rules.py kiểm tra đủ sáu outcome, fallback IN_PROGRESS và không pitch tiếp với DO_NOT_CALL/HARD_REJECTION.",
        ],
    )
    append_row(
        appendix_table,
        [
            "project/bigquery/create_serving_views.sql",
            "Tạo view vw_customer_outcome_scripts phục vụ truy vấn kịch bản cùng ngữ cảnh khách hàng, offer và fact cuộc gọi.",
            "Chương 3.5.4, Chương 4.6.5",
            "BigQuery query trên view trả 23.447 dòng; join dùng ON tường minh để tránh lỗi ambiguous customer_id/offer_id.",
        ],
    )

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
