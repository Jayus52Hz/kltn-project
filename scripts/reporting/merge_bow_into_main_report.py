# -*- coding: utf-8 -*-
from pathlib import Path
from shutil import copyfile

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[2]
REPORTS_DIR = ROOT / "docs" / "reports"

MAIN_REPORT = REPORTS_DIR / "Report KLTN - 22133056 - Nguyen Quoc Thinh - revised model theory.docx"
CANONICAL_REPORT = REPORTS_DIR / "Report KLTN - 22133056 - Nguyen Quoc Thinh.docx"


def insert_after(paragraph, text="", style_name=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style_name:
        new_para.style = style_name
    if text:
        new_para.add_run(text)
    return new_para


def set_text(paragraph, text):
    paragraph.clear()
    paragraph.add_run(text)


def find_idx(doc, text):
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.style.name.startswith("Heading") and paragraph.text.strip() == text:
            return idx
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == text:
            return idx
    raise ValueError(f"Paragraph not found: {text}")


def replace_between(doc, start_text, end_text, items):
    start_idx = find_idx(doc, start_text)
    end_idx = find_idx(doc, end_text)
    start = doc.paragraphs[start_idx]
    end = doc.paragraphs[end_idx]

    removing = []
    cursor = start._p.getnext()
    while cursor is not None and cursor is not end._p:
        removing.append(cursor)
        cursor = cursor.getnext()

    for element in removing:
        element.getparent().remove(element)

    anchor = start
    for style_name, text in items:
        anchor = insert_after(anchor, text, style_name)


def replace_contains(doc, needle, replacement):
    count = 0
    for paragraph in doc.paragraphs:
        if needle in paragraph.text:
            set_text(paragraph, replacement)
            count += 1
    return count


def remove_duplicate_references(doc):
    seen = set()
    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        if not text.startswith("["):
            continue
        if text in seen:
            paragraph._p.getparent().remove(paragraph._p)
        else:
            seen.add(text)


doc = Document(MAIN_REPORT)

replace_contains(
    doc,
    "Mô hình NLP gồm baseline CountVectorizer + Logistic Regression và mô hình chính RoBERTa fine-tuned",
    "Mô hình NLP chính trong pipeline là Bag-of-Words kết hợp Logistic Regression theo chiến lược One-vs-Rest; RoBERTa fine-tuned được giữ làm baseline học sâu để so sánh chất lượng và phân tích trade-off vận hành.",
)

replace_contains(
    doc,
    "Tích hợp mô hình NLP chính RoBERTa fine-tuned để phân loại transcript hội thoại thành nhãn nghiệp vụ có thể phân tích; BoW được giữ làm baseline so sánh.",
    "Tích hợp mô hình NLP BoW + Logistic Regression để phân loại transcript hội thoại thành call_code nghiệp vụ; RoBERTa được giữ làm baseline học sâu để đối chiếu chất lượng.",
)

replace_contains(
    doc,
    "Call transcript là dữ liệu phi cấu trúc có giá trị cao vì phản ánh nhu cầu, lý do từ chối và mức độ quan tâm của khách hàng. Trong đề tài, RoBERTa fine-tuned là mô hình NLP chính",
    "Call transcript là dữ liệu phi cấu trúc có giá trị cao vì phản ánh nhu cầu, lý do từ chối và mức độ quan tâm của khách hàng. Trong đề tài, BoW + Logistic Regression là mô hình NLP chính trong pipeline production",
)

replace_between(
    doc,
    "1.14.1. Vai trò của mô hình NLP",
    "1.15. Mô hình Star Schema cho dashboard vận hành",
    [
        (
            "Normal",
            "Trong đề tài, mô hình NLP được dùng để biến transcript cuộc gọi thành tín hiệu phân tích có cấu trúc. Đầu vào của mô hình là call_transcript, đầu ra là danh sách call_code phản ánh nội dung, trạng thái và hành vi xuất hiện trong cuộc gọi. call_code gốc sinh trong giai đoạn tạo dữ liệu chỉ đóng vai trò nhãn huấn luyện có giám sát; sau khi pipeline vận hành, source of truth cho các tầng phân tích là call_code do mô hình dự đoán.",
        ),
        (
            "Normal",
            "Bài toán này là multi-label text classification. Một transcript có thể đồng thời chứa nhiều tín hiệu nghiệp vụ, ví dụ vừa có PRODUCT_PITCH, vừa có OBJECTION_HANDLING, vừa có FOLLOW_UP_EMAIL_REQUESTED. Vì vậy mô hình không chọn duy nhất một lớp như bài toán single-label, mà dự đoán độc lập nhiều nhãn trên cùng một mẫu.",
        ),
        (
            "Normal",
            "Về kiến trúc dữ liệu, bước NLP giúp chuyển dữ liệu phi cấu trúc thành trường phân tích có thể tổng hợp trong Lakehouse. Transcript chỉ cần tồn tại ở nơi phục vụ inference và kiểm tra kỹ thuật tại Silver; Gold và BigQuery không lưu transcript để giảm rủi ro PII, chỉ giữ call_code đã được mô hình hóa thành nhãn nghiệp vụ.",
        ),
        ("Heading 3", "1.14.2. Mô hình chính BoW + Logistic Regression"),
        (
            "Normal",
            "Mô hình chính của hệ thống là Bag-of-Words kết hợp Logistic Regression theo chiến lược One-vs-Rest. Toàn bộ mô hình được triển khai bằng scikit-learn: CountVectorizer tạo đặc trưng văn bản, MultiLabelBinarizer mã hóa tập call_code thành vector multi-hot, còn OneVsRestClassifier huấn luyện một Logistic Regression riêng cho từng nhãn.",
        ),
        (
            "Normal",
            "Trong artifact bow_model.pkl, mô hình được lưu dưới dạng dictionary gồm vectorizer, classifier và mlb. Cách đóng gói này giúp Silver job nạp một artifact duy nhất bằng joblib, biến transcript thành ma trận sparse rồi suy luận call_code bằng classifier. Đây là lựa chọn phù hợp với pipeline Spark batch vì nhẹ, dễ tái lập, không cần GPU và không yêu cầu runtime deep learning nặng.",
        ),
        ("Normal", "Đoạn mã 1.11. Nạp artifact BoW và sinh call_code trong Silver"),
        ("Report Code", "bundle = joblib.load(BOW_MODEL_PATH)"),
        ("Report Code", "X = bundle[\"vectorizer\"].transform(transcripts.fillna(\"\"))"),
        ("Report Code", "y_pred = bundle[\"classifier\"].predict(X)"),
        ("Report Code", "label_list = bundle[\"mlb\"].inverse_transform(y_pred)"),
        (
            "Normal",
            "Bag-of-Words biểu diễn văn bản bằng cách đếm sự xuất hiện của token hoặc n-gram. Với mỗi transcript, CountVectorizer chuẩn hóa chữ, tách token, tạo các unigram và bigram, sau đó ánh xạ transcript thành vector có số chiều bằng kích thước vocabulary. Trong mô hình của đề tài, vocabulary được giới hạn ở 5.000 đặc trưng, ngram_range=(1,2), stop_words='english'.",
        ),
        (
            "Normal",
            "Ưu điểm của BoW là tính minh bạch: trọng số của Logistic Regression có thể được hiểu như mức đóng góp của từng token/ngram vào xác suất xuất hiện một call_code. Với dữ liệu telesales tổng hợp có nhiều cụm từ nghiệp vụ lặp lại, cách biểu diễn n-gram giúp bắt được các mẫu như 'not interested', 'call back', 'send email' hoặc 'price too high' mà không cần mô hình ngữ cảnh sâu.",
        ),
        (
            "Normal",
            "Logistic Regression mô hình hóa xác suất một nhãn xuất hiện theo công thức sigmoid của tổ hợp tuyến tính giữa vector đặc trưng và trọng số. Vì bài toán có nhiều nhãn, One-vs-Rest tách bài toán thành nhiều bộ phân loại nhị phân độc lập: với mỗi call_code, nhãn đó là lớp dương, toàn bộ nhãn còn lại là lớp âm. Kết quả cuối cùng là tập các nhãn được dự đoán dương cho transcript.",
        ),
        ("Normal", "Đoạn mã 1.12. Cấu hình huấn luyện BoW trong notebook"),
        ("Report Code", "counter = CountVectorizer(max_features=5000, ngram_range=(1, 2), stop_words='english')"),
        ("Report Code", "base_lr = LogisticRegression(solver='liblinear', class_weight='balanced', random_state=42)"),
        ("Report Code", "counting_model = OneVsRestClassifier(base_lr)"),
        ("Report Code", "counting_model.fit(X_train_count, y_train_full)"),
        (
            "Normal",
            "Điểm yếu của BoW là không hiểu thứ tự dài và ngữ cảnh sâu như Transformer. Tuy nhiên, sau thực nghiệm full rebuild trên môi trường lakehouse CPU-only, tính ổn định và tốc độ của BoW quan trọng hơn mức cải thiện chưa đủ rõ rệt của RoBERTa. Do đó BoW được chọn làm mô hình chính để sinh call_code trong pipeline production.",
        ),
        ("Heading 3", "1.14.3. RoBERTa như baseline học sâu"),
        (
            "Normal",
            "RoBERTa vẫn được trình bày trong đề tài với vai trò baseline học sâu. Về nguồn gốc, RoBERTa là biến thể được tối ưu từ BERT, được công bố trong nghiên cứu RoBERTa: A Robustly Optimized BERT Pretraining Approach của Liu et al. Mô hình kế thừa kiến trúc Transformer encoder, sử dụng self-attention để tạo biểu diễn ngữ cảnh hóa cho từng token.",
        ),
        (
            "Normal",
            "Trong notebook, checkpoint nền được dùng là roberta-base thông qua thư viện Hugging Face Transformers. AutoTokenizer chịu trách nhiệm token hóa transcript, AutoModelForSequenceClassification nạp trọng số pretrained và gắn classification head cho bài toán multi-label. Đầu ra là vector logits 32 chiều, được đưa qua sigmoid để quyết định từng call_code độc lập.",
        ),
        ("Normal", "Đoạn mã 1.13. Cấu hình RoBERTa baseline"),
        ("Report Code", "model = AutoModelForSequenceClassification.from_pretrained("),
        ("Report Code", "    \"roberta-base\","),
        ("Report Code", "    num_labels=num_labels,"),
        ("Report Code", "    problem_type=\"multi_label_classification\","),
        ("Report Code", ")"),
        (
            "Normal",
            "RoBERTa có lợi thế về biểu diễn ngữ cảnh, đặc biệt với những câu có phủ định, chuyển ý hoặc phụ thuộc dài. Tuy nhiên trong môi trường vận hành của đề tài, checkpoint roberta_saved lớn hơn nhiều so với artifact BoW, inference cần PyTorch và thời gian xử lý dài hơn đáng kể. Vì vậy RoBERTa phù hợp để làm baseline chất lượng và hướng phát triển sau này hơn là mô hình production mặc định.",
        ),
        ("Heading 3", "1.14.4. Tiêu chí đánh giá và quyết định mô hình"),
        (
            "Normal",
            "Bài toán được đánh giá bằng F1 Micro, F1 Macro, Precision, Recall và Exact Match. Trên test set, BoW đạt F1 Micro 70,16%, F1 Macro 62,08%, Precision 65,92%, Recall 74,98% và Exact Match 17,89%. RoBERTa đạt F1 Micro 73,46%, F1 Macro 56,14%, Precision 80,39%, Recall 67,62% và Exact Match 17,71%. RoBERTa nhỉnh hơn ở F1 Micro và Precision, nhưng Macro F1 và Exact Match không tốt hơn rõ rệt.",
        ),
        (
            "Normal",
            "Khi xét cùng chi phí vận hành, mức cải thiện của RoBERTa không đủ thuyết phục để trở thành mô hình chính trong pipeline. BoW chạy nhanh, ổn định, artifact nhỏ, dễ debug và phù hợp hơn với mục tiêu của prototype là chứng minh luồng Hybrid Data Lakehouse end-to-end. Vì vậy quyết định cuối cùng là BoW làm mô hình production, RoBERTa làm baseline học sâu.",
        ),
        ("Normal", "Đoạn mã 1.14. Cấu hình mô hình production trong pipeline"),
        ("Report Code", "NLP_MODEL_TYPE = os.getenv(\"NLP_MODEL_TYPE\", \"bow\").lower()"),
        ("Report Code", "BOW_MODEL_PATH = os.path.join(MODELS_PATH, \"bow_model.pkl\")"),
        ("Report Code", "ROBERTA_PATH = os.path.join(MODELS_PATH, \"roberta_saved\")  # experimental baseline"),
    ],
)

replace_between(
    doc,
    "2.6.1. Mục tiêu mô hình",
    "2.7. Data quality, đạo đức dữ liệu và giới hạn của dữ liệu tổng hợp",
    [
        (
            "Normal",
            "Mục tiêu của mô hình NLP là dự đoán call_code từ call_transcript để biến hội thoại phi cấu trúc thành tín hiệu phân tích. call_code ban đầu trong dữ liệu tổng hợp chỉ dùng làm nhãn học có giám sát khi huấn luyện; sau huấn luyện, pipeline không sử dụng call_code gốc cho Gold hoặc BigQuery. Trường call_code ở Silver/Gold là kết quả inference từ mô hình.",
        ),
        (
            "Normal",
            "Thiết kế này phản ánh đúng bối cảnh vận hành: hệ thống nhận transcript mới, mô hình đọc transcript và sinh nhãn nghiệp vụ, sau đó các bước phân tích dùng nhãn đã suy luận thay vì phụ thuộc vào nhãn có sẵn trong dữ liệu huấn luyện.",
        ),
        ("Heading 3", "2.6.2. Mô hình chính BoW + Logistic Regression"),
        (
            "Normal",
            "BoW + Logistic Regression là mô hình chính được triển khai trong Silver job. Mô hình dùng CountVectorizer với max_features=5000, ngram_range=(1,2) và stop_words='english'. Sau khi transcript được chuyển thành vector sparse, OneVsRestClassifier huấn luyện LogisticRegression cho từng call_code. class_weight='balanced' được dùng để giảm ảnh hưởng của mất cân bằng nhãn.",
        ),
        (
            "Normal",
            "Artifact bow_model.pkl gồm vectorizer, classifier và mlb. vectorizer chịu trách nhiệm biến transcript thành đặc trưng, classifier dự đoán vector multi-label, còn mlb chuyển vector dự đoán về danh sách tên call_code. Cấu trúc này giúp pipeline chỉ cần nạp một file model duy nhất trong Spark.",
        ),
        ("Report Code", "bow_broadcast = spark.sparkContext.broadcast(joblib.load(BOW_MODEL_PATH))"),
        ("Report Code", "@F.pandas_udf(ArrayType(StringType()))"),
        ("Report Code", "def predict_call_codes(transcripts: pd.Series) -> pd.Series:"),
        ("Report Code", "    bundle = bow_broadcast.value"),
        ("Report Code", "    X = bundle[\"vectorizer\"].transform(transcripts.fillna(\"\"))"),
        ("Report Code", "    y_pred = bundle[\"classifier\"].predict(X)"),
        ("Report Code", "    return pd.Series([list(labels) for labels in bundle[\"mlb\"].inverse_transform(y_pred)])"),
        ("Heading 3", "2.6.3. RoBERTa baseline"),
        (
            "Normal",
            "RoBERTa được giữ làm baseline học sâu để so sánh. Notebook sử dụng roberta-base từ Hugging Face, token hóa transcript với max_length=512, cấu hình AutoModelForSequenceClassification với problem_type='multi_label_classification', sau đó fine-tune bằng PyTorch/Trainer. Mỗi nhãn được xử lý như một bài toán nhị phân độc lập thông qua sigmoid và BCEWithLogitsLoss.",
        ),
        (
            "Normal",
            "Về lý thuyết, RoBERTa có khả năng biểu diễn ngữ cảnh tốt hơn BoW nhờ self-attention. Mô hình có thể khai thác quan hệ giữa các token trong toàn bộ chuỗi, phù hợp với các tình huống hội thoại có phủ định, chuyển ý hoặc ý nghĩa phụ thuộc vào nhiều câu. Đây là lý do RoBERTa được đưa vào notebook như baseline học sâu nghiêm túc, không chỉ là mô hình tham khảo đơn giản.",
        ),
        ("Report Code", "encoded = tokenizer(batch[\"call_transcript\"], truncation=True, padding=\"max_length\", max_length=512)"),
        ("Report Code", "loss = torch.nn.BCEWithLogitsLoss()(outputs.logits, labels.float())"),
        ("Report Code", "probs = torch.sigmoid(torch.tensor(logits)).numpy()"),
        ("Report Code", "predicted_labels = mlb.inverse_transform((probs >= 0.5).astype(int))"),
        ("Report Caption", "Hình 2.6. So sánh BoW production model và RoBERTa baseline trong bài toán multi-label call_code"),
        ("Heading 3", "2.6.4. Kết quả notebook và quyết định ban đầu"),
        (
            "Normal",
            "Trong notebook, BoW đạt F1 Micro 70,16%, Macro F1 62,08%, Precision 65,92%, Recall 74,98% và Exact Match 17,89%. RoBERTa đạt F1 Micro 73,46%, Macro F1 56,14%, Precision 80,39%, Recall 67,62% và Exact Match 17,71%. Nếu chỉ nhìn Precision và F1 Micro, RoBERTa có vẻ hấp dẫn hơn. Tuy nhiên Macro F1 và Exact Match không cải thiện rõ, còn Recall thấp hơn BoW.",
        ),
        (
            "Normal",
            "Sau khi đưa vào thực nghiệm end-to-end, quyết định mô hình được điều chỉnh theo tiêu chí vận hành. Pipeline lakehouse cần một mô hình đủ tốt, ổn định, có thể chạy lại toàn bộ hệ thống trong thời gian hợp lý. Với tiêu chí đó, BoW phù hợp hơn RoBERTa cho bản triển khai hiện tại.",
        ),
        ("Heading 3", "2.6.5. Tích hợp BoW vào Silver job"),
        (
            "Normal",
            "Repository đã được đồng bộ theo quyết định này: silver_job.py mặc định đọc NLP_MODEL_TYPE=bow, nạp bow_model.pkl và sinh call_code trực tiếp từ call_transcript bằng Pandas UDF. RoBERTa vẫn có thể bật lại bằng NLP_MODEL_TYPE=roberta để phục vụ thí nghiệm hoặc so sánh, nhưng không còn là cấu hình production.",
        ),
        ("Report Code", "NLP_MODEL_TYPE = os.getenv(\"NLP_MODEL_TYPE\", \"bow\").lower()"),
        ("Report Code", "if NLP_MODEL_TYPE == \"bow\":"),
        ("Report Code", "    print(f\"Using BoW production model from {BOW_MODEL_PATH} ...\")"),
        ("Report Code", "elif NLP_MODEL_TYPE == \"roberta\":"),
        ("Report Code", "    print(f\"Using RoBERTa experimental baseline from {ROBERTA_PATH} ...\")"),
        ("Report Caption", "Hình 2.7. Luồng BoW inference ghi call_code vào Silver"),
        ("Heading 3", "2.6.6. Ranh giới dữ liệu sau inference"),
        (
            "Normal",
            "Sau inference, Silver giữ call_transcript để phục vụ kiểm tra kỹ thuật và khả năng tái xử lý cục bộ. Gold và BigQuery không nhận transcript, không nhận call_code_original và không nhận call_code_predicted. Các tầng phân tích chỉ giữ call_code do mô hình sinh ra, cùng các cờ nghiệp vụ được suy ra từ call_code như has_successful_sale, has_hard_rejection, has_soft_rejection, has_do_not_call và has_objection.",
        ),
    ],
)

replace_between(
    doc,
    "4.5.1. Kết quả mô hình trong notebook",
    "4.6. Kiểm thử khả năng chạy lại và phục hồi lỗi",
    [
        (
            "Normal",
            "Notebook NLP_model.ipynb cho thấy BoW đạt F1 Micro 70,16%, Macro F1 62,08%, Precision 65,92%, Recall 74,98% và Exact Match 17,89%. RoBERTa đạt F1 Micro 73,46%, Macro F1 56,14%, Precision 80,39%, Recall 67,62% và Exact Match 17,71%. Chênh lệch F1 Micro giữa hai mô hình là 3,30 điểm phần trăm, còn Exact Match của RoBERTa thấp hơn nhẹ so với BoW.",
        ),
        (
            "Normal",
            "Kết quả notebook cho thấy RoBERTa có Precision cao hơn, nhưng không tạo ra cải thiện toàn diện. Macro F1 của RoBERTa thấp hơn BoW, Recall thấp hơn BoW, và Exact Match gần như tương đương. Vì vậy cần đánh giá thêm chi phí vận hành trong pipeline lakehouse thay vì chỉ chọn mô hình dựa trên một chỉ số đơn lẻ.",
        ),
        ("Heading 3", "4.5.2. Thực nghiệm RoBERTa trong full rebuild"),
        (
            "Normal",
            "Trong full rebuild ngày 01/06/2026, RoBERTa được thử nghiệm trực tiếp trong Silver job trên 23.447 bản ghi call_logs. Lần thử đầu dùng batch_size=16 và một Spark task inference; job chạy từ 05:53:48 đến 06:10:07 nhưng vẫn chưa vượt qua Silver. Spark UI cho thấy một task active kéo dài trên toàn bộ tập call_logs, phản ánh nút thắt inference CPU.",
        ),
        (
            "Normal",
            "Lần thử thứ hai tăng song song hóa bằng ROBERTA_NUM_PARTITIONS=4, ROBERTA_BATCH_SIZE=16 và TORCH_NUM_THREADS=1. CPU tăng lên khoảng 380%, nhưng Python workers bị crash với lỗi Python worker exited unexpectedly, EOFException và spark-submit exit -9. Điều này cho thấy tăng partition giúp dùng CPU tốt hơn nhưng làm áp lực bộ nhớ/runtime PyTorch cao hơn, khiến pipeline không ổn định.",
        ),
        (
            "Normal",
            "Lần thử thứ ba giảm xuống ROBERTA_NUM_PARTITIONS=2, ROBERTA_BATCH_SIZE=8 và TORCH_NUM_THREADS=1. Cấu hình này ổn định hơn nhưng vẫn chưa hoàn tất Silver đến 07:24:59, nên quá lâu so với nhu cầu rebuild và kiểm thử end-to-end. Với kết quả notebook không vượt trội rõ rệt, chi phí vận hành này không phù hợp cho mô hình production của prototype.",
        ),
        ("Report Code", "ROBERTA_BATCH_SIZE = int(os.getenv(\"ROBERTA_BATCH_SIZE\", \"8\"))"),
        ("Report Code", "ROBERTA_NUM_PARTITIONS = int(os.getenv(\"ROBERTA_NUM_PARTITIONS\", \"2\"))"),
        ("Report Code", "torch.set_num_threads(int(os.getenv(\"TORCH_NUM_THREADS\", \"1\")))"),
        ("Heading 3", "4.5.3. Thực nghiệm BoW sau khi thay mô hình"),
        (
            "Normal",
            "Sau khi chuyển NLP_MODEL_TYPE về bow, Silver chạy từ 07:27:32 đến 07:28:12 và xử lý đủ 23.447 bản ghi call_logs. Gold tiếp tục chạy thành công từ 07:28:14 đến 07:28:42. BigQuery sync chạy thành công từ 07:32:36 đến 07:33:15, đồng bộ dim_customer 4.344 dòng, dim_offer 5.072 dòng, dim_date 2 dòng và fact_telesales_calls 23.447 dòng.",
        ),
        (
            "Normal",
            "Kết quả này cho thấy BoW đáp ứng tốt hơn yêu cầu vận hành của pipeline: inference nhanh, ít phụ thuộc tài nguyên, dễ chạy lại và phù hợp với mục tiêu chứng minh kiến trúc Hybrid Data Lakehouse end-to-end. RoBERTa được giữ lại như baseline học sâu và hướng nâng cấp khi có model serving/GPU hoặc nhu cầu chất lượng cao hơn.",
        ),
        ("Report Code", "Using BoW production model from /opt/spark/work-dir/batch-etl/models/bow_model.pkl"),
        ("Report Code", "MERGE INTO lakehouse.silver.call_logs completed (23,447 source records)"),
        ("Report Code", "Silver job completed successfully."),
        ("Heading 3", "4.5.4. Kiểm tra schema sau khi đổi mô hình"),
        (
            "Normal",
            "Sau full rebuild, Silver call_logs có call_code là nhãn do mô hình sinh và vẫn giữ call_transcript để phục vụ inference nội bộ. Gold fact_telesales_calls có call_code và các cờ nghiệp vụ, nhưng không có call_transcript, call_code_original hoặc call_code_predicted. BigQuery sync cũng loại các cột nhạy cảm trước khi ghi, bảo đảm transcript không đi vào lớp serving phân tích.",
        ),
        ("Report Code", "BLOCKED_BIGQUERY_COLUMNS = {\"call_transcript\", \"call_code_original\", \"call_code_predicted\", \"full_name\", \"address\"}"),
        ("Report Code", "df = df.drop(*[col for col in BLOCKED_BIGQUERY_COLUMNS if col in df.columns])"),
    ],
)

replace_contains(
    doc,
    "Mức độ tích hợp NLP được định hướng theo RoBERTa fine-tuned để nâng chất lượng dự đoán call_code. BoW chỉ là baseline chứng minh mức cải thiện và có thể dùng trong demo nhẹ.",
    "Mức độ tích hợp NLP được định hướng theo BoW + Logistic Regression để bảo đảm pipeline chạy nhanh và ổn định end-to-end. RoBERTa được dùng làm baseline học sâu để chứng minh trade-off giữa chất lượng mô hình và chi phí vận hành.",
)

replace_contains(
    doc,
    "BoW triển khai nhẹ nhưng không hiểu sâu ngữ cảnh, sarcasm hoặc các quan hệ dài trong hội thoại. RoBERTa tốt hơn về F1 Micro và Precision nên được chọn làm mô hình chính, nhưng cần",
    "BoW triển khai nhẹ, ổn định và phù hợp với pipeline hiện tại, nhưng chưa hiểu sâu ngữ cảnh, sắc thái hoặc quan hệ dài trong hội thoại. RoBERTa tốt hơn về Precision nhưng quá nặng trong full rebuild CPU-only, nên được giữ làm baseline và hướng phát triển sau này.",
)

replace_contains(
    doc,
    "Hướng thứ hai là nâng cấp NLP. RoBERTa hoặc một transformer nhỏ hơn có thể được triển khai thành service riêng",
    "Hướng thứ hai là nâng cấp NLP. RoBERTa hoặc một transformer nhỏ hơn có thể được triển khai thành service riêng khi có hạ tầng model serving phù hợp",
)

replace_contains(
    doc,
    '        F.array_contains(F.col("call_code_original"), "SUCCESSFUL_SALE").alias("has_successful_sale"),',
    '        F.array_contains(F.col("call_code"), "SUCCESSFUL_SALE").alias("has_successful_sale"),',
)

replace_contains(
    doc,
    '        F.when(F.array_contains(F.col("call_code_original"), "SUCCESSFUL_SALE"), "SALE")',
    '        F.when(F.array_contains(F.col("call_code"), "SUCCESSFUL_SALE"), "SALE")',
)

replace_contains(
    doc,
    '         .when(F.array_contains(F.col("call_code_original"), "DO_NOT_CALL_REQUEST"), "DO_NOT_CALL")',
    '         .when(F.array_contains(F.col("call_code"), "DO_NOT_CALL_REQUEST"), "DO_NOT_CALL")',
)

replace_contains(
    doc,
    "f.call_code_original, f.call_code,",
    "f.call_code,",
)

replace_contains(
    doc,
    "Outcome priority là một quyết định thiết kế nghiệp vụ. Nếu call_code_original chứa SUCCESSFUL_SALE thì outcome_category là SALE. Nếu có DO_NOT_CALL_REQUEST thì là DO_NOT_CALL. Sau đó mới đến HARD_REJECTION, SOFT_REJECTION, WARM_LEAD và cuối cùng IN_PROGRESS. Cách ưu tiên này tránh một cuộc gọi có nhiều nhãn bị đếm sai khi có đồng thời objection và kết quả cuối.",
    "Outcome priority là một quyết định thiết kế nghiệp vụ dựa trên call_code do mô hình sinh. Nếu call_code chứa SUCCESSFUL_SALE thì outcome_category là SALE. Nếu có DO_NOT_CALL_REQUEST thì là DO_NOT_CALL. Sau đó mới đến HARD_REJECTION, SOFT_REJECTION, WARM_LEAD và cuối cùng IN_PROGRESS. Cách ưu tiên này tránh một cuộc gọi có nhiều nhãn bị đếm sai khi có đồng thời objection và kết quả cuối.",
)

replace_contains(
    doc,
    "Silver job có nhiều điểm cần kiểm thử hơn Bronze. Với cust, cần kiểm tra phone_number_masked và national_id_masked đã che đúng. Với offer, cần kiểm tra product_name, lead_source, loan_amount và interest_rate parse đúng kiểu. Với call_logs, cần kiểm tra call_timestamp được convert sang timestamp, call_code_original là array và call_code được mô hình RoBERTa trả về.",
    "Silver job có nhiều điểm cần kiểm thử hơn Bronze. Với cust, cần kiểm tra phone_number_masked và national_id_masked đã che đúng. Với offer, cần kiểm tra product_name, lead_source, loan_amount và interest_rate parse đúng kiểu. Với call_logs, cần kiểm tra call_timestamp được convert sang timestamp, call_transcript chỉ dùng ở Silver và call_code được mô hình BoW trả về.",
)

replace_contains(
    doc,
    "Đối với NLP inference, không nên chỉ kiểm tra cột tồn tại. Cần lấy một vài transcript mẫu và so sánh call_code_original với call_code để diễn giải. Nếu predicted thiếu hoặc thừa nhãn, có thể liên hệ lại với precision/recall của mô hình. Điều này giúp phần NLP không bị trình bày như black box.",
    "Đối với NLP inference, không nên chỉ kiểm tra cột tồn tại. Cần lấy một vài transcript mẫu, đọc call_code do mô hình sinh và diễn giải vì sao nhãn đó hợp lý theo nội dung hội thoại. Nếu mô hình thiếu hoặc thừa nhãn, có thể liên hệ lại với precision/recall của mô hình. Điều này giúp phần NLP không bị trình bày như black box.",
)

replace_contains(
    doc,
    "Do đó kết luận mô hình của đề tài là: RoBERTa là mô hình chính; BoW là baseline và fallback nhẹ.",
    "Do đó kết luận mô hình của đề tài là: BoW + Logistic Regression là mô hình production; RoBERTa là baseline học sâu và hướng nâng cấp.",
)

replace_contains(
    doc,
    "Silver sử dụng RoBERTa khi mục tiêu là bám đúng thiết kế mô hình của đề tài.",
    "Silver sử dụng BoW trong cấu hình production vì đáp ứng tốt hơn yêu cầu rebuild nhanh, ổn định và ít phụ thuộc tài nguyên.",
)

for paragraph in doc.paragraphs:
    text = paragraph.text
    replacements = {
        "mô hình chính RoBERTa fine-tuned": "mô hình chính BoW + Logistic Regression",
        "RoBERTa là mô hình chính": "BoW + Logistic Regression là mô hình chính",
        "BoW chỉ là baseline": "RoBERTa là baseline học sâu",
        "NLP_MODEL_TYPE=roberta": "NLP_MODEL_TYPE=bow",
        "NLP_MODEL_TYPE = os.getenv(\"NLP_MODEL_TYPE\", \"roberta\").lower()": "NLP_MODEL_TYPE = os.getenv(\"NLP_MODEL_TYPE\", \"bow\").lower()",
        "Hình 2.7. Cấu hình huấn luyện RoBERTa và luồng đánh giá train/valid/test": "Hình 2.7. Luồng BoW inference ghi call_code vào Silver",
        "Hình 2.8. Luồng RoBERTa inference ghi call_code vào Silver": "Hình 2.8. RoBERTa baseline và cấu hình thí nghiệm so sánh",
        "Hình 2.6. So sánh baseline BoW và RoBERTa trong bài toán multi-label call_code": "Hình 2.6. So sánh BoW production model và RoBERTa baseline trong bài toán multi-label call_code",
    }
    new_text = text
    for old, new in replacements.items():
        new_text = new_text.replace(old, new)
    if new_text != text:
        set_text(paragraph, new_text)

references = [
    "[23] scikit-learn, CountVectorizer documentation, https://scikit-learn.org/stable/modules/generated/sklearn.feature_extraction.text.CountVectorizer.html",
    "[24] scikit-learn, LogisticRegression documentation, https://scikit-learn.org/stable/modules/generated/sklearn.linear_model.LogisticRegression.html",
    "[25] scikit-learn, OneVsRestClassifier documentation, https://scikit-learn.org/stable/modules/generated/sklearn.multiclass.OneVsRestClassifier.html",
    "[26] scikit-learn, MultiLabelBinarizer documentation, https://scikit-learn.org/stable/modules/generated/sklearn.preprocessing.MultiLabelBinarizer.html",
    "[27] Hugging Face Transformers, RoBERTa documentation, https://huggingface.co/docs/transformers/main/model_doc/roberta",
    "[28] PyTorch, BCEWithLogitsLoss documentation, https://pytorch.org/docs/stable/generated/torch.nn.BCEWithLogitsLoss.html",
]

ref_anchor = None
for p in doc.paragraphs:
    if p.text.strip().startswith("[22]"):
        ref_anchor = p
        break

if ref_anchor is not None:
    anchor = ref_anchor
    existing = {p.text.strip() for p in doc.paragraphs}
    for ref in references:
        if ref not in existing:
            anchor = insert_after(anchor, ref, "Normal")

remove_duplicate_references(doc)

doc.save(MAIN_REPORT)
copyfile(MAIN_REPORT, CANONICAL_REPORT)

print(f"Updated {MAIN_REPORT}")
print(f"Updated {CANONICAL_REPORT}")
